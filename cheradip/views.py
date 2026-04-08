from django.shortcuts import render
from rest_framework import generics, status, viewsets, filters
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from .models import (
    Item,
    Customer,
    CustomerToken,
    CreatedQuestionSet,
    PendingQuestion,
    OrderDetail,
    Transaction,
    Notification,
    Country,
    Location,
    JsonData,
)
from .serializers import (
    ItemSerializer,
    CustomerSignupSerializer,
    CustomerUpdateSerializer,
    CustomerSerializer,
    NotificationSerializer,
    CountrySerializer,
    CountryListSerializer,
)
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.authentication import BaseAuthentication
from .permissions import IsSuperUserOrStaff, PublicAccess
from .location import Bangladesh
from django.http import Http404, HttpResponse, HttpResponseRedirect, JsonResponse
from io import BytesIO
from django.contrib.auth import authenticate
from django.contrib.auth.hashers import check_password
from django.views.decorators.csrf import csrf_exempt
import logging, random, string, json, requests, os, re, csv, time, zipfile
from html import escape
from urllib import parse as urllib_parse
from urllib.parse import quote
from pathlib import Path
from rest_framework.decorators import action
from django.conf import settings
from django.db import connections
from django.db.models import Q
from django.db.models.expressions import RawSQL
from django.db.utils import ProgrammingError, OperationalError

from .subject_question_tables import subject_question_table_name, next_qid_for_chapter_topic

try:
    from reportlab.lib.pagesizes import A4, A3, A5, letter, legal
    from reportlab.lib.units import mm as rl_mm
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    pdfmetrics = None
    TTFont = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Mm as DocxMm, Pt as DocxPt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


# ==============================================================================
# COUNTRY VIEWSET
# ==============================================================================

class CountryViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet for Country model with search, autocomplete, and geolocation detection.
    
    Endpoints:
    - GET /api/countries/ - List all countries (lightweight, unpaginated)
    - GET /api/countries/{country_code}/ - Get single country (full details)
    - GET /api/countries/?search=bang - Search countries
    - GET /api/countries/?featured=true - Get featured countries
    - GET /api/countries/detect/ - Detect country from IP
    """
    queryset = Country.objects.filter(is_active=True)
    serializer_class = CountrySerializer
    lookup_field = 'country_code'
    pagination_class = None  # Return full list for dropdowns/options
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['country_name', 'country_name_native', 'country_code', 'phone_code']
    ordering_fields = ['display_order', 'country_name']
    ordering = ['display_order', 'country_name']
    
    def get_serializer_class(self):
        """Use lightweight serializer for list, full for detail"""
        if self.action == 'list':
            return CountryListSerializer
        return CountrySerializer
    
    def get_queryset(self):
        queryset = super().get_queryset()
        # Avoid loading datetime columns (MySQL/PyMySQL can return them as str → 'utcoffset' error)
        queryset = queryset.defer('created_at', 'updated_at')

        # Filter by featured
        featured = self.request.query_params.get('featured')
        if featured and featured.lower() == 'true':
            queryset = queryset.filter(is_featured=True)
        
        # Filter by continent
        continent = self.request.query_params.get('continent')
        if continent:
            queryset = queryset.filter(continent__iexact=continent)
        
        # Custom search parameter
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(
                Q(country_name__icontains=search) |
                Q(country_name_native__icontains=search) |
                Q(country_code__iexact=search) |
                Q(phone_code__icontains=search)
            )
        
        return queryset
    
    @action(detail=False, methods=['get'])
    def detect(self, request):
        """
        Detect user's country based on IP address.
        Uses multiple fallback services for reliability.
        """
        # Get client IP
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0].strip()
        else:
            ip = request.META.get('REMOTE_ADDR')
        
        # Skip detection for localhost
        if ip in ['127.0.0.1', 'localhost', '::1']:
            # Return Bangladesh as default for localhost
            try:
                country = Country.objects.defer('created_at', 'updated_at').get(country_code='BD')
                serializer = CountrySerializer(country)
                return Response({
                    'detected': True,
                    'ip': ip,
                    'country': serializer.data,
                    'source': 'default'
                })
            except Country.DoesNotExist:
                return Response({
                    'detected': False,
                    'ip': ip,
                    'error': 'Default country not found'
                })
        
        # Try multiple geolocation services
        country_code = None
        source = None
        
        # Service 1: ipapi.co
        try:
            response = requests.get(f'https://ipapi.co/{ip}/country/', timeout=5)
            if response.status_code == 200:
                country_code = response.text.strip().upper()
                source = 'ipapi.co'
        except:
            pass
        
        # Service 2: ip-api.com (fallback)
        if not country_code:
            try:
                response = requests.get(f'http://ip-api.com/json/{ip}?fields=countryCode', timeout=5)
                if response.status_code == 200:
                    data = response.json()
                    country_code = data.get('countryCode', '').upper()
                    source = 'ip-api.com'
            except:
                pass
        
        # Return detected country
        if country_code:
            try:
                country = Country.objects.defer('created_at', 'updated_at').get(country_code=country_code)
                serializer = CountrySerializer(country)
                return Response({
                    'detected': True,
                    'ip': ip,
                    'country': serializer.data,
                    'source': source
                })
            except Country.DoesNotExist:
                return Response({
                    'detected': False,
                    'ip': ip,
                    'country_code': country_code,
                    'error': 'Country not in database'
                })
        
        # Fallback to default (Bangladesh)
        try:
            country = Country.objects.defer('created_at', 'updated_at').get(country_code='BD')
            serializer = CountrySerializer(country)
            return Response({
                'detected': False,
                'ip': ip,
                'country': serializer.data,
                'source': 'default',
                'error': 'Could not detect country'
            })
        except Country.DoesNotExist:
            return Response({
                'detected': False,
                'ip': ip,
                'error': 'Detection failed and no default country'
            })
    
    @action(detail=False, methods=['get'])
    def featured(self, request):
        """Get featured countries for quick selection"""
        countries = Country.objects.filter(is_active=True, is_featured=True).order_by('display_order').defer('created_at', 'updated_at')
        serializer = CountryListSerializer(countries, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def by_continent(self, request):
        """Get countries grouped by continent"""
        continents = Country.objects.filter(is_active=True).values_list('continent', flat=True).distinct()
        result = {}
        for continent in continents:
            if continent:
                countries = Country.objects.filter(is_active=True, continent=continent).order_by('country_name').defer('created_at', 'updated_at')
                result[continent] = CountryListSerializer(countries, many=True).data
        return Response(result)


class AllCountriesView(APIView):
    """GET /api/country/ - Return all active countries as array (for <option> dropdowns)."""
    def get(self, request):
        countries = Country.objects.filter(is_active=True).order_by('display_order', 'country_name').defer('created_at', 'updated_at')
        serializer = CountryListSerializer(countries, many=True)
        return Response(serializer.data)


class LocationDivisionsView(APIView):
    """
    GET /api/locations/divisions/?country_code=BD
    Returns distinct division names from cheradip_location for the given country.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        if not country_code:
            return Response([], status=status.HTTP_200_OK)
        qs = Location.objects.filter(country_id=country_code).exclude(
            division__isnull=True
        ).exclude(division='').values_list('division', flat=True).distinct().order_by('division')
        return Response(list(qs))


class LocationDistrictsView(APIView):
    """
    GET /api/locations/districts/?country_code=BD&division=Dhaka
    Returns distinct district names from cheradip_location for the given country and division.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        division = (request.query_params.get('division') or '').strip()
        if not country_code or not division:
            return Response([], status=status.HTTP_200_OK)
        qs = Location.objects.filter(
            country_id=country_code, division=division
        ).exclude(district__isnull=True).exclude(district='').values_list(
            'district', flat=True
        ).distinct().order_by('district')
        return Response(list(qs))


class LocationThanasView(APIView):
    """
    GET /api/locations/thanas/?country_code=BD&division=Dhaka&district=Dhaka
    Returns distinct thana names from cheradip_location for the given country, division, district.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        division = (request.query_params.get('division') or '').strip()
        district = (request.query_params.get('district') or '').strip()
        if not country_code or not division or not district:
            return Response([], status=status.HTTP_200_OK)
        qs = Location.objects.filter(
            country_id=country_code, division=division, district=district
        ).exclude(thana__isnull=True).exclude(thana='').values_list(
            'thana', flat=True
        ).distinct().order_by('thana')
        return Response(list(qs))


# ==============================================================================
# JOB DB VIEWSETS (cheradip_job – NTRCA merit/vacancy/recommend, institutes, banbeis, token)
# ==============================================================================

from .models import (
    Merit5, Merit6, Merit7,
    Vacancy5, Vacancy6, Vacancy7,
    Recommend5, Recommend6, Recommend7,
    Banbeis, Institutes, Token,
)
from .serializers import (
    Merit5Serializer, Merit6Serializer, Merit7Serializer,
    Vacancy5Serializer, Vacancy6Serializer, Vacancy7Serializer,
    Recommend5Serializer, Recommend6Serializer, Recommend7Serializer,
    BanbeisSerializer, InstitutesSerializer, TokenSerializer,
)
from rest_framework.pagination import PageNumberPagination


class JobListPagination(PageNumberPagination):
    page_size = 100
    page_size_query_param = 'page_size'
    max_page_size = 500


class _MeritViewSetMixin:
    """Filter merit list by query param code= (designation code)."""
    pagination_class = JobListPagination
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = super().get_queryset()
        code = self.request.query_params.get('code')
        if code is not None and code != '':
            qs = qs.filter(Code=code)
        return qs.order_by('SL')

    @action(detail=False, url_path='total_table_count')
    def total_table_count(self, request):
        """GET /api/merit5|6|7/total_table_count/ – total rows in table (no filters)."""
        count = self.queryset.count()
        return Response({'count': count})


class Merit5ViewSet(_MeritViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Merit5.objects.all()
    serializer_class = Merit5Serializer


class Merit6ViewSet(_MeritViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Merit6.objects.all()
    serializer_class = Merit6Serializer


class Merit7ViewSet(_MeritViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Merit7.objects.all()
    serializer_class = Merit7Serializer


class _VacancyViewSetMixin:
    pagination_class = JobListPagination
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = super().get_queryset()
        designation = self.request.query_params.get('designation')
        if designation:
            qs = qs.filter(Designation__iexact=designation)
        subject = self.request.query_params.get('subject')
        if subject:
            qs = qs.filter(Subject__iexact=subject)
        districts = self.request.query_params.getlist('district')
        if districts:
            qs = qs.filter(District__in=districts)
        return qs

    @action(detail=False, url_path='total_table_count')
    def total_table_count(self, request):
        """GET /api/vacancy5|6|7/total_table_count/ – total rows in table (no filters)."""
        count = self.queryset.count()
        return Response({'count': count})


class Vacancy5ViewSet(_VacancyViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Vacancy5.objects.all()
    serializer_class = Vacancy5Serializer


class Vacancy6ViewSet(_VacancyViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Vacancy6.objects.all()
    serializer_class = Vacancy6Serializer


class Vacancy7ViewSet(_VacancyViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Vacancy7.objects.all()
    serializer_class = Vacancy7Serializer


class _RecommendViewSetMixin:
    pagination_class = JobListPagination
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = super().get_queryset()
        code = self.request.query_params.get('code')
        if code is not None and str(code).strip() != '':
            try:
                qs = qs.filter(Code=int(str(code).strip()))
            except (ValueError, TypeError):
                pass
        districts = self.request.query_params.getlist('district')
        if districts:
            qs = qs.filter(District__in=districts)
        thanas = self.request.query_params.getlist('thana')
        if thanas:
            qs = qs.filter(Thana__in=thanas)
        return qs

    @action(detail=False, url_path='unique_districts')
    def unique_districts(self, request):
        qs = self.get_queryset()
        code = request.query_params.get('code')
        if code is not None and str(code).strip() != '':
            try:
                qs = qs.filter(Code=int(str(code).strip()))
            except (ValueError, TypeError):
                pass
        districts = qs.values_list('District', flat=True).distinct().order_by('District')
        districts = [d for d in districts if d]
        return Response(list(districts))

    @action(detail=False, url_path='unique_thanas')
    def unique_thanas(self, request):
        """GET /api/recommend5|6|7/unique_thanas/?district=DHAKA&district=KISHOREGANJ&code=201 – thanas under selected districts (optional code)."""
        qs = self.get_queryset()
        code = request.query_params.get('code')
        if code is not None and str(code).strip() != '':
            try:
                qs = qs.filter(Code=int(str(code).strip()))
            except (ValueError, TypeError):
                pass
        district_list = request.query_params.getlist('district')
        district_list = [d.strip() for d in district_list if d and str(d).strip()]
        if district_list:
            qs = qs.filter(District__in=district_list)
        thanas = qs.values_list('Thana', flat=True).distinct().order_by('Thana')
        thanas = [t for t in thanas if t]
        return Response(list(thanas))

    @action(detail=False, url_path='total_table_count')
    def total_table_count(self, request):
        """GET /api/recommend5|6|7/total_table_count/ – total rows in table (no filters)."""
        count = self.queryset.count()
        return Response({'count': count})


class Recommend5ViewSet(_RecommendViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Recommend5.objects.all()
    serializer_class = Recommend5Serializer


class Recommend6ViewSet(_RecommendViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Recommend6.objects.all()
    serializer_class = Recommend6Serializer


class Recommend7ViewSet(_RecommendViewSetMixin, viewsets.ReadOnlyModelViewSet):
    queryset = Recommend7.objects.all()
    serializer_class = Recommend7Serializer


class BanbeisViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Banbeis.objects.all()
    serializer_class = BanbeisSerializer
    pagination_class = JobListPagination
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = super().get_queryset()
        eiin = self.request.query_params.get('eiin')
        if eiin is not None and str(eiin).strip() != '':
            try:
                eiin_int = int(str(eiin).strip())
                qs = qs.filter(EIIN=eiin_int)
            except (ValueError, TypeError):
                pass
        return qs


class InstitutesViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Institutes.objects.all()
    serializer_class = InstitutesSerializer
    pagination_class = JobListPagination
    permission_classes = [AllowAny]
    filter_backends = [filters.SearchFilter]
    search_fields = [
        'instituteName', 'instituteNameBn', 'eiinNo', 'districtName', 'thanaName',
        'mobile', 'mobileAlternate', 'divisionName', 'divisionNameBn', 'districtNameBn',
        'thanaNameBn', 'instituteTypeName', 'instituteTypeNameBn', 'mouzaName', 'mouzaNameBn', 'email'
    ]

    def get_queryset(self):
        qs = super().get_queryset()
        # Optional search query param ?q= – match if any of these fields contains the term
        q = self.request.query_params.get('q')
        if q and str(q).strip():
            q = str(q).strip()
            q_any = (
                Q(eiinNo__icontains=q) |
                Q(mobile__icontains=q) |
                Q(mobileAlternate__icontains=q) |
                Q(instituteName__icontains=q) |
                Q(instituteNameBn__icontains=q) |
                Q(divisionName__icontains=q) |
                Q(divisionNameBn__icontains=q) |
                Q(districtName__icontains=q) |
                Q(districtNameBn__icontains=q) |
                Q(thanaName__icontains=q) |
                Q(thanaNameBn__icontains=q) |
                Q(instituteTypeName__icontains=q) |
                Q(instituteTypeNameBn__icontains=q) |
                Q(mouzaName__icontains=q) |
                Q(mouzaNameBn__icontains=q) |
                Q(email__icontains=q)
            )
            if q.isdigit():
                q_any = q_any | Q(id=int(q)) | Q(year=int(q))
            q_lower = q.lower()
            if q_lower in ('true', 'govt', 'government', 'yes', '1'):
                q_any = q_any | Q(isGovt=True)
            elif q_lower in ('false', 'non govt', 'no', '0'):
                q_any = q_any | Q(isGovt=False)
            qs = qs.filter(q_any)
        # Filter by type/division/district/thana (e.g. from institute filter UI)
        for param, field in [
            ('instituteTypeName', 'instituteTypeName'),
            ('divisionName', 'divisionName'),
            ('districtName', 'districtName'),
            ('thanaName', 'thanaName'),
        ]:
            values = self.request.query_params.getlist(param)
            values = [v.strip() for v in values if v and str(v).strip()]
            if values:
                qs = qs.filter(**{f'{field}__in': values})
        return qs

    @action(detail=False, url_path='unique_types', methods=['get'])
    def unique_types(self, request):
        """GET /api/institutes/unique_types/ – return all unique instituteTypeName from cheradip_institutes."""
        values = (
            Institutes.objects
            .values_list('instituteTypeName', flat=True)
            .distinct()
            .order_by('instituteTypeName')
        )
        # Exclude null/blank, return as list
        types = [v for v in values if v and str(v).strip()]
        return Response(types)

    @action(detail=False, url_path='unique_divisions', methods=['get'])
    def unique_divisions(self, request):
        """GET /api/institutes/unique_divisions/ – return all unique divisionName from cheradip_institutes."""
        values = (
            Institutes.objects
            .values_list('divisionName', flat=True)
            .distinct()
            .order_by('divisionName')
        )
        divisions = [v for v in values if v and str(v).strip()]
        return Response(divisions)

    @action(detail=False, url_path='unique_districts', methods=['get'])
    def unique_districts(self, request):
        """GET /api/institutes/unique_districts/?divisionName=Rangpur – unique districtName for that divisionName."""
        division_names = request.query_params.getlist('divisionName')
        division_names = [d.strip() for d in division_names if d and str(d).strip()]
        qs = Institutes.objects.all()
        if division_names:
            qs = qs.filter(divisionName__in=division_names)
        values = qs.values_list('districtName', flat=True).distinct().order_by('districtName')
        districts = [v for v in values if v and str(v).strip()]
        return Response(districts)

    @action(detail=False, url_path='unique_thanas', methods=['get'])
    def unique_thanas(self, request):
        """GET /api/institutes/unique_thanas/?districtName=Dinajpur – unique thanaName for that districtName."""
        district_names = request.query_params.getlist('districtName')
        district_names = [d.strip() for d in district_names if d and str(d).strip()]
        qs = Institutes.objects.all()
        if district_names:
            qs = qs.filter(districtName__in=district_names)
        values = qs.values_list('thanaName', flat=True).distinct().order_by('thanaName')
        thanas = [v for v in values if v and str(v).strip()]
        return Response(thanas)


class InstituteDetailView(APIView):
    """GET /api/institute/?eiin=XXX – return one institute (Institutes or Banbeis by EIIN) for unlock details."""
    permission_classes = [AllowAny]

    def get(self, request):
        eiin = request.query_params.get('eiin')
        if not eiin:
            return Response({}, status=status.HTTP_400_BAD_REQUEST)
        eiin_str = str(eiin).strip()
        # Try Institutes first (PK eiinNo) – routed to job DB
        try:
            inst = Institutes.objects.get(eiinNo=eiin_str)
            serializer = InstitutesSerializer(inst)
            return Response({'results': [serializer.data]})
        except Institutes.DoesNotExist:
            pass
        # Fallback: Banbeis by EIIN (bigint) – routed to job DB
        try:
            eiin_int = int(eiin_str)
            banbeis = Banbeis.objects.filter(EIIN=eiin_int).first()
            if banbeis:
                serializer = BanbeisSerializer(banbeis)
                return Response({'results': [serializer.data]})
        except (ValueError, TypeError):
            pass
        return Response({}, status=status.HTTP_404_NOT_FOUND)


SITEMAP_URLS_PER_FILE = 50000  # Sitemap spec limit

# XSL so browsers render sitemap XML as a readable page (removes "no style information" message)
SITEMAP_XSL = '''<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0" xmlns:xsl="http://www.w3.org/1999/XSL/Transform" xmlns:sm="http://www.sitemaps.org/schemas/sitemap/0.9">
  <xsl:output method="html" doctype-system="about:legacy-compat" encoding="UTF-8"/>
  <xsl:template match="/">
    <html>
      <head><meta charset="UTF-8"/><title>Sitemap</title>
        <style>
          body { font-family: system-ui, sans-serif; margin: 1rem 2rem; background: #f5f5f5; }
          h1 { color: #333; }
          table { border-collapse: collapse; background: #fff; box-shadow: 0 1px 3px rgba(0,0,0,.1); }
          th, td { padding: .5rem .75rem; text-align: left; border-bottom: 1px solid #eee; }
          th { background: #1976d2; color: #fff; }
          a { color: #1976d2; }
          tr:hover { background: #f9f9f9; }
        </style>
      </head>
      <body>
        <h1>Sitemap</h1>
        <xsl:choose>
          <xsl:when test="sm:sitemapindex">
            <table><tr><th>#</th><th>Sitemap</th></tr>
              <xsl:for-each select="sm:sitemapindex/sm:sitemap">
                <tr><td><xsl:value-of select="position()"/></td>
                  <td><a><xsl:attribute name="href"><xsl:value-of select="sm:loc"/></xsl:attribute><xsl:value-of select="sm:loc"/></a></td></tr>
              </xsl:for-each>
            </table>
          </xsl:when>
          <xsl:otherwise>
            <table><tr><th>#</th><th>URL</th></tr>
              <xsl:for-each select="sm:urlset/sm:url">
                <tr><td><xsl:value-of select="position()"/></td>
                  <td><a><xsl:attribute name="href"><xsl:value-of select="sm:loc"/></xsl:attribute><xsl:value-of select="sm:loc"/></a></td></tr>
              </xsl:for-each>
            </table>
          </xsl:otherwise>
        </xsl:choose>
      </body>
    </html>
  </xsl:template>
</xsl:stylesheet>'''


def _sitemap_base_url(request):
    """Base URL for sitemap. Set SITEMAP_BASE_URL in settings for production (e.g. https://cheradip.com)."""
    base = getattr(settings, 'SITEMAP_BASE_URL', None)
    if base:
        return base.rstrip('/')
    # Local dev: point to Angular app so sitemap URLs are openable (default 4200)
    if getattr(settings, 'DEBUG', False):
        return 'http://localhost:4200'
    return request.build_absolute_uri('/').rstrip('/') or 'https://cheradip.com'


def _sitemap_static_pages(base_url):
    """Yield <url> lines for all navigable app routes (fcheradip), in hierarchical order."""
    # Grouped so sitemap is easy to scan; comments mark sections.
    sections = [
        ('Home', ['index']),
        ('Products / Shop', ['packages', 'books', 'cart', 'choice', 'order']),
        ('User / Account', ['about_us', 'faqs', 'live_chat', 'login', 'auth', 'admin', 'profile', 'password', 'mobile', 'myorder']),
        ('NTRCA', ['ntrca', 'institute', 'institutes', 'vacant7', 'vacant5', 'vacant6', 'merit7', 'merit5', 'merit6', 'recommend7', 'recommend5', 'recommend6']),
        ('Student', ['student', 'student/dashboard', 'student/liveexam', 'student/archive', 'student/report', 'student/stats', 'student/leaderboard', 'student/tutor']),
        ('Question Bank', ['question']),
        ('Other', ['scrape']),
    ]
    for label, paths in sections:
        yield f'  <!-- {label} -->'
        for path in paths:
            loc = f'{base_url}/{path}' if path else base_url
            yield f'  <url><loc>{loc}</loc><changefreq>weekly</changefreq><priority>0.8</priority></url>'


def _sitemap_slug_for_url(slug):
    """Encode slug for URL in sitemap: keep Bengali/Unicode as literal; encode space, %, &, etc."""
    if not slug:
        return slug
    return slug.replace('%', '%25').replace('&', '%26').replace(' ', '%20')


def _sitemap_institute_url_entries(base_url):
    """Yield <url>...</url> lines for all institutes: EIIN-only, EIIN-English name, EIIN-Bengali name.
    Bengali and other Unicode appear as literal characters in the XML (UTF-8) so they display correctly."""
    for inst in Institutes.objects.all().values('eiinNo', 'instituteName', 'instituteNameBn').iterator(chunk_size=5000):
        eiin = (inst.get('eiinNo') or '').strip()
        if not eiin:
            continue
        name_en = (inst.get('instituteName') or '').strip()
        name_bn = (inst.get('instituteNameBn') or '').strip()
        yield f'  <url><loc>{base_url}/institutes/{_sitemap_slug_for_url(eiin)}</loc><changefreq>weekly</changefreq></url>'
        if name_en:
            slug_en = f"{eiin}-{name_en}"
            yield f'  <url><loc>{base_url}/institutes/{_sitemap_slug_for_url(slug_en)}</loc><changefreq>weekly</changefreq></url>'
        if name_bn:
            slug_bn = f"{eiin}-{name_bn}"
            yield f'  <url><loc>{base_url}/institutes/{_sitemap_slug_for_url(slug_bn)}</loc><changefreq>weekly</changefreq></url>'


def sitemap_xsl(request):
    """Serve XSL stylesheet so browsers render sitemap XML as a readable page."""
    return HttpResponse(SITEMAP_XSL.encode('utf-8'), content_type='application/xml; charset=utf-8')


def sitemap_pages(request):
    """Serve sitemap_pages.xml: all static app URLs from fcheradip (hierarchical)."""
    base = _sitemap_base_url(request)
    url_lines = list(_sitemap_static_pages(base))
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<?xml-stylesheet type="text/xsl" href="sitemap.xsl"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        + '\n'.join(url_lines) + '\n'
        '</urlset>'
    )
    return HttpResponse(xml.encode('utf-8'), content_type='application/xml; charset=utf-8')


def sitemap_institutes(request):
    """Serve sitemap index at sitemap.xml: pages first, then institute parts (111k+ URLs)."""
    base = _sitemap_base_url(request)
    # Hierarchy: 1) App pages, 2) Institute detail URLs (3 parts)
    index_lines = [
        f'  <sitemap><loc>{base}/sitemap_pages.xml</loc></sitemap>',
        f'  <sitemap><loc>{base}/sitemap_institutes_1.xml</loc></sitemap>',
        f'  <sitemap><loc>{base}/sitemap_institutes_2.xml</loc></sitemap>',
        f'  <sitemap><loc>{base}/sitemap_institutes_3.xml</loc></sitemap>',
    ]
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<?xml-stylesheet type="text/xsl" href="sitemap.xsl"?>\n'
        '<sitemapindex xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        + '\n'.join(index_lines) + '\n'
        '</sitemapindex>'
    )
    return HttpResponse(xml.encode('utf-8'), content_type='application/xml; charset=utf-8')


def sitemap_institutes_part(request, page):
    """Serve one sitemap part (page 1, 2, or 3) with up to 50,000 institute URLs."""
    if page not in (1, 2, 3):
        return HttpResponse(status=404)
    base = _sitemap_base_url(request)
    skip = (page - 1) * SITEMAP_URLS_PER_FILE
    take = SITEMAP_URLS_PER_FILE
    url_lines = []
    try:
        for line in _sitemap_institute_url_entries(base):
            if skip > 0:
                skip -= 1
                continue
            url_lines.append(line)
            if len(url_lines) >= take:
                break
    except Exception as e:
        logger.warning("Sitemap part %s generation failed: %s", page, e)
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<?xml-stylesheet type="text/xsl" href="sitemap.xsl"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        + '\n'.join(url_lines) + '\n'
        '</urlset>'
    )
    return HttpResponse(xml.encode('utf-8'), content_type='application/xml; charset=utf-8')


class TokenViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Token.objects.all()
    serializer_class = TokenSerializer
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = super().get_queryset()
        token_val = self.request.query_params.get('token')
        if token_val is not None and token_val != '':
            try:
                qs = qs.filter(Token=int(token_val))
            except (ValueError, TypeError):
                qs = qs.none()
        return qs

    def list(self, request, *args, **kwargs):
        token_val = request.query_params.get('token')
        if token_val is not None and token_val != '':
            # Validate single token: return one result + success/counter for back.ts and vacant6
            qs = self.get_queryset()
            obj = qs.first()
            try:
                counter_int = int(obj.Counter) if obj and obj.Counter not in (None, '') else 0
            except (ValueError, TypeError):
                counter_int = 0
            status_ok = getattr(obj, 'Status', 0) == 0 if obj else False
            success = bool(obj and status_ok and counter_int > 0)
            serializer = self.get_serializer(qs, many=True)
            data = {
                'count': 1 if obj else 0,
                'results': serializer.data,
                'success': success,
                'counter': counter_int if success else 0,
            }
            return Response(data)
        return super().list(request, *args, **kwargs)

    @action(detail=True, methods=['post'], url_path='update_status')
    def update_status(self, request, pk=None):
        """Set Status=1 (used) for this token. Body: { \"Status\": 1 }."""
        token_obj = self.get_object()
        token_obj.Status = 1
        token_obj.save(update_fields=['Status'])
        return Response({'success': True})

    def create(self, request, *args, **kwargs):
        """Use one token unlock: POST { token, eiin }. Decrements Counter, returns { success, remaining }."""
        token_val = request.data.get('token')
        eiin = request.data.get('eiin')
        if not token_val:
            return Response({'success': False, 'remaining': 0}, status=status.HTTP_400_BAD_REQUEST)
        try:
            token_int = int(token_val)
        except (ValueError, TypeError):
            return Response({'success': False, 'remaining': 0}, status=status.HTTP_400_BAD_REQUEST)
        obj = Token.objects.filter(Token=token_int).first()
        if not obj:
            return Response({'success': False, 'remaining': 0})
        try:
            counter_int = int(obj.Counter) if obj.Counter not in (None, '') else 0
        except (ValueError, TypeError):
            counter_int = 0
        if getattr(obj, 'Status', 0) != 0 or counter_int <= 0:
            return Response({'success': False, 'remaining': 0})
        counter_int -= 1
        obj.Counter = str(counter_int)
        obj.save(update_fields=['Counter'])
        return Response({'success': True, 'remaining': counter_int})


class ItemListCreateView(generics.ListCreateAPIView):
    queryset = Item.objects.all()
    serializer_class = ItemSerializer

    def get_permissions(self):
        if self.request.method in ['GET', 'HEAD', 'OPTIONS']:
            return [PublicAccess()]
        else:
            return [IsSuperUserOrStaff()]
            
    def list(self, request, *args, **kwargs):
        response = super().list(request, *args, **kwargs)

        # Update image URLs to include 'manage' prefix
        for item_data in response.data:
            if 'image' in item_data and item_data['image']:
                item_data['image'] = f'{settings.HOST_URL}/manage/media/{item_data["image"].split("/media/")[-1]}'

        return response

@api_view(['GET'])
@permission_classes([IsSuperUserOrStaff])
def item_list(request):
    items = Item.objects.all()
    serializer = ItemSerializer(items, many=True)
    return Response(serializer.data) 


class CustomerCreateView(APIView):
    """
    Signup: create one Customer row in cheradip_customers for Student, Teacher, or Job Seeker.
    Same table and same endpoint for all account types.
    """
    permission_classes = [AllowAny]
    authentication_classes = []

    def _get(self, raw, key, default=None):
        v = raw.get(key, default)
        return v[0] if isinstance(v, (list, tuple)) and len(v) else v

    @staticmethod
    def _normalize_acctype(value):
        """Normalize acctype to model choice: Student, Teacher, JobSeeker."""
        if not value or not str(value).strip():
            return 'Student'
        v = str(value).strip()
        if v == 'Job Seeker':
            return 'JobSeeker'
        return v

    @staticmethod
    def _date_to_iso(value):
        """Accept DD/MM/YYYY or YYYY-MM-DD; return YYYY-MM-DD for serializer, or None if invalid/empty."""
        if value is None or (isinstance(value, str) and not value.strip()):
            return None
        s = str(value).strip()
        if not s:
            return None
        # Already ISO (YYYY-MM-DD)
        if len(s) == 10 and s[4] == '-' and s[7] == '-':
            return s
        # DD/MM/YYYY
        parts = s.split('/')
        if len(parts) == 3:
            try:
                d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= d <= 31 and 1 <= m <= 12 and 1900 <= y <= 2100:
                    return f'{y:04d}-{m:02d}-{d:02d}'
            except (ValueError, TypeError):
                pass
        return None

    def post(self, request, *args, **kwargs):
        raw = request.data
        acctype = self._normalize_acctype(self._get(raw, 'acctype', 'Student'))
        country_code = self._get(raw, 'country_code') or self._get(raw, 'countryCode') or 'US'
        date_of_birth = self._date_to_iso(self._get(raw, 'date_of_birth'))
        user_data = {
            'acctype': acctype,  # Student | Teacher | JobSeeker
            'fullName': self._get(raw, 'fullName', ''),
            'username': self._get(raw, 'username', ''),
            'password': self._get(raw, 'password', ''),
            'country_code': country_code,
            'date_of_birth': date_of_birth,
            'class_name': self._get(raw, 'class_name'),
            'group': self._get(raw, 'group'),
            'department': self._get(raw, 'department'),
            'teacher_level': self._get(raw, 'teacher_level'),
            'teacher_subject_code': self._get(raw, 'teacher_subject_code'),
            'teacher_department_code': self._get(raw, 'teacher_department_code'),
            'teacher_department_name': self._get(raw, 'teacher_department_name'),
            'gender': self._get(raw, 'gender', 'Male'),
            'email': self._get(raw, 'email'),
        }
        serializer = CustomerSignupSerializer(data=user_data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = serializer.save()
            if acctype == 'Teacher':
                tcode = self._get(raw, 'teacher_department_code', '')
                tname = self._get(raw, 'teacher_department_name', '')
                if (tcode or '').strip().upper() == 'OTHER' and (tname or '').strip():
                    _append_department_to_json((tname or '').strip(), None)
            token = self.generate_unique_key()
            CustomerToken.objects.create(key=token, customer=user)
            return Response({'authToken': token}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Signup save failed: %s', e)
            err_msg = 'Signup failed. Please try again.'
            if getattr(settings, 'DEBUG', False):
                err_msg = str(e)
            return Response({'detail': err_msg}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def generate_unique_key(self):
        length = 40
        characters = string.ascii_letters + string.digits
        return ''.join(random.choice(characters) for _ in range(length))


class SignupProfileView(APIView):
    """
    GET /api/signup_profile/?username=xxx&acctype=Teacher|Student|JobSeeker
    Returns profile data from Customer table. Password is never returned.
    """
    permission_classes = [AllowAny]
    authentication_classes = []

    def get(self, request):
        username = request.query_params.get('username')
        acctype = (request.query_params.get('acctype') or '').strip()
        if not username:
            return Response({'detail': 'username is required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            q = Customer.objects.filter(username=username)
            if acctype and acctype in ('Teacher', 'Student', 'JobSeeker'):
                q = q.filter(acctype=acctype)
            obj = q.first()
            if not obj:
                return Response({'detail': 'Profile not found for this username and account type.'}, status=status.HTTP_404_NOT_FOUND)
            data = {
                'acctype': obj.acctype,
                'fullName': obj.fullName,
                'username': obj.username,
                'date_of_birth': obj.date_of_birth.isoformat() if obj.date_of_birth else None,
                'class_name': getattr(obj, 'class_name', None),
                'group': obj.group,
                'department': getattr(obj, 'department', None),
                'teacher_level': getattr(obj, 'teacher_level', None),
                'teacher_subject_code': getattr(obj, 'teacher_subject_code', None),
                'teacher_department_code': getattr(obj, 'teacher_department_code', None),
                'teacher_department_name': getattr(obj, 'teacher_department_name', None),
                'gender': obj.gender,
                'email': obj.email,
                'country_code': getattr(obj, 'country_code', None),
                'division': obj.division,
                'district': obj.district,
                'thana': obj.thana,
                'union': obj.union,
                'village': obj.village,
                'date_joined': obj.date_joined.isoformat() if obj.date_joined else None,
            }
            return Response(data, status=status.HTTP_200_OK)
        except Customer.DoesNotExist:
            return Response({'detail': 'Profile not found for this username and account type.'}, status=status.HTTP_404_NOT_FOUND)


class CustomerRetrieveView(APIView):
    """Login: authenticate against Customer only. Optional country_code filter on username lookup."""
    permission_classes = [AllowAny]
    authentication_classes = []

    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        raw_country = request.data.get('countryCode') or request.data.get('country_code')
        if hasattr(raw_country, 'get'):
            country_code = (raw_country.get('country_code') or raw_country.get('countryCode') or '').strip().upper() or None
        else:
            country_code = (str(raw_country).strip().upper() or None) if raw_country else None
        if not username or not password:
            return Response({'error': 'Username and password are required.'}, status=status.HTTP_400_BAD_REQUEST)

        user = None
        try:
            if country_code:
                user = Customer.objects.filter(username=username, country_code=country_code).first()
                if user and not user.check_password(password):
                    user = None
            else:
                user = authenticate(request, username=username, password=password)
        except (ProgrammingError, OperationalError):
            user = None
        if user is not None:
            acctype = getattr(user, 'acctype', None)
            fullName = getattr(user, 'fullName', None)
            group = getattr(user, 'group', None)
            gender = getattr(user, 'gender', None)
            division = getattr(user, 'division', None)
            district = getattr(user, 'district', None)
            thana = getattr(user, 'thana', None)
            union = getattr(user, 'union', None)
            village = getattr(user, 'village', None)
            token = self.generate_unique_key()
            CustomerToken.objects.filter(customer=user).delete()
            CustomerToken.objects.create(key=token, customer=user)
            return Response({
                'authToken': token,
                'acctype': acctype,
                'fullName': fullName,
                'username': user.username,
                'group': group,
                'gender': gender,
                'division': division,
                'district': district,
                'thana': thana,
                'union': union,
                'village': village,
            }, status=status.HTTP_200_OK)

        return Response({'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)

    def generate_unique_key(self):
        length = 40
        characters = string.ascii_letters + string.digits
        return ''.join(random.choice(characters) for _ in range(length))

    def get(self, request, *args, **kwargs):
        if request.user.is_authenticated:
            customer_data = {
                'username': request.user.username,
                'fullName': getattr(request.user, 'fullName', getattr(request.user, 'full_name', '')),
            }
            return Response(customer_data, status=status.HTTP_200_OK)
        return Response({'error': 'Not authenticated'}, status=status.HTTP_401_UNAUTHORIZED)



class BearerTokenAuthentication(BaseAuthentication):
    """Authenticate by Authorization: Bearer <token>; look up CustomerToken and set request.user."""
    def authenticate(self, request):
        auth = request.META.get('HTTP_AUTHORIZATION')
        if not auth or not auth.startswith('Bearer '):
            return None
        token = auth[7:].strip()
        if not token:
            return None
        try:
            ct = CustomerToken.objects.select_related('customer').get(key=token)
            return (ct.customer, token)
        except CustomerToken.DoesNotExist:
            return None


class CustomerSettingsView(APIView):
    """GET/POST customer settings (JSON). Requires Bearer token. Used for export_format etc."""
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        settings_dict = getattr(request.user, 'settings', None)
        if settings_dict is None:
            settings_dict = {}
        return Response({'settings': settings_dict}, status=status.HTTP_200_OK)

    def post(self, request):
        new_settings = request.data.get('settings')
        if new_settings is not None and not isinstance(new_settings, dict):
            return Response({'error': 'settings must be a JSON object'}, status=status.HTTP_400_BAD_REQUEST)
        user = request.user
        current = getattr(user, 'settings', None) or {}
        if new_settings is not None:
            current = {**current, **new_settings}
            user.settings = current
            user.save(update_fields=['settings'])
        return Response({'settings': current}, status=status.HTTP_200_OK)


# Bengali-capable font for PDF (reportlab). Try Windows Vrinda, Nirmala UI, or project fonts.
_PDF_BENGALI_FONT_REGISTERED = None

def _get_pdf_bengali_font():
    """Register and return a Bengali-capable font name, or None to use Helvetica."""
    global _PDF_BENGALI_FONT_REGISTERED
    if _PDF_BENGALI_FONT_REGISTERED is not None:
        return _PDF_BENGALI_FONT_REGISTERED
    if not REPORTLAB_AVAILABLE or pdfmetrics is None or TTFont is None:
        return None
    font_name = 'BengaliPDF'
    paths_to_try = []
    if hasattr(settings, 'BASE_DIR'):
        base = getattr(settings, 'BASE_DIR', None)
        if base:
            paths_to_try.extend([
                os.path.join(base, 'cheradip', 'static', 'fonts', 'NotoSansBengali-Regular.ttf'),
                os.path.join(base, 'static', 'fonts', 'NotoSansBengali-Regular.ttf'),
            ])
    windir = os.environ.get('WINDIR', 'C:\\Windows')
    paths_to_try.extend([
        os.path.join(windir, 'Fonts', 'vrinda.ttf'),
        os.path.join(windir, 'Fonts', 'NirmalaUI.ttf'),
        os.path.join(windir, 'Fonts', 'Nirmala.ttf'),
    ])
    for path in paths_to_try:
        if path and os.path.isfile(path):
            try:
                font = TTFont(font_name, path)
                pdfmetrics.registerFont(font)
                _PDF_BENGALI_FONT_REGISTERED = font_name
                return font_name
            except Exception as e:
                logger.debug('Could not register PDF font %s: %s', path, e)
    _PDF_BENGALI_FONT_REGISTERED = False
    return None


# Page size (width_pt, height_pt) for reportlab. 1 inch = 72 pt, 1 mm = 72/25.4 pt.
def _export_page_size_pt(name):
    from reportlab.lib.pagesizes import A4, A3, A5, letter, legal
    pt_per_mm = 72 / 25.4
    sizes = {
        'A4': A4,
        'A3': A3,
        'A5': A5,
        'Letter': letter,
        'Legal': legal,
        'B4': (250 * pt_per_mm, 353 * pt_per_mm),
        'B5': (176 * pt_per_mm, 250 * pt_per_mm),
        'Tabloid': (279.4 * pt_per_mm, 431.8 * pt_per_mm),
    }
    return sizes.get(name, A4)


def _export_page_size_mm(name):
    """Return (width_mm, height_mm) for python-docx."""
    sizes_mm = {
        'A4': (210, 297),
        'A3': (297, 420),
        'A5': (148, 210),
        'B4': (250, 353),
        'B5': (176, 250),
        'Letter': (215.9, 279.4),
        'Legal': (215.9, 355.6),
        'Tabloid': (279.4, 431.8),
    }
    return sizes_mm.get(name, (210, 297))


def _export_page_size_pt_from_mm(w_mm, h_mm):
    pt_per_mm = 72 / 25.4
    return (float(w_mm) * pt_per_mm, float(h_mm) * pt_per_mm)


def _export_resolve_page_mm(data):
    """Paper size in mm from request: named size, optional Custom (inches), optional landscape swap."""
    name = (data.get('pageSize') or data.get('page_size') or 'A4').strip() or 'A4'
    orient = (data.get('pageOrientation') or data.get('page_orientation') or 'portrait').strip().lower()
    landscape = orient in ('landscape', 'l', 'land')
    if name.lower() == 'custom':
        try:
            wi = float(data.get('customPageWidthIn') or data.get('custom_page_width_in') or 8.5)
            hi = float(data.get('customPageHeightIn') or data.get('custom_page_height_in') or 11)
        except (TypeError, ValueError):
            wi, hi = 8.5, 11
        wi = max(2.0, min(48.0, wi))
        hi = max(2.0, min(48.0, hi))
        w_mm = wi * 25.4
        h_mm = hi * 25.4
    else:
        w_mm, h_mm = _export_page_size_mm(name)
    if landscape:
        w_mm, h_mm = h_mm, w_mm
    return w_mm, h_mm


def _export_page_sections(data):
    try:
        n = int(data.get('pageSections') or data.get('page_sections') or 1)
    except (TypeError, ValueError):
        n = 1
    return max(1, min(10, n))


def _export_layout_section_gap_px(data):
    try:
        g = int(data.get('sectionGapPx') or data.get('section_gap_px') or 14)
    except (TypeError, ValueError):
        g = 14
    return max(1, min(100, g))


def _docx_apply_section_columns(section, num_cols, space_twips=210, show_sep=False):
    """Apply N newspaper-style columns to a python-docx section (2–10). No-op if unavailable or num_cols < 2."""
    if not DOCX_AVAILABLE or num_cols <= 1:
        return
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        n = max(2, min(10, int(num_cols)))
        stw = max(20, min(1600, int(space_twips)))
        sect_pr = section._sectPr
        cols_el = sect_pr.find(qn('w:cols'))
        if cols_el is None:
            cols_el = OxmlElement('w:cols')
            sect_pr.append(cols_el)
        cols_el.set(qn('w:num'), str(n))
        cols_el.set(qn('w:space'), str(stw))
        if show_sep:
            cols_el.set(qn('w:sep'), '1')
        elif qn('w:sep') in cols_el.attrib:
            del cols_el.attrib[qn('w:sep')]
    except Exception:
        logger.debug('docx section columns not applied', exc_info=True)


def _export_layout_column_gap_px(data):
    try:
        g = int(data.get('layoutColumnGapPx') or data.get('layout_column_gap_px') or 14)
    except (TypeError, ValueError):
        g = 14
    return max(1, min(100, g))


def _export_show_column_divider(data):
    raw = data.get('showColumnDivider')
    if raw is None:
        raw = data.get('show_column_divider')
    if raw is None:
        return True
    if isinstance(raw, str):
        return raw.strip().lower() in ('1', 'true', 'yes', 'on')
    return bool(raw)


def _pdf_px_to_pt(px):
    return float(px) * 72.0 / 96.0


def _pdf_paint_column_dividers(canvas, left, bottom_y, top_y, col_w, gap_pt, ncols):
    """Vertical lines in the middle of each column gutter (multi-column PDF)."""
    if ncols <= 1 or top_y <= bottom_y:
        return
    canvas.saveState()
    canvas.setStrokeColorRGB(0.78, 0.78, 0.78)
    canvas.setLineWidth(0.5)
    for k in range(1, ncols):
        x = left + k * col_w + (k - 0.5) * gap_pt
        canvas.line(x, bottom_y, x, top_y)
    canvas.restoreState()


def _pdf_paint_horizontal_section_dividers(canvas, left, right, bottom_y, top_y, n_sections, gap_pt=0):
    """Horizontal lines in the middle of each section gutter (PDF y increases upward)."""
    if n_sections <= 1 or top_y <= bottom_y:
        return
    span = float(top_y - bottom_y)
    n = float(n_sections)
    g = max(0.0, float(gap_pt))
    band = (span - (n - 1) * g) / n
    if band <= 0:
        return
    canvas.saveState()
    canvas.setStrokeColorRGB(0.78, 0.78, 0.78)
    canvas.setLineWidth(0.5)
    for k in range(1, n_sections):
        y = top_y - k * band - (k - 0.5) * g
        canvas.line(left, y, right, y)
    canvas.restoreState()


class ExportQuestionsView(APIView):
    """POST: generate PDF or DOCX from questions list. Requires Bearer auth. Body: questions, questionHeader, pageSize, marginTop, marginRight, marginBottom, marginLeft, format ('pdf'|'docx'), filename (optional)."""
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def _playwright_local_font_face_css(self):
        """
        Build @font-face rules from local font files so Playwright PDF output is
        stable across Windows machines.
        """
        base_dir = Path(getattr(settings, 'BASE_DIR', os.getcwd()))
        candidates = [
            base_dir / 'fonts',
            base_dir / 'static' / 'fonts',
            base_dir / 'static_src' / 'fonts',
            base_dir / 'cheradip' / 'fonts',
            base_dir.parent / 'fcheradip' / 'src' / 'assets' / 'fonts',
        ]
        fonts_dir = None
        for c in candidates:
            if c.exists() and c.is_dir():
                fonts_dir = c
                break
        if fonts_dir is None:
            return ''

        # filename -> (family, weight, style)
        wanted = {
            'Roboto-Regular.ttf': ('Roboto', '400', 'normal'),
            'Roboto-Thin.ttf': ('Roboto', '300', 'normal'),
            'NotoSansBengali-Regular.ttf': ('Noto Sans Bengali', '400', 'normal'),
            'NotoSansBengali-Bold.ttf': ('Noto Sans Bengali', '700', 'normal'),
            'NotoSerifBengali-Regular.ttf': ('Noto Serif Bengali', '400', 'normal'),
            'NotoSerifBengali-Bold.ttf': ('Noto Serif Bengali', '700', 'normal'),
            'SolaimanLipi.ttf': ('SolaimanLipi', '400', 'normal'),
            'Kalpurush.ttf': ('Kalpurush', '400', 'normal'),
            'NotoSans-Regular.ttf': ('Noto Sans', '400', 'normal'),
            'NotoSans-Bold.ttf': ('Noto Sans', '700', 'normal'),
            'NotoSerif-Regular.ttf': ('Noto Serif', '400', 'normal'),
            'NotoSerif-Bold.ttf': ('Noto Serif', '700', 'normal'),
            'STIXTwoMath-Regular.otf': ('STIX Two Math', '400', 'normal'),
        }
        rules = []
        for fn, spec in wanted.items():
            fp = fonts_dir / fn
            if not fp.exists():
                continue
            family, weight, style = spec
            ext = fp.suffix.lower()
            fmt = 'opentype' if ext == '.otf' else 'truetype'
            rules.append(
                "@font-face { "
                f"font-family: '{family}'; "
                f"src: url('{fp.resolve().as_uri()}') format('{fmt}'); "
                f"font-weight: {weight}; "
                f"font-style: {style}; "
                "font-display: swap; "
                "}"
            )
        return '\n'.join(rules)

    def post(self, request):
        data = request.data
        questions = data.get('questions') or []
        if not isinstance(questions, list) or len(questions) == 0:
            return Response({'error': 'questions list is required and must be non-empty'}, status=status.HTTP_400_BAD_REQUEST)
        fmt = (data.get('format') or '').strip().lower()
        if fmt not in ('pdf', 'docx'):
            return Response({'error': 'format must be pdf or docx'}, status=status.HTTP_400_BAD_REQUEST)
        question_header = (data.get('questionHeader') or '').strip()
        margin_top = float(data.get('marginTop') or 25.4)
        margin_right = float(data.get('marginRight') or 25.4)
        margin_bottom = float(data.get('marginBottom') or 25.4)
        margin_left = float(data.get('marginLeft') or 25.4)
        filename_base = (data.get('filename') or 'questions').strip() or 'questions'
        filename_base = re.sub(r'[^\w\-_.\s]', '_', filename_base)[:120]
        layout_columns = max(1, min(10, int(data.get('layoutColumns') or data.get('layout_columns') or 1)))
        layout_gap_px = _export_layout_column_gap_px(data)
        show_col_div = _export_show_column_divider(data)
        w_mm, h_mm = _export_resolve_page_mm(data)
        page_sections = _export_page_sections(data)
        section_gap_px = _export_layout_section_gap_px(data)
        layout_settings = data.get('layout_settings')
        if not isinstance(layout_settings, dict):
            layout_settings = {}
        wants_preview_style_pdf = (
            bool(layout_settings)
            or any(
                k in data
                for k in (
                    'previewQuestionsFontPx',
                    'previewQuestionsFontPxCreative',
                    'previewQuestionsFontPxMcq',
                    'previewQuestionsLineHeight',
                    'previewQuestionsLineHeightCreative',
                    'previewQuestionsLineHeightMcq',
                    'cqPageOrientation',
                    'mcqPageOrientation',
                    'questionsGap',
                    'questionsGapCreative',
                    'questionsPadding',
                    'headerLineFontSizes',
                    'layoutColumnsCreative',
                )
            )
        )

        if fmt == 'pdf':
            buf = None
            if PLAYWRIGHT_AVAILABLE:
                try:
                    buf = self._build_pdf_playwright(
                        questions=questions,
                        question_header=question_header,
                        margin_top=margin_top,
                        margin_right=margin_right,
                        margin_bottom=margin_bottom,
                        margin_left=margin_left,
                        page_w_mm=w_mm,
                        page_h_mm=h_mm,
                        layout_columns=layout_columns,
                        layout_column_gap_px=layout_gap_px,
                        show_column_divider=show_col_div,
                        page_sections=page_sections,
                        section_gap_px=section_gap_px,
                        raw_data=data,
                        layout_settings=layout_settings,
                    )
                except Exception:
                    logger.exception('Playwright PDF render failed')
                    if wants_preview_style_pdf:
                        return Response(
                            {
                                'error': (
                                    'Preview-style PDF rendering failed on server. '
                                    'Run "playwright install chromium" in backend venv and ensure fonts are installed.'
                                )
                            },
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR
                        )
            elif wants_preview_style_pdf:
                return Response(
                    {
                        'error': (
                            'Preview-style PDF requires Playwright. '
                            'Install dependency and run "playwright install chromium".'
                        )
                    },
                    status=status.HTTP_503_SERVICE_UNAVAILABLE
                )
            if buf is None:
                if not REPORTLAB_AVAILABLE:
                    return Response({'error': 'PDF generation not available (reportlab/playwright not installed)'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
                buf = self._build_pdf(
                    questions, question_header, margin_top, margin_right, margin_bottom, margin_left,
                    page_w_mm=w_mm,
                    page_h_mm=h_mm,
                    layout_columns=layout_columns,
                    layout_column_gap_px=layout_gap_px,
                    show_column_divider=show_col_div,
                    page_sections=page_sections,
                    section_gap_px=section_gap_px,
                )
            resp = HttpResponse(buf.getvalue(), content_type='application/pdf')
            resp['Content-Disposition'] = 'attachment; filename="%s.pdf"' % filename_base.replace('"', '_')
            return resp
        else:
            if not DOCX_AVAILABLE:
                return Response({'error': 'DOCX generation not available (python-docx not installed)'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
            buf = self._build_docx(
                questions, question_header, margin_top, margin_right, margin_bottom, margin_left,
                page_w_mm=w_mm,
                page_h_mm=h_mm,
                layout_columns=layout_columns,
                layout_column_gap_px=layout_gap_px,
                show_column_divider=show_col_div,
            )
            resp = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            resp['Content-Disposition'] = 'attachment; filename="%s.docx"' % filename_base.replace('"', '_')
            return resp

    def _build_pdf_playwright(
        self,
        questions,
        question_header,
        margin_top,
        margin_right,
        margin_bottom,
        margin_left,
        page_w_mm,
        page_h_mm,
        layout_columns=1,
        layout_column_gap_px=14,
        show_column_divider=True,
        page_sections=1,
        section_gap_px=14,
        raw_data=None,
        layout_settings=None,
    ):
        if not PLAYWRIGHT_AVAILABLE:
            return None
        d = raw_data if isinstance(raw_data, dict) else {}
        ls = layout_settings if isinstance(layout_settings, dict) else {}

        def num(v, fallback):
            try:
                n = float(v)
                if n != n:
                    return float(fallback)
                return n
            except Exception:
                return float(fallback)

        def intval(v, fallback):
            try:
                return int(round(float(v)))
            except Exception:
                return int(fallback)

        def pick(key, fallback):
            if key in d and d.get(key) is not None:
                return d.get(key)
            return ls.get(key, fallback)

        def page_mm_for_orientation(orientation_name):
            page_name = str(pick('pageSize', 'A4') or 'A4').strip() or 'A4'
            if page_name.lower() == 'custom':
                wi = num(pick('customPageWidthIn', 8.5), 8.5)
                hi = num(pick('customPageHeightIn', 11), 11)
                wi = max(2.0, min(48.0, wi))
                hi = max(2.0, min(48.0, hi))
                base_w = wi * 25.4
                base_h = hi * 25.4
            else:
                base_w, base_h = _export_page_size_mm(page_name)
            o = str(orientation_name or 'portrait').strip().lower()
            if o in ('landscape', 'l', 'land'):
                return (base_h, base_w)
            return (base_w, base_h)

        cq_orient = str(pick('cqPageOrientation', pick('pageOrientation', 'portrait')) or 'portrait').strip().lower()
        mcq_orient = str(pick('mcqPageOrientation', pick('pageOrientation', 'portrait')) or 'portrait').strip().lower()
        cq_w_mm, cq_h_mm = page_mm_for_orientation(cq_orient)
        mcq_w_mm, mcq_h_mm = page_mm_for_orientation(mcq_orient)

        q_font_global = num(pick('previewQuestionsFontPx', 16), 16)
        q_font_cq = num(pick('previewQuestionsFontPxCreative', q_font_global), q_font_global)
        q_font_mcq = num(pick('previewQuestionsFontPxMcq', q_font_global), q_font_global)
        q_lh_global = num(pick('previewQuestionsLineHeight', 1.4), 1.4)
        q_lh_cq = num(pick('previewQuestionsLineHeightCreative', q_lh_global), q_lh_global)
        q_lh_mcq = num(pick('previewQuestionsLineHeightMcq', q_lh_global), q_lh_global)
        h_lh = num(pick('previewHeaderLineHeight', 1.25), 1.25)
        q_pad = max(0, num(pick('questionsPadding', 2), 2))
        q_gap_mcq = max(0, num(pick('questionsGap', 2), 2))
        q_gap_cq = max(0, num(pick('questionsGapCreative', 4), 4))
        mcq_extra_bottom_mm = max(0.0, num(pick('mcqExtraBottomMarginMm', 0), 0))
        options_cols = max(1, min(4, intval(pick('optionsColumns', 2), 2)))
        cols_mcq = max(1, min(10, intval(pick('layoutColumns', layout_columns), layout_columns)))
        cols_cq = max(1, min(10, intval(pick('layoutColumnsCreative', cols_mcq), cols_mcq)))
        col_gap = max(1, min(100, intval(pick('layoutColumnGapPx', layout_column_gap_px), layout_column_gap_px)))
        show_div = bool(pick('showColumnDivider', show_column_divider))

        hfs_raw = pick('headerLineFontSizes', [16, 18, 16, 16, 14, 14, 14, 12])
        if isinstance(hfs_raw, list):
            hfs = [max(8, min(64, num(x, 14))) for x in hfs_raw]
        else:
            hfs = [16, 18, 16, 16, 14, 14, 14, 12]

        def sanitize_header_html(text):
            raw = str(text or '')
            # Keep only a small safe subset used by preview header editing.
            allowed = [
                '<br>', '<br/>', '<br />',
                '<hr>', '<hr/>', '<hr />',
                '<b>', '</b>',
                '<strong>', '</strong>',
                '<i>', '</i>',
                '<u>', '</u>',
            ]
            tokenized = raw
            tokens = {}
            for i, tag in enumerate(allowed):
                token = '__HDR_TAG_%d__' % i
                tokenized = re.sub(re.escape(tag), token, tokenized, flags=re.IGNORECASE)
                tokens[token] = tag
            out = escape(tokenized)
            for token, tag in tokens.items():
                out = out.replace(token, tag)
            return out

        def normalize_bengali_digits(s):
            return str(s or '').translate(str.maketrans('0123456789', '০১২৩৪৫৬৭৮৯'))

        def mcq_set_letter_from_line_or_settings(line_txt):
            """Match preview/export line `… সেট : ক` or layout_settings.mcqSetLetter."""
            t = str(line_txt or '')
            m = re.search(r'সেট\s*[:ঃ]\s*([কখগঘ])', t)
            if m:
                return m.group(1)
            v = pick('mcqSetLetter', None)
            if v is None:
                return None
            s = str(v).strip()
            if s in ('ক', 'খ', 'গ', 'ঘ'):
                return s
            return None

        def render_subject_code_row(line, font_px, line_h, extra_class='', for_mcq_header=False):
            txt = str(line or '')
            digits = re.findall(r'[0-9০-৯]', txt)
            digs = [normalize_bengali_digits(d) for d in digits[:3]]
            while len(digs) < 3:
                digs.append('&nbsp;')
            set_letter = mcq_set_letter_from_line_or_settings(txt) if for_mcq_header else None
            if set_letter:
                sl = escape(set_letter)
                return (
                    '<div class="hline hline-code-row-wrap%s" style="font-size:%.2fpx; line-height:%.3f;">'
                    '<span class="q-code-grid q-code-grid--mcq-set">'
                    '<span class="q-code-label">বিষয় কোড</span>'
                    '<span class="q-code-colon">:</span>'
                    '<span class="q-code-cell">%s</span>'
                    '<span class="q-code-cell">%s</span>'
                    '<span class="q-code-cell">%s</span>'
                    '<span class="q-code-label q-code-label--set">সেট</span>'
                    '<span class="q-code-colon">:</span>'
                    '<span class="q-code-cell q-code-cell--filler"></span>'
                    '<span class="q-code-cell q-code-cell--set-letter">%s</span>'
                    '<span class="q-code-cell q-code-cell--filler"></span>'
                    '</span>'
                    '</div>'
                ) % (
                    extra_class,
                    font_px,
                    line_h,
                    digs[0],
                    digs[1],
                    digs[2],
                    sl,
                )
            return (
                '<div class="hline hline-code-row-wrap%s" style="font-size:%.2fpx; line-height:%.3f;">'
                '<span class="q-code-grid">'
                '<span class="q-code-label">বিষয় কোড</span>'
                '<span class="q-code-colon">:</span>'
                '<span class="q-code-cell">%s</span>'
                '<span class="q-code-cell">%s</span>'
                '<span class="q-code-cell">%s</span>'
                '</span>'
                '</div>'
            ) % (extra_class, font_px, line_h, digs[0], digs[1], digs[2])

        def _cq_header_line_is_notice(line):
            """True for [দ্রষ্টব্য : …] rows — must not sit under the floating subject-code grid."""
            t = str(line or '').strip()
            if not t:
                return False
            if re.search(r'^\s*\[?\s*দ্রষ্টব্য', t):
                return True
            if t.startswith('[') and 'দ্রষ্টব্য' in t[:200]:
                return True
            return False

        def compile_playwright_header_html(qh_source, cq_header=False):
            """
            Compile header text into q-header HTML.
            The first "বিষয় কোড" line is converted to the subject-code grid and removed from flow.
            cq_header: Creative PDF — keep দ্রষ্টব্য lines out of q-header-band so the grid aligns with
            subject/title lines (e.g. বিষয় name), not over the notice paragraph.
            """
            header_lines = [ln for ln in str(qh_source or '').replace('\r\n', '\n').split('\n')]
            header_html = ''
            if any(s.strip() for s in header_lines):
                code_line_index = None
                for _ci, _ln in enumerate(header_lines):
                    if re.search(r'বিষ[য়য]\s*কোড', str(_ln or '').strip()):
                        code_line_index = _ci
                        break
                before_hr_chunks = []
                band_before_chunks = []
                notice_before_chunks = []
                after_hr_chunks = []
                after_raw_lines = []
                code_row_html = ''
                code_row_source_line = ''
                code_row_font_px = hfs[5] if len(hfs) > 5 else hfs[-1]
                for i, line in enumerate(header_lines):
                    fz = hfs[i] if i < len(hfs) else hfs[-1]
                    line_txt = str(line or '').strip()
                    if re.search(r'বিষ[য়য]\s*কোড', line_txt):
                        if code_row_html:
                            continue
                        code_row_source_line = line_txt
                        code_row_html = render_subject_code_row(
                            line_txt, fz, h_lh, '', not cq_header
                        )
                        code_row_font_px = fz
                        continue
                    if re.fullmatch(r'<hr\s*/?>', line_txt, flags=re.IGNORECASE):
                        after_hr_chunks.append('<hr class="hline-hr" />')
                        if cq_header:
                            after_raw_lines.append(None)
                        continue
                    rendered = sanitize_header_html(line if line else '')
                    line_html = '<div class="hline" style="font-size:%.2fpx; line-height:%.3f;">%s</div>' % (
                        fz,
                        h_lh,
                        rendered,
                    )
                    if code_line_index is not None and i < code_line_index:
                        if cq_header:
                            if _cq_header_line_is_notice(line):
                                notice_before_chunks.append(line_html)
                            else:
                                band_before_chunks.append(line_html)
                        else:
                            before_hr_chunks.append(line_html)
                    else:
                        after_hr_chunks.append(line_html)
                        if cq_header:
                            after_raw_lines.append(line)
                if code_row_html:
                    # Overlay code grid on the band of lines above the বিষয় কোড row (preview-like), not a
                    # separate full-width row. Works with or without <hr>; band = lines before code index only.
                    if cq_header:
                        if len(band_before_chunks) >= 1:
                            floating_code = render_subject_code_row(
                                code_row_source_line or 'বিষয় কোড',
                                code_row_font_px,
                                h_lh,
                                extra_class=' hline-code-row-wrap--floating',
                                for_mcq_header=False,
                            )
                            band_html = '<div class="q-header-band">%s%s</div>' % (
                                ''.join(band_before_chunks),
                                floating_code,
                            )
                            header_html = '<div class="q-header">%s%s%s</div>' % (
                                band_html,
                                ''.join(notice_before_chunks),
                                ''.join(after_hr_chunks),
                            )
                        elif after_hr_chunks:
                            rest = list(after_hr_chunks)
                            raws = list(after_raw_lines)
                            leading_hr = []
                            while rest and (rest[0] or '').strip().startswith('<hr') and 'hline-hr' in (
                                rest[0] or ''
                            ):
                                leading_hr.append(rest.pop(0))
                                if raws:
                                    raws.pop(0)
                            after_notices = []
                            while (
                                rest
                                and raws
                                and raws[0] is not None
                                and _cq_header_line_is_notice(raws[0])
                            ):
                                after_notices.append(rest.pop(0))
                                raws.pop(0)
                            if rest and (not raws or raws[0] is not None):
                                floating_code = render_subject_code_row(
                                    code_row_source_line or 'বিষয় কোড',
                                    code_row_font_px,
                                    h_lh,
                                    extra_class=' hline-code-row-wrap--floating',
                                    for_mcq_header=False,
                                )
                                band_line = rest.pop(0)
                                if raws:
                                    raws.pop(0)
                                band_html = '<div class="q-header-band">%s%s</div>' % (
                                    band_line,
                                    floating_code,
                                )
                                header_html = '<div class="q-header">%s%s%s%s%s</div>' % (
                                    ''.join(leading_hr),
                                    band_html,
                                    ''.join(notice_before_chunks),
                                    ''.join(after_notices),
                                    ''.join(rest),
                                )
                            else:
                                all_chunks = (
                                    notice_before_chunks
                                    + after_hr_chunks
                                )
                                all_chunks.insert(min(5, len(all_chunks)), code_row_html)
                                header_html = '<div class="q-header">%s</div>' % ''.join(all_chunks)
                        else:
                            all_chunks = notice_before_chunks + after_hr_chunks
                            all_chunks.insert(min(5, len(all_chunks)), code_row_html)
                            header_html = '<div class="q-header">%s</div>' % ''.join(all_chunks)
                    elif len(before_hr_chunks) >= 1:
                        floating_code = render_subject_code_row(
                            code_row_source_line or 'বিষয় কোড',
                            code_row_font_px,
                            h_lh,
                            extra_class=' hline-code-row-wrap--floating',
                            for_mcq_header=True,
                        )
                        band_html = '<div class="q-header-band">%s%s</div>' % (
                            ''.join(before_hr_chunks),
                            floating_code,
                        )
                        header_html = '<div class="q-header">%s%s</div>' % (
                            band_html,
                            ''.join(after_hr_chunks),
                        )
                    else:
                        all_chunks = before_hr_chunks + after_hr_chunks
                        all_chunks.insert(min(5, len(all_chunks)), code_row_html)
                        header_html = '<div class="q-header">%s</div>' % ''.join(all_chunks)
                else:
                    if cq_header:
                        header_html = '<div class="q-header">%s%s%s</div>' % (
                            ''.join(band_before_chunks),
                            ''.join(notice_before_chunks),
                            ''.join(after_hr_chunks),
                        )
                    else:
                        header_html = '<div class="q-header">%s%s</div>' % (
                            ''.join(before_hr_chunks),
                            ''.join(after_hr_chunks),
                        )
            return header_html

        qh_root = str(question_header or '').strip()
        qh_creative = str(pick('questionHeaderCreative', qh_root) or '').strip() or qh_root
        qh_mcq = str(pick('questionHeaderMcq', qh_root) or '').strip() or qh_root
        header_html_creative = compile_playwright_header_html(qh_creative, cq_header=True)
        header_html_mcq = compile_playwright_header_html(qh_mcq)

        def is_creative(q):
            t = str((q or {}).get('type') or '').strip().lower()
            return ('সৃজন' in t) or ('creative' in t)

        creative_questions = []
        mcq_questions = []
        for idx, q in enumerate(questions):
            if is_creative(q):
                creative_questions.append({'idx': idx, 'q': q})
            else:
                mcq_questions.append({'idx': idx, 'q': q})

        def wrap_roman_lines_html(text):
            roman_line = re.compile(r'^\s*(i|ii|iii|I|II|III)\.')
            bn_paren_line = re.compile(r'^\s*\([কখগঘ]\)')
            lines = str(text or '').splitlines()
            parts = []
            for line in lines:
                cls = 'topic-question-line'
                if roman_line.match(line):
                    cls = 'topic-question-line topic-question-roman-line'
                elif bn_paren_line.match(line):
                    cls = 'topic-question-line topic-question-bn-paren-line'
                parts.append('<span class="%s">%s</span>' % (cls, escape(line)))
            return ''.join(parts)

        def question_display_text(raw_text, creative):
            s = str(raw_text or '').strip()
            if not s or not creative:
                return s
            with_newlines = re.sub(r'\s+(ক\.|খ\.|গ\.|ঘ\.)', r'\n\1', s)
            with_newlines = re.sub(r'([।,])\s*(ক\.|খ\.|গ\.|ঘ\.)', r'\1\n\2', with_newlines)
            return (
                with_newlines
                .replace('ক.', '(ক)')
                .replace('খ.', '(খ)')
                .replace('গ.', '(গ)')
                .replace('ঘ.', '(ঘ)')
            )

        def question_display_structure(raw_text, creative):
            full = question_display_text(raw_text, creative)
            if not creative or not full:
                return {'intro': full, 'parts': []}
            lines = [ln.strip() for ln in full.split('\n') if ln.strip()]
            if len(lines) <= 1:
                return {'intro': full, 'parts': []}
            return {'intro': lines[0], 'parts': lines[1:]}

        def to_bengali_digits(n):
            return str(n).translate(str.maketrans('0123456789', '০১২৩৪৫৬৭৮৯'))

        serial_raw = pick('previewSerialByIndex', {})
        serial_by_index = serial_raw if isinstance(serial_raw, dict) else {}

        def item_serial(item_idx, fallback_num):
            if item_idx is None:
                return to_bengali_digits(fallback_num)
            s = serial_by_index.get(str(item_idx))
            if s is None:
                s = serial_by_index.get(item_idx)
            try:
                sn = int(s)
                if sn > 0:
                    return to_bengali_digits(sn)
            except Exception:
                pass
            return to_bengali_digits(fallback_num)

        def render_items_html(items, start_num=1):
            out = []
            for idx, q in enumerate(items):
                i = start_num + idx
                item_idx = None
                if isinstance(q, dict) and 'q' in q:
                    try:
                        item_idx = int(q.get('idx'))
                    except Exception:
                        item_idx = None
                    qq = q.get('q') if isinstance(q.get('q'), dict) else {}
                else:
                    qq = q if isinstance(q, dict) else {}
                creative = is_creative(qq)
                qcls = 'q-item q-cq' if creative else 'q-item q-mcq'
                fz = q_font_cq if creative else q_font_mcq
                q_lh = q_lh_cq if creative else q_lh_mcq
                q_gap = q_gap_cq if creative else q_gap_mcq
                style = (
                    'font-size: %.2fpx; '
                    '--preview-question-lh: %.3f; '
                    '--preview-q-bn-paren-inset: %.2fpx; '
                    '--preview-q-subpart-pl: %.2fpx; '
                    'padding-top: %.2fpx; '
                    'padding-bottom: %.2fpx; '
                    'margin-bottom: %.2fpx;'
                ) % (fz, q_lh, 2 * fz - 2, 2 * fz - 4, q_pad, q_pad, q_gap)

                struct = question_display_structure(qq.get('question') or '', creative)
                intro_html = wrap_roman_lines_html(struct.get('intro') or '')
                if struct.get('parts'):
                    parts_html = ''.join(
                        '<div class="q-subpart">%s</div>' % wrap_roman_lines_html(p)
                        for p in struct.get('parts') or []
                    )
                    stem_html = (
                        '<span class="q-stem-with-parts">'
                        '<span class="q-text q-intro">%s</span>'
                        '</span>%s'
                    ) % (intro_html, parts_html)
                else:
                    stem_html = '<span class="q-text">%s</span>' % intro_html

                options = []
                if qq.get('option_1') or qq.get('option_2'):
                    for key, lab in [('option_1', '(ক)'), ('option_2', '(খ)'), ('option_3', '(গ)'), ('option_4', '(ঘ)')]:
                        ov = qq.get(key)
                        if ov:
                            txt = str(ov).strip()
                            if txt:
                                options.append('<span class="q-opt">%s %s</span>' % (lab, escape(txt)))
                opts_html = '<div class="q-options">%s</div>' % ''.join(options) if options else ''
                out.append(
                    '<div class="%s" style="%s"><div class="q-content"><label class="q-label">'
                    '<strong class="qn">%s।</strong>%s</label>%s</div></div>'
                    % (qcls, style, item_serial(item_idx, i), stem_html if stem_html else '&nbsp;', opts_html)
                )
            return ''.join(out)

        def render_fixed_columns_html(columns_items, include_header_in_first_col=False, hdr_html=None):
            cols = columns_items if isinstance(columns_items, list) and columns_items else [[]]
            total = len(cols)
            chunks = []
            for ci, col_items in enumerate(cols):
                col = col_items if isinstance(col_items, list) else []
                col_html = render_items_html(col, 1)
                if include_header_in_first_col and ci == 0 and hdr_html:
                    col_html = (
                        hdr_html.replace('class="q-header"', 'class="q-header q-header--landscape-first-col"', 1)
                        + col_html
                    )
                rule_cls = ' q-col--rule' if show_div and total > 1 and ci < total - 1 else ''
                chunks.append('<div class="q-col%s">%s</div>' % (rule_cls, col_html))
            # Grid + explicit column count matches preview column-major packing. Flex + align-items:stretch
            # made both columns the height of the taller one and caused bad print fragmentation (col1
            # spilling to page 2 while col2 still had room on page 1).
            ncols = max(1, total)
            return (
                '<div class="q-wrap-fixed" style="grid-template-columns: repeat(%d, minmax(0, 1fr));">%s</div>'
                % (ncols, ''.join(chunks))
            )

        cq_items_html = render_items_html(creative_questions, 1)
        mcq_items_html = render_items_html(mcq_questions, 1)
        lead_empty_first_page = bool(pick('leadEmptyFirstPageActive', False))
        raw_lead_binding = pick('leadBindingItemIndexes', [])
        lead_binding_idx = set()
        if isinstance(raw_lead_binding, list):
            for x in raw_lead_binding:
                try:
                    v = int(x)
                    if v >= 0:
                        lead_binding_idx.add(v)
                except Exception:
                    continue
        first_section_kind = 'cq' if creative_questions else 'mcq'
        lead_empty_cq = bool(
            lead_empty_first_page
            and first_section_kind == 'cq'
            and cols_cq > 1
            and cq_orient in ('landscape', 'l', 'land')
        )
        lead_empty_mcq = bool(lead_empty_first_page and cols_mcq > 1 and mcq_orient in ('landscape', 'l', 'land'))
        if first_section_kind != 'mcq':
            lead_empty_mcq = False
        cq_render_cols = (cols_cq - 1) if lead_empty_cq else cols_cq
        mcq_render_cols = (cols_mcq - 1) if lead_empty_mcq else cols_mcq

        creative_by_idx = {int(it.get('idx')): it for it in creative_questions if isinstance(it, dict) and it.get('idx') is not None}
        mcq_by_idx = {int(it.get('idx')): it for it in mcq_questions if isinstance(it, dict) and it.get('idx') is not None}
        plan_raw = pick('exportPreviewPagePlan', [])
        use_plan = isinstance(plan_raw, list) and len(plan_raw) > 0 and page_sections <= 1
        sections = []
        if use_plan:
            for pi, pg in enumerate(plan_raw):
                if not isinstance(pg, dict):
                    continue
                kind = str(pg.get('kind') or '').strip().lower()
                if kind not in ('creative', 'mcq'):
                    continue
                hdr_for_page = header_html_creative if kind == 'creative' else header_html_mcq
                pool = creative_by_idx if kind == 'creative' else mcq_by_idx
                cols_idx = pg.get('questionColumnIndexes')
                lead_idx = pg.get('leadBindingIndexes')
                lead_empty = bool(pg.get('leadEmpty'))
                header_visible = bool(pg.get('headerVisible', True))
                header_in_col = bool((pg.get('headerInFirstColumn') or lead_empty) and header_visible)
                if not isinstance(cols_idx, list):
                    cols_idx = []
                columns_items = []
                for col in cols_idx:
                    if not isinstance(col, list):
                        continue
                    col_items = []
                    for x in col:
                        try:
                            iv = int(x)
                        except Exception:
                            continue
                        it = pool.get(iv)
                        if it is not None:
                            col_items.append(it)
                    columns_items.append(col_items)
                if not columns_items:
                    columns_items = [[]]
                main_html = render_fixed_columns_html(columns_items, include_header_in_first_col=header_in_col, hdr_html=hdr_for_page)
                if lead_empty and isinstance(lead_idx, list):
                    lead_items = []
                    for x in lead_idx:
                        try:
                            iv = int(x)
                        except Exception:
                            continue
                        it = pool.get(iv)
                        if it is not None:
                            lead_items.append(it)
                    lead_html = '<div class="lead-empty-col-content">%s</div>' % render_items_html(lead_items, 1)
                    body_html = (
                        '<div class="lead-empty-grid%s" style="--lead-cols:%d; --lead-main-cols:%d; --lead-gap:%dpx;">'
                        '<div class="lead-empty-col">%s</div>'
                        '<div class="lead-empty-main">%s</div>'
                        '</div>'
                    ) % (
                        ' lead-empty-grid--ruled' if show_div and len(columns_items) > 0 else '',
                        len(columns_items) + 1,
                        len(columns_items),
                        col_gap,
                        lead_html,
                        main_html,
                    )
                    header_html_page = ''
                else:
                    body_html = main_html
                    header_html_page = '' if (header_in_col or not header_visible) else hdr_for_page
                paper_cls = 'paper-cq' if kind == 'creative' else 'paper-mcq'
                break_cls = ' paper-break' if pi > 0 else ''
                sections.append('<section class="paper %s%s">%s%s</section>' % (paper_cls, break_cls, header_html_page, body_html))
        else:
            if creative_questions:
                creative_main_items = creative_questions
                creative_lead_items = []
                if lead_empty_cq and lead_binding_idx:
                    creative_lead_items = [it for it in creative_questions if int(it.get('idx', -1)) in lead_binding_idx]
                    if creative_lead_items:
                        creative_main_items = [it for it in creative_questions if int(it.get('idx', -1)) not in lead_binding_idx]
                cq_items_html = render_items_html(creative_main_items, 1)
                cq_body_html = '<div class="q-wrap q-wrap-cq">%s</div>' % cq_items_html
                if lead_empty_cq:
                    lead_header_html = header_html_creative.replace(
                        'class="q-header"',
                        'class="q-header q-header--lead-first-col"',
                        1
                    )
                    lead_binding_html = ''
                    if creative_lead_items:
                        lead_binding_html = '<div class="lead-empty-col-content">%s</div>' % render_items_html(creative_lead_items, 1)
                    cq_body_html = (
                        '<div class="lead-empty-grid%s" style="--lead-cols:%d; --lead-main-cols:%d; --lead-gap:%dpx;">'
                        '<div class="lead-empty-col" aria-hidden="true">%s</div>'
                        '<div class="lead-empty-main">%s%s</div>'
                        '</div>'
                    ) % (
                        ' lead-empty-grid--ruled' if show_div and cols_cq > 1 else '',
                        cols_cq,
                        cq_render_cols,
                        col_gap,
                        lead_binding_html,
                        lead_header_html,
                        cq_body_html,
                    )
                    cq_header_html = ''
                else:
                    cq_header_html = header_html_creative
                sections.append(
                    '<section class="paper paper-cq">%s%s</section>'
                    % (cq_header_html, cq_body_html)
                )
            if mcq_questions:
                mcq_main_items = mcq_questions
                mcq_lead_items = []
                if lead_empty_mcq and lead_binding_idx:
                    mcq_lead_items = [it for it in mcq_questions if int(it.get('idx', -1)) in lead_binding_idx]
                    if mcq_lead_items:
                        mcq_main_items = [it for it in mcq_questions if int(it.get('idx', -1)) not in lead_binding_idx]
                mcq_items_html = render_items_html(mcq_main_items, 1)
                mcq_body_html = '<div class="q-wrap q-wrap-mcq">%s</div>' % mcq_items_html
                if lead_empty_mcq:
                    lead_header_html = header_html_mcq.replace(
                        'class="q-header"',
                        'class="q-header q-header--lead-first-col"',
                        1
                    )
                    lead_binding_html = ''
                    if mcq_lead_items:
                        lead_binding_html = '<div class="lead-empty-col-content">%s</div>' % render_items_html(mcq_lead_items, 1)
                    mcq_body_html = (
                        '<div class="lead-empty-grid%s" style="--lead-cols:%d; --lead-main-cols:%d; --lead-gap:%dpx;">'
                        '<div class="lead-empty-col" aria-hidden="true">%s</div>'
                        '<div class="lead-empty-main">%s%s</div>'
                        '</div>'
                    ) % (
                        ' lead-empty-grid--ruled' if show_div and cols_mcq > 1 else '',
                        cols_mcq,
                        mcq_render_cols,
                        col_gap,
                        lead_binding_html,
                        lead_header_html,
                        mcq_body_html,
                    )
                    mcq_header_html = ''
                else:
                    mcq_header_html = header_html_mcq
                sections.append(
                    '<section class="paper paper-mcq%s">%s%s</section>'
                    % (' paper-break' if creative_questions else '', mcq_header_html, mcq_body_html)
                )
            if not sections:
                sections.append('<section class="paper paper-mcq"><div class="q-wrap q-wrap-mcq"></div></section>')

        divider_css_mcq = 'column-rule: 1px solid #c8c8c8;' if show_div and mcq_render_cols > 1 else ''
        divider_css_cq = 'column-rule: 1px solid #c8c8c8;' if show_div and cq_render_cols > 1 else ''
        local_font_face_css = self._playwright_local_font_face_css()

        html = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <style>
    {local_font_face_css}
    @page default {{
      size: {float(page_w_mm):.3f}mm {float(page_h_mm):.3f}mm;
      margin: {float(margin_top):.3f}mm {float(margin_right):.3f}mm {float(margin_bottom):.3f}mm {float(margin_left):.3f}mm;
    }}
    @page cq {{
      size: {float(cq_w_mm):.3f}mm {float(cq_h_mm):.3f}mm;
      margin: {float(margin_top):.3f}mm {float(margin_right):.3f}mm {float(margin_bottom):.3f}mm {float(margin_left):.3f}mm;
    }}
    @page mcq {{
      size: {float(mcq_w_mm):.3f}mm {float(mcq_h_mm):.3f}mm;
      margin: {float(margin_top):.3f}mm {float(margin_right):.3f}mm {float(margin_bottom + mcq_extra_bottom_mm):.3f}mm {float(margin_left):.3f}mm;
    }}
    html, body {{ margin: 0; padding: 0; }}
    body {{
      font-family: "Roboto", sans-serif;
      font-size: {q_font_global:.2f}px;
      line-height: {q_lh_global:.3f};
      color: #111;
    }}
    .paper {{ page: default; }}
    .paper-cq {{ page: cq; }}
    .paper-mcq {{ page: mcq; }}
    .paper-break {{ break-before: page; page-break-before: always; }}
    .q-header {{ margin: 0 0 8px 0; text-align: center; }}
    .hline-hr {{
      border: 0;
      border-top: 1px solid #222;
      margin: 6px 0;
    }}
    .hline-code-row-wrap {{
      display: inline-flex;
      flex-wrap: wrap;
      align-items: center;
      justify-content: flex-end;
      width: 100%;
      margin-bottom: 7px;
      box-sizing: border-box;
    }}
    .q-header-band {{
      position: relative;
      width: 100%;
      box-sizing: border-box;
    }}
    .hline-code-row-wrap--floating {{
      position: absolute;
      right: 0;
      bottom: 7px;
      width: auto;
      max-width: 100%;
      margin-bottom: 0;
      z-index: 2;
      pointer-events: none;
    }}
    .q-code-grid {{
      display: grid;
      grid-template-columns: repeat(5, auto);
      align-items: center;
      justify-items: center;
      border: 1px solid #333;
      padding: 6px 6px;
      box-sizing: border-box;
    }}
    .q-code-label {{ white-space: nowrap; }}
    .q-code-colon {{ padding: 0 2px; white-space: nowrap; }}
    .q-code-cell {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      min-width: 1.35em;
      padding: 3px 8px;
      text-align: center;
      box-sizing: border-box;
      border: 1px solid #333;
    }}
    .q-code-grid .q-code-cell:nth-child(3),
    .q-code-grid .q-code-cell:nth-child(4) {{
      margin-right: -3px;
    }}
    /* MCQ: second row — সেট : (letter), matches preview two-row code table. */
    .q-code-grid--mcq-set {{
      grid-template-rows: auto auto;
      row-gap: 15px;
    }}
    .q-code-label--set {{ white-space: nowrap; }}
    .q-code-cell--filler {{
      border: none !important;
      min-width: 0.35em;
      padding: 0;
      background: transparent;
    }}
    .q-code-cell--set-letter {{
      border-width: 2px;
      min-width: 1.35em;
      padding: 3px 5px;
    }}
    /* Creative: no outer frame around the subject-code row (digits keep cell borders). */
    .paper-cq .q-code-grid {{
      border: none;
      padding: 0;
    }}
    .q-header--lead-first-col {{
      width: calc(
        (100% - ((var(--lead-main-cols, 1) - 1) * var(--lead-gap, {col_gap}px)))
        / var(--lead-main-cols, 1)
      );
      max-width: 100%;
      margin-left: 0;
      margin-right: 0;
    }}
    .hline {{ margin: 0; }}
    .lead-empty-grid {{
      display: grid;
      grid-template-columns: repeat({cols_mcq}, minmax(0, 1fr));
      gap: {col_gap}px;
      align-items: stretch;
      position: relative;
    }}
    .lead-empty-col {{
      min-height: 100%;
    }}
    .lead-empty-col-content {{
      min-width: 0;
      width: 100%;
    }}
    .q-wrap-fixed {{
      display: grid;
      gap: {col_gap}px;
      align-items: start;
      width: 100%;
      box-sizing: border-box;
    }}
    .q-col {{
      min-width: 0;
      min-height: 0;
      display: flex;
      flex-direction: column;
      align-items: stretch;
      align-self: start;
      box-sizing: border-box;
    }}
    .q-col--rule {{
      border-right: 1px solid #c8c8c8;
    }}
    .q-header--landscape-first-col {{
      flex-shrink: 0;
      width: 100%;
      margin-bottom: 8px;
      box-sizing: border-box;
    }}
    .lead-empty-main {{
      grid-column: 2 / -1;
      min-width: 0;
      position: relative;
      z-index: 1;
    }}
    /* Single divider between lead column and question block (same weight/color as .q-col--rule / column-rule). */
    .lead-empty-grid--ruled .lead-empty-col {{
      border-right: 1px solid #c8c8c8;
      box-sizing: border-box;
      position: relative;
      z-index: 2;
    }}
    .q-wrap {{
      column-count: {cols_mcq};
      column-gap: {col_gap}px;
      {divider_css_mcq}
    }}
    .q-wrap-cq {{
      column-count: {cq_render_cols};
      column-gap: {col_gap}px;
      {divider_css_cq}
    }}
    .q-wrap-mcq {{
      column-count: {mcq_render_cols};
      column-gap: {col_gap}px;
      {divider_css_mcq}
    }}
    .q-item {{
      break-inside: auto;
      page-break-inside: auto;
    }}
    .q-content {{
      min-width: 0;
      width: 100%;
      text-align: justify;
      position: relative;
      padding-right: 2px;
      box-sizing: border-box;
      break-inside: auto;
      page-break-inside: auto;
    }}
    .q-label {{ display: flow-root; margin: 0; }}
    .qn {{
      float: left;
      font-weight: 700;
      margin-right: 0.2em;
      line-height: var(--preview-question-lh, 1.4);
    }}
    .q-text {{
      display: block;
      overflow: hidden;
      min-width: 0;
      font-size: 1em;
      line-height: var(--preview-question-lh, 1.4);
      color: #333;
    }}
    .topic-question-line {{ display: block; box-sizing: border-box; }}
    .topic-question-line.topic-question-roman-line {{
      padding-left: 10px;
      text-indent: -10px;
    }}
    .topic-question-line.topic-question-bn-paren-line {{
      padding-left: var(--preview-q-bn-paren-inset, 18px);
      text-indent: calc(0px - var(--preview-q-bn-paren-inset, 18px));
    }}
    .q-stem-with-parts {{ display: flow-root; margin-bottom: 4px; }}
    .q-intro {{ display: block; overflow: hidden; min-width: 0; }}
    .q-subpart {{
      padding-left: var(--preview-q-subpart-pl, 14px);
      font-size: 1em;
      line-height: var(--preview-question-lh, 1.4);
      color: #333;
      margin-top: 2px;
      box-sizing: border-box;
      clear: both;
    }}
    .q-options {{
      margin-top: 3px;
      margin-bottom: 3px;
      margin-left: 0;
      padding-left: var(--preview-q-subpart-pl, 14px);
      box-sizing: border-box;
      font-size: calc(13 / 14 * 1em);
      line-height: var(--preview-question-lh, 1.4);
      color: #555;
      display: grid;
      grid-template-columns: repeat({options_cols}, minmax(0, 1fr));
      gap: 4px 1.5em;
      align-items: start;
      justify-content: start;
      break-inside: auto;
      page-break-inside: auto;
    }}
    .q-opt {{
      display: block;
      white-space: normal;
      min-width: 0;
      padding-left: 16px;
      text-indent: -16px;
      box-sizing: border-box;
      text-align: left;
    }}
  </style>
</head>
<body>
  {''.join(sections)}
</body>
</html>"""
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            try:
                page = browser.new_page()
                page.set_content(html, wait_until='networkidle')
                pdf_bytes = page.pdf(
                    print_background=True,
                    prefer_css_page_size=True,
                    margin={'top': '0mm', 'right': '0mm', 'bottom': '0mm', 'left': '0mm'},
                )
            finally:
                browser.close()
        buf = BytesIO()
        buf.write(pdf_bytes)
        buf.seek(0)
        return buf

    def _build_pdf(
        self,
        questions,
        question_header,
        margin_top,
        margin_right,
        margin_bottom,
        margin_left,
        page_w_mm,
        page_h_mm,
        layout_columns=1,
        layout_column_gap_px=14,
        show_column_divider=True,
        page_sections=1,
        section_gap_px=14,
    ):
        layout_columns = max(1, min(10, int(layout_columns or 1)))
        page_sections = max(1, min(10, int(page_sections or 1)))
        gap_pt = _pdf_px_to_pt(max(1, min(100, int(layout_column_gap_px or 14))))
        section_gap_pt = _pdf_px_to_pt(max(1, min(100, int(section_gap_px or 14))))
        width_pt, height_pt = _export_page_size_pt_from_mm(page_w_mm, page_h_mm)
        buf = BytesIO()
        c = pdf_canvas.Canvas(buf, pagesize=(width_pt, height_pt))
        w_pt = width_pt
        h_pt = height_pt
        left = margin_left * rl_mm
        bottom_effective = margin_bottom * rl_mm
        right = w_pt - (margin_right * rl_mm)
        top_effective = h_pt - (margin_top * rl_mm)
        pdf_font = _get_pdf_bengali_font() or 'Helvetica'
        c.setFont(pdf_font, 12)
        content_w = right - left
        if layout_columns <= 1:
            col_w = content_w
            ncols = 1
        else:
            ncols = layout_columns
            col_w = (content_w - gap_pt * (ncols - 1)) / ncols

        def paint_overlays():
            if show_column_divider and ncols > 1:
                _pdf_paint_column_dividers(c, left, bottom_effective, top_effective, col_w, gap_pt, ncols)
            if page_sections > 1:
                _pdf_paint_horizontal_section_dividers(
                    c, left, right, bottom_effective, top_effective, page_sections, section_gap_pt
                )

        if ncols == 1:
            paint_overlays()
            y = top_effective
            if question_header:
                c.drawString(left, y, question_header[:200])
                y -= 18
            for i, q in enumerate(questions):
                qtext = (q.get('question') or '').strip() or ' '
                line = '%s. %s' % (i + 1, qtext[:500])
                if y < bottom_effective + 20:
                    c.showPage()
                    c.setFont(pdf_font, 12)
                    paint_overlays()
                    y = top_effective
                c.drawString(left, y, line[:100])
                y -= 14
                opts = [q.get('option_1'), q.get('option_2'), q.get('option_3'), q.get('option_4')]
                opts = [str(o).strip() for o in opts if o]
                for opt in opts[:4]:
                    if y < bottom_effective + 14:
                        c.showPage()
                        c.setFont(pdf_font, 12)
                        paint_overlays()
                        y = top_effective
                    c.drawString(left + 20, y, (opt[:80]))
                    y -= 14
                y -= 8
            c.save()
            buf.seek(0)
            return buf

        if question_header:
            c.drawString(left, top_effective, question_header[:200])
            start_y = top_effective - 18
        else:
            start_y = top_effective
        ys = [start_y] * ncols

        def x_for_col(ci):
            return left + ci * (col_w + gap_pt)

        paint_overlays()

        qi = 0
        _guard = 0
        while qi < len(questions):
            _guard += 1
            if _guard > 10000:
                break
            ci = max(range(ncols), key=lambda j: ys[j])
            x0 = x_for_col(ci)
            y = ys[ci]
            q = questions[qi]
            qtext = (q.get('question') or '').strip() or ' '
            line = '%s. %s' % (qi + 1, qtext[:500])
            restart = False
            if y < bottom_effective + 20:
                c.showPage()
                c.setFont(pdf_font, 12)
                ys = [top_effective] * ncols
                paint_overlays()
                restart = True
            if restart:
                continue
            c.drawString(x0, y, line[:100])
            y -= 14
            opts = [q.get('option_1'), q.get('option_2'), q.get('option_3'), q.get('option_4')]
            opts = [str(o).strip() for o in opts if o]
            for opt in opts[:4]:
                if y < bottom_effective + 14:
                    c.showPage()
                    c.setFont(pdf_font, 12)
                    ys = [top_effective] * ncols
                    paint_overlays()
                    restart = True
                    break
                c.drawString(x0 + 20, y, opt[:80])
                y -= 14
            if restart:
                continue
            y -= 8
            ys[ci] = y
            qi += 1
        c.save()
        buf.seek(0)
        return buf

    def _build_docx(
        self,
        questions,
        question_header,
        margin_top,
        margin_right,
        margin_bottom,
        margin_left,
        page_w_mm,
        page_h_mm,
        layout_columns=1,
        layout_column_gap_px=14,
        show_column_divider=True,
    ):
        doc = DocxDocument()
        section = doc.sections[0]
        w_mm, h_mm = float(page_w_mm), float(page_h_mm)
        section.page_width = DocxMm(w_mm)
        section.page_height = DocxMm(h_mm)
        section.top_margin = DocxMm(margin_top)
        section.right_margin = DocxMm(margin_right)
        section.bottom_margin = DocxMm(margin_bottom)
        section.left_margin = DocxMm(margin_left)
        gap_px = max(1, min(100, int(layout_column_gap_px or 14)))
        space_twips = int(round(gap_px * 15))
        _docx_apply_section_columns(
            section, layout_columns, space_twips=space_twips, show_sep=bool(show_column_divider)
        )
        if question_header:
            p = doc.add_paragraph(question_header)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, q in enumerate(questions):
            qtext = (q.get('question') or '').strip() or ' '
            doc.add_paragraph('%s. %s' % (i + 1, qtext))
            for key, label in [('option_1', 'A.'), ('option_2', 'B.'), ('option_3', 'C.'), ('option_4', 'D.')]:
                opt = q.get(key)
                if opt:
                    doc.add_paragraph('   %s %s' % (label, str(opt).strip()), style='List Bullet')
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


class ExportQuestionsBulkView(APIView):
    """POST: body { items: [ { questions, questionHeader, filename, pageSize?, marginTop?, ... }, ... ] }. Returns a single ZIP of all PDFs. One download, no repeated prompts."""
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data
        items = data.get('items')
        if not isinstance(items, list) or len(items) == 0:
            return Response({'error': 'items must be a non-empty list'}, status=status.HTTP_400_BAD_REQUEST)
        if not REPORTLAB_AVAILABLE:
            return Response({'error': 'PDF generation not available'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
        exporter = ExportQuestionsView()
        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for i, item in enumerate(items):
                questions = item.get('questions') or []
                if not isinstance(questions, list):
                    continue
                question_header = (item.get('questionHeader') or '')[:255]
                filename_base = (item.get('filename') or 'questions_%s' % (i + 1)).strip()[:120]
                filename_base = re.sub(r'[^\w\-_.\s]', '_', filename_base)
                margin_top = float(item.get('marginTop') or 25.4)
                margin_right = float(item.get('marginRight') or 25.4)
                margin_bottom = float(item.get('marginBottom') or 25.4)
                margin_left = float(item.get('marginLeft') or 25.4)
                w_mm, h_mm = _export_resolve_page_mm(item)
                lc = max(1, min(10, int(item.get('layoutColumns') or item.get('layout_columns') or 1)))
                gap_px = _export_layout_column_gap_px(item)
                show_div = _export_show_column_divider(item)
                page_sec = _export_page_sections(item)
                sec_gap_px = _export_layout_section_gap_px(item)
                pdf_buf = exporter._build_pdf(
                    questions, question_header,
                    margin_top, margin_right, margin_bottom, margin_left,
                    page_w_mm=w_mm,
                    page_h_mm=h_mm,
                    layout_columns=lc,
                    layout_column_gap_px=gap_px,
                    show_column_divider=show_div,
                    page_sections=page_sec,
                    section_gap_px=sec_gap_px,
                )
                zf.writestr(filename_base + '.pdf', pdf_buf.getvalue())
        zip_buf.seek(0)
        resp = HttpResponse(zip_buf.getvalue(), content_type='application/zip')
        resp['Content-Disposition'] = 'attachment; filename="created_questions_all.zip"'
        return resp


class CreatedQuestionSetListCreateView(APIView):
    """GET: list current user's created question sets. POST: create one (name, question_header, questions). Counter is per-name: same name gets _1, _2, _3."""
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            qs = CreatedQuestionSet.objects.filter(customer=request.user).order_by('-created_at')
        except Exception as e:
            logger.exception('CreatedQuestionSet list failed: %s', e)
            return Response(
                {'error': 'Could not load created questions. Run migrations: python manage.py migrate'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        data = []
        for o in qs:
            try:
                created_at = o.created_at.isoformat() if getattr(o, 'created_at', None) else None
            except Exception:
                created_at = None
            ls = getattr(o, 'layout_settings', None)
            if not isinstance(ls, dict):
                ls = {}
            data.append({
                'id': o.id,
                'name': o.name,
                'question_header': o.question_header or '',
                'questions': o.questions if isinstance(o.questions, list) else [],
                'counter': o.counter,
                'file_name_base': '%s_%s' % (str(o.name).replace(' ', '_'), o.counter),
                'created_at': created_at,
                'layout_settings': ls,
            })
        return Response(data, status=status.HTTP_200_OK)

    def post(self, request):
        name = (request.data.get('name') or '').strip() or 'questions'
        name = name[:200]
        question_header = (request.data.get('question_header') or '')[:255]
        questions = request.data.get('questions')
        if not isinstance(questions, list):
            return Response({'error': 'questions must be a list'}, status=status.HTTP_400_BAD_REQUEST)
        layout_settings = request.data.get('layout_settings')
        if not isinstance(layout_settings, dict):
            layout_settings = {}
        from django.db.models import Max
        # Per-name counter: same name gets _1, _2, _3, ...
        next_counter = (
            CreatedQuestionSet.objects.filter(customer=request.user, name=name).aggregate(Max('counter'))['counter__max'] or 0
        ) + 1
        try:
            obj = CreatedQuestionSet.objects.create(
                customer=request.user,
                name=name,
                question_header=question_header,
                questions=questions,
                counter=next_counter,
                layout_settings=layout_settings,
            )
        except Exception as e:
            logger.exception('CreatedQuestionSet create failed: %s', e)
            return Response(
                {'error': 'Could not save. Run migrations: python manage.py migrate'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        created_at = obj.created_at.isoformat() if getattr(obj, 'created_at', None) else None
        return Response({
            'id': obj.id,
            'name': obj.name,
            'question_header': obj.question_header,
            'counter': obj.counter,
            'file_name_base': '%s_%s' % (obj.name.replace(' ', '_'), obj.counter),
            'created_at': created_at,
            'layout_settings': obj.layout_settings if isinstance(getattr(obj, 'layout_settings', None), dict) else {},
        }, status=status.HTTP_201_CREATED)


class CreatedQuestionSetDetailView(APIView):
    """GET: one set. PATCH: rename (name). DELETE: remove."""
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def _get_obj(self, request, pk):
        try:
            return CreatedQuestionSet.objects.get(pk=pk, customer=request.user)
        except CreatedQuestionSet.DoesNotExist:
            return None

    def get(self, request, pk):
        obj = self._get_obj(request, pk)
        if not obj:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
        ls = getattr(obj, 'layout_settings', None)
        if not isinstance(ls, dict):
            ls = {}
        return Response({
            'id': obj.id,
            'name': obj.name,
            'question_header': obj.question_header,
            'questions': obj.questions,
            'counter': obj.counter,
            'file_name_base': '%s_%s' % (obj.name.replace(' ', '_'), obj.counter),
            'created_at': obj.created_at.isoformat() if obj.created_at else None,
            'layout_settings': ls,
        }, status=status.HTTP_200_OK)

    def patch(self, request, pk):
        obj = self._get_obj(request, pk)
        if not obj:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
        new_name = request.data.get('name')
        if new_name is not None:
            obj.name = (str(new_name).strip() or obj.name)[:200]
            obj.save(update_fields=['name'])
        return Response({
            'id': obj.id,
            'name': obj.name,
            'file_name_base': '%s_%s' % (obj.name.replace(' ', '_'), obj.counter),
        }, status=status.HTTP_200_OK)

    def delete(self, request, pk):
        obj = self._get_obj(request, pk)
        if not obj:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
        obj.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class CustomerUpdateView(APIView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(request, username=username, password=password)
        if user is None:
            return Response({'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)
        serializer = CustomerUpdateSerializer(user, data=request.data, partial=True)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        serializer.save()
        token = self.generate_unique_key()
        CustomerToken.objects.filter(customer=user).delete()
        CustomerToken.objects.create(key=token, customer=user)
        return Response({'authToken': token, 'token': token}, status=status.HTTP_200_OK)

    def generate_unique_key(self):
        length = 40
        characters = string.ascii_letters + string.digits
        return ''.join(random.choice(characters) for _ in range(length))

class CustomerResetView(APIView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None:
            serializer = CustomerSerializer(user, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response({'message': 'Password reset successfully'}, status=status.HTTP_200_OK)
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        return Response({'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)


class NotificationViewSet(viewsets.ReadOnlyModelViewSet):
    """List notifications for the header marquee. GET /api/notification/. Returns last 10 by id (newest first)."""
    serializer_class = NotificationSerializer
    pagination_class = None

    def get_queryset(self):
        return Notification.objects.all().order_by('-id')[:10]


class MobileNumberExistsView(APIView):
    """Check username (mobile) in Customer table. Optional country_code filter. Returns exists and found_in='customer' when found."""
    permission_classes = [AllowAny]
    authentication_classes = []

    def get(self, request, *args, **kwargs):
        username = request.query_params.get('username')
        country_code = (request.query_params.get('countryCode') or request.query_params.get('country_code') or '').strip().upper() or None

        if not username:
            return Response({'exists': False}, status=status.HTTP_200_OK)
        q = Q(username=username)
        if country_code:
            q &= Q(country_code=country_code)
        if Customer.objects.filter(q).exists():
            return Response({'exists': True, 'found_in': 'customer'}, status=status.HTTP_200_OK)
        return Response({'exists': False}, status=status.HTTP_200_OK)


class PasswordExistsView(APIView):
    """Check username+password in Customer table. Optional country_code filter."""
    permission_classes = [AllowAny]
    authentication_classes = []

    def get(self, request, *args, **kwargs):
        username = request.query_params.get('username')
        password = request.query_params.get('password')
        country_code = (request.query_params.get('countryCode') or request.query_params.get('country_code') or '').strip().upper() or None

        if not username or not password:
            return Response({'exists': False}, status=status.HTTP_200_OK)
        q = Q(username=username)
        if country_code:
            q &= Q(country_code=country_code)
        try:
            customer = Customer.objects.filter(q).first()
            if customer and customer.check_password(password):
                return Response({'exists': True}, status=status.HTTP_200_OK)
        except (ProgrammingError, OperationalError):
            pass
        return Response({'exists': False}, status=status.HTTP_200_OK)
    

@csrf_exempt
def save_json_data(request):
    """Store JSON payload in JsonData; optionally link/update Transaction by trxid+paidFrom."""
    if request.method != 'POST':
        return JsonResponse({'message': 'Invalid request method'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        trxid = body.pop('trxid', None)
        paidFrom = body.pop('paidFrom', None)
        username = body.get('username', None)
        transaction = None
        if trxid and paidFrom:
            try:
                transaction = Transaction.objects.get(trxid=trxid, paidFrom=paidFrom)
            except Transaction.DoesNotExist:
                pass
        if transaction and username and transaction.username != username:
            transaction.username = username
            transaction.save(update_fields=['username'])
        JsonData.objects.create(
            data=body,
            data_type='order_submission',
            description=username or trxid or 'save_json_data',
        )
        return JsonResponse({'message': 'Data saved successfully'}, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
def run_scraper(request):
    """Proxy scraper: fetch paginated API (from frontend config), return combined JSON."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        import time
        body = json.loads(request.body.decode('utf-8'))
        base_url = (body.get('base_url') or '').strip()
        params = body.get('params') or {}
        headers = body.get('headers') or {}
        page_param = (body.get('page_param') or 'page').strip() or 'page'
        delay_seconds = float(body.get('delay_seconds') or 1)
        delay_seconds = max(0, min(10, delay_seconds))
        chapter_name = body.get('chapter_name') or 'scraped_data'

        if not base_url:
            return JsonResponse({'error': 'base_url is required'}, status=400)

        # Normalize params: ensure page_param is not in params yet for first request
        req_params = dict(params) if isinstance(params, dict) else {}
        if not isinstance(params, dict):
            req_params = {}
            for item in params if isinstance(params, list) else []:
                k = item.get('key') or item.get('name')
                v = item.get('value')
                if k:
                    req_params[k] = v

        all_data = []
        page = 1
        while True:
            req_params[page_param] = page
            resp = requests.get(base_url, headers=headers, params=req_params, timeout=30)
            if resp.status_code != 200:
                return JsonResponse({
                    'error': f'HTTP {resp.status_code} at page {page}',
                    'data': all_data,
                    'chapter_name': chapter_name
                }, status=200)
            try:
                data = resp.json()
            except Exception:
                return JsonResponse({
                    'error': f'Non-JSON response at page {page}',
                    'data': all_data,
                    'chapter_name': chapter_name
                }, status=200)
            if data is None or (isinstance(data, list) and len(data) == 0) or (isinstance(data, dict) and not data):
                break
            all_data.append(data)
            time.sleep(delay_seconds)
            page += 1
            if page > 500:
                break

        return JsonResponse({'data': all_data, 'chapter_name': chapter_name})
    except json.JSONDecodeError as e:
        return JsonResponse({'error': f'Invalid JSON: {e}'}, status=400)
    except Exception as e:
        logger.exception('run_scraper failed')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def run_scraper_page(request):
    """POST: fetch a single page. Returns { data, has_more } for progress bar."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        import time
        body = json.loads(request.body.decode('utf-8'))
        base_url = (body.get('base_url') or '').strip()
        params = body.get('params') or {}
        headers = body.get('headers') or {}
        page_param = (body.get('page_param') or 'page').strip() or 'page'
        page_number = int(body.get('page_number') or 1)
        delay_seconds = float(body.get('delay_seconds') or 1)
        delay_seconds = max(0, min(10, delay_seconds))

        if not base_url:
            return JsonResponse({'error': 'base_url is required'}, status=400)

        req_params = dict(params) if isinstance(params, dict) else {}
        if not isinstance(params, dict):
            req_params = {}
            for item in params if isinstance(params, list) else []:
                k = item.get('key') or item.get('name')
                v = item.get('value')
                if k:
                    req_params[k] = v
        req_params[page_param] = page_number

        if page_number > 1:
            time.sleep(delay_seconds)

        resp = requests.get(base_url, headers=headers, params=req_params, timeout=30)
        if resp.status_code != 200:
            return JsonResponse({'error': f'HTTP {resp.status_code}', 'data': None, 'has_more': False})

        try:
            data = resp.json()
        except Exception:
            return JsonResponse({'error': 'Non-JSON response', 'data': None, 'has_more': False})

        has_more = data is not None and (
            (isinstance(data, list) and len(data) > 0) or
            (isinstance(data, dict) and bool(data))
        )
        return JsonResponse({'data': data, 'has_more': has_more})
    except json.JSONDecodeError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.exception('run_scraper_page failed')
        return JsonResponse({'error': str(e)}, status=500)


def _scraper_desktop_path():
    """Return user's Desktop directory (Windows: USERPROFILE\\Desktop, else ~/Desktop). Works on any computer."""
    home = os.environ.get('USERPROFILE') or os.environ.get('HOME') or os.path.expanduser('~')
    return os.path.join(home, 'Desktop')


def _scraper_root_config_path():
    """Path to project config file that stores custom scraper root (one line). Enables different root per computer."""
    return os.path.join(settings.BASE_DIR, 'scraper_root.txt')


def _scraper_get_root_override():
    """Return custom root folder from config file, or None to use default Desktop."""
    path = _scraper_root_config_path()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, 'r', encoding='utf-8') as f:
            line = (f.read() or '').strip()
        if line:
            return line
    except Exception:
        pass
    return None


def _scraper_base_path():
    """Scraper folder: {root}/Scraper. Root = custom from scraper_root.txt or Desktop. Creates folder if not exists."""
    root_override = _scraper_get_root_override()
    if root_override:
        base = os.path.join(root_override, 'Scraper')
    else:
        base = os.path.join(_scraper_desktop_path(), 'Scraper')
    try:
        os.makedirs(base, exist_ok=True)
    except Exception:
        pass
    return base


def _scraper_site_dir(website, group):
    """Return directory path Desktop/Scraper/{website}/{group}/. Creates Scraper and subdirs if needed. Group is uppercase (HSC, JSC)."""
    root = _scraper_base_path()
    safe_website = re.sub(r'[^\w\-]', '', (website or 'daricomma').strip()).lower() or 'daricomma'
    safe_group = (re.sub(r'[^\w\-]', '', (group or 'default').strip()) or 'default').upper()
    dir_path = os.path.join(root, safe_website, safe_group)
    try:
        os.makedirs(dir_path, exist_ok=True)
    except Exception:
        pass
    return dir_path


def _scraper_helper_json():
    """Path to Scraper/helper.json. Default: {root}/Scraper/helper.json (e.g. C:\\Users\\sasha\\Desktop\\Scraper\\helper.json).
    Tries: 1) SCRAPER_HELPER_JSON env (exact path), 2) {root}/Scraper/helper.json, 3) project BASE_DIR/helper.json."""
    env_path = (os.environ.get('SCRAPER_HELPER_JSON') or '').strip()
    if env_path and os.path.isfile(env_path):
        return os.path.normpath(env_path)
    path = os.path.normpath(os.path.join(_scraper_base_path(), 'helper.json'))
    if os.path.isfile(path):
        return path
    fallback = os.path.normpath(os.path.join(settings.BASE_DIR, 'helper.json'))
    if os.path.isfile(fallback):
        return fallback
    return path


def _scraper_api_txt():
    """Path to Desktop/Scraper/api.txt – one API URL per line (script-compatible). Do not clear on start; clear only when all tasks done."""
    return os.path.join(_scraper_base_path(), 'api.txt')


def _scraper_load_seen_apis(api_base_prefix=None):
    """Load already-used API URLs from Desktop/Scraper/api.txt. Do not clear file on new run (same as script)."""
    seen = set()
    path = _scraper_api_txt()
    if not os.path.isfile(path):
        return seen
    try:
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                url = line.strip()
                if not url or url.startswith('#'):
                    continue
                if api_base_prefix is None or url.startswith(api_base_prefix):
                    seen.add(url)
    except Exception:
        pass
    return seen


def _scraper_add_api_to_file(api_url, api_base_prefix=None):
    """Append one API URL to Desktop/Scraper/api.txt (script add_api_to_file)."""
    if not api_url or not isinstance(api_url, str):
        return
    api_url = api_url.strip()
    if api_base_prefix and not api_url.startswith(api_base_prefix):
        return
    try:
        _scraper_base_path()
        with open(_scraper_api_txt(), 'a', encoding='utf-8') as f:
            f.write(api_url + '\n')
    except Exception:
        pass


def _scraper_clear_api_file():
    """Remove all URLs from Desktop/Scraper/api.txt – call only after all tasks completed (script clear_api_file)."""
    try:
        _scraper_base_path()
        open(_scraper_api_txt(), 'w', encoding='utf-8').close()
    except Exception:
        pass


def _normalize_library_for_response(lib):
    """Ensure library dict has camelCase keys so frontend (loginUrl, apiBaseUrl, etc.) always receives them.
    Accepts both camelCase and snake_case keys from helper.json."""
    if not isinstance(lib, dict):
        return lib
    # (camelCase_key, snake_case_alias_or_same)
    key_map = [
        ('loginUrl', 'login_url'), ('username',), ('password',), ('groups',),
        ('apiBaseUrl', 'api_base_url'), ('apiUrlTemplate', 'api_url_template'),
        ('bearerToken', 'bearer_token'), ('questionPerPage', 'question_per_page'),
    ]
    out = {}
    for names in key_map:
        camel = names[0]
        val = None
        for key in names:
            if key in lib:
                val = lib[key]
                break
        if val is not None:
            out[camel] = val
        elif camel == 'groups':
            out[camel] = []
        elif camel == 'questionPerPage':
            out[camel] = 200
        else:
            out[camel] = ''
    if not isinstance(out.get('groups'), list):
        out['groups'] = [{'name': 'Default', 'urls': []}]
    return out


def _default_libraries():
    """Default library entries per site (daricomma, other). HSC group includes the 4 chapter URLs."""
    _HSC_CHAPTER_URLS = [
        'https://daricomma.com/academic/HSC%20-%20ICT/chapter/default&page=1',
        'https://daricomma.com/academic/HSC%20-%20%E0%A6%89%E0%A6%9A%E0%A7%8D%E0%A6%9A%E0%A6%A4%E0%A6%B0%20%E0%A6%97%E0%A6%A3%E0%A6%BF%E0%A6%A4%20%E0%A7%A7%E0%A6%AE%20%E0%A6%AA%E0%A6%A4%E0%A7%8D%E0%A6%B0/chapter/default&page=1',
        'https://daricomma.com/academic/HSC%20-%20%E0%A6%89%E0%A7%8E%E0%A6%AA%E0%A6%BE%E0%A6%A6%E0%A6%A8%20%E0%A6%AC%E0%A7%8D%E0%A6%AF%E0%A6%AC%E0%A6%B8%E0%A7%8D%E0%A6%A5%E0%A6%BE%E0%A6%AA%E0%A6%A8%E0%A6%BE%20%E0%A6%93%20%E0%A6%AC%E0%A6%BF%E0%A6%AA%E0%A6%A3%E0%A6%A8%20%E0%A7%A8%E0%A7%9F%20%E0%A6%AA%E0%A6%A4%E0%A7%8D%E0%A6%B0/chapter/default&page=1',
        'https://daricomma.com/academic/HSC%20-%20%E0%A6%85%E0%A6%B0%E0%A7%8D%E0%A6%A5%E0%A6%A8%E0%A7%80%E0%A6%A4%E0%A6%BF%20%E0%A7%A8%E0%A7%9F%20%E0%A6%AA%E0%A6%A4%E0%A7%8D%E0%A6%B0/chapter/default&page=1',
    ]
    _base = {
        'loginUrl': '',
        'username': '',
        'password': '',
        'groups': [{'name': 'Default', 'urls': []}],
        'apiBaseUrl': '',
        'apiUrlTemplate': '',
        'bearerToken': '',
        'questionPerPage': 200,
    }
    return {
        'daricomma': {
            'loginUrl': 'https://www.daricomma.com/sign-in',
            'username': '',
            'password': '',
            'groups': [{'name': 'HSC', 'urls': _HSC_CHAPTER_URLS}],
            'apiBaseUrl': 'https://api.daricomma.com/v2/question/',
            'apiUrlTemplate': 'https://api.daricomma.com/v2/question/7e93e529-3405-40ad-b003-895dacf21e9f',
            'bearerToken': '',
            'questionPerPage': 200,
        },
        'chorcha': {**_base},
        'eprosnobank': {**_base},
        'livemcq': {**_base},
        'other': {**_base},
    }


@csrf_exempt
def scraper_helper(request):
    """GET: return { lastSite, libraries: { daricomma: {...}, other: {...} } }. POST: save body { lastSite, libraries }."""
    if request.method == 'GET':
        try:
            _scraper_base_path()
            defaults = _default_libraries()
            with open(_scraper_helper_json(), 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Migrate old format { groups } to new format
            if 'libraries' not in data:
                groups = data.get('groups') if isinstance(data.get('groups'), list) else []
                data = {
                    'lastSite': data.get('lastSite', 'daricomma'),
                    'libraries': {
                        'daricomma': {**defaults['daricomma'], 'groups': groups or defaults['daricomma']['groups']},
                        'other': defaults['other'],
                    },
                }
            last_site = (data.get('lastSite') or 'daricomma').strip().lower() or 'daricomma'
            if last_site not in defaults:
                last_site = 'daricomma'
            libs_raw = data.get('libraries') or {}
            # Normalize keys to lowercase so frontend (sitePreset='daricomma') always finds the library
            libs = {}
            for k, v in libs_raw.items():
                if isinstance(v, dict):
                    libs[str(k).lower()] = v
            for key in defaults:
                if key not in libs or not isinstance(libs[key], dict):
                    libs[key] = defaults[key]
                else:
                    for k, v in defaults[key].items():
                        if k not in libs[key]:
                            libs[key][k] = v
            # If daricomma HSC group has no URLs, fill with default 4 chapter URLs
            if libs.get('daricomma') and isinstance(libs['daricomma'].get('groups'), list) and libs['daricomma']['groups']:
                first = libs['daricomma']['groups'][0]
                if isinstance(first, dict) and first.get('name') == 'HSC' and not (first.get('urls') or []):
                    first['urls'] = list(defaults['daricomma']['groups'][0]['urls'])
            # Normalize each library to camelCase so frontend always gets loginUrl, apiBaseUrl, bearerToken, apiUrlTemplate
            libs = {k: _normalize_library_for_response(v) for k, v in libs.items()}
            return JsonResponse({'lastSite': last_site, 'libraries': libs})
        except FileNotFoundError:
            return JsonResponse({'lastSite': 'daricomma', 'libraries': _default_libraries()})
        except Exception as e:
            logger.exception('scraper_helper GET failed')
            return JsonResponse({'error': str(e), 'lastSite': 'daricomma', 'libraries': _default_libraries()})
    if request.method == 'POST':
        try:
            body = json.loads(request.body.decode('utf-8'))
            last_site = (body.get('lastSite') or 'daricomma').strip() or 'daricomma'
            libraries = body.get('libraries')
            if not isinstance(libraries, dict):
                return JsonResponse({'error': 'libraries must be an object'}, status=400)
            defaults = _default_libraries()
            for key in defaults:
                if key not in libraries or not isinstance(libraries[key], dict):
                    libraries[key] = defaults[key]
            _scraper_base_path()
            with open(_scraper_helper_json(), 'w', encoding='utf-8') as f:
                json.dump({'lastSite': last_site, 'libraries': libraries}, f, ensure_ascii=False, indent=2)
            return JsonResponse({'ok': True})
        except (json.JSONDecodeError, TypeError) as e:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.exception('scraper_helper POST failed')
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Method not allowed'}, status=405)


@csrf_exempt
def scraper_clear_api_file(request):
    """POST: Clear Scrape/api.txt (call only after all tasks completed – same as script clear_api_file())."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        _scraper_clear_api_file()
        return JsonResponse({'ok': True})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def scraper_root_get(request):
    """GET: Return current scraper root folder (custom or default Desktop). Used so UI can show and change it."""
    if request.method != 'GET':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    root = _scraper_get_root_override()
    if not root:
        root = _scraper_desktop_path()
    return JsonResponse({'path': root})


@csrf_exempt
def scraper_default_root_get(request):
    """GET: Return default Desktop path for this machine. Use for 'Use default (Desktop)' button."""
    if request.method != 'GET':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    return JsonResponse({'path': _scraper_desktop_path()})


@csrf_exempt
def scraper_root_post(request):
    """POST: Set custom scraper root folder. Body: { path: "C:\\Users\\...\\Desktop" }. Creates scraper_root.txt in project."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        path = (body.get('path') or '').strip()
        if not path:
            config_path = _scraper_root_config_path()
            if os.path.isfile(config_path):
                try:
                    os.remove(config_path)
                except Exception:
                    pass
            return JsonResponse({'ok': True})
        config_path = _scraper_root_config_path()
        with open(config_path, 'w', encoding='utf-8') as f:
            f.write(path)
        return JsonResponse({'ok': True})
    except (json.JSONDecodeError, TypeError) as e:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def _scraper_safe_request(url, headers=None, params=None, max_retry=5, timeout=30):
    """Fetch URL with retries (matches script safe_request)."""
    headers = headers or {}
    params = params or {}
    for attempt in range(max_retry):
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=timeout)
            if resp.status_code == 200:
                return resp
            _scraper_progress('API error {}, retrying...'.format(resp.status_code))
        except Exception as e:
            _scraper_progress('Network error: {}, retrying...'.format(e))
        if attempt < max_retry - 1:
            time.sleep(3)
    return None


def _scraper_parse_academic_url(url):
    """Parse daricomma academic URL; return (subject_from_url, chapter_no, chapter_slug)."""
    try:
        if '/academic/' not in url or '/chapter/' not in url:
            return ('', '', '')
        parts = url.split('/chapter/', 1)
        if len(parts) != 2:
            return ('', '', '')
        subject_encoded = parts[0].rstrip('/').split('/academic/')[-1]
        rest = parts[1].split('&')[0].split('?')[0]
        subject_from_url = urllib_parse.unquote(subject_encoded)
        chapter_slug = urllib_parse.unquote(rest)
        chapter_no = _scraper_extract_chapter_no_from_slug(chapter_slug)
        return (subject_from_url, chapter_no, chapter_slug)
    except Exception:
        return ('', '', '')


def _scraper_extract_chapter_no_from_slug(chapter_slug):
    """Extract chapter number from URL slug (e.g. অধ্যায়-০১ঃ or 1.)."""
    if not chapter_slug or not isinstance(chapter_slug, str):
        return ''
    s = chapter_slug.strip()
    if 'অধ্যায়' in s:
        parts = s.split(None, 1)
        return parts[0].strip() if parts else (s.split()[0] if s.split() else '')
    m = re.match(r'^(\d+\.)', s)
    return m.group(1) if m else ''


def _scraper_extract_chapter_no_from_option_text(level2_text):
    """Get ChapterNo from Level2 dropdown option text (e.g. 'অধ্যায়-০১ঃ ম্যাট্রিক্স' or '1. Text Book')."""
    if not level2_text or not isinstance(level2_text, str):
        return ''
    s = level2_text.strip()
    if 'অধ্যায়' in s:
        parts = s.split(None, 1)
        return parts[0].strip() if parts else (s.split()[0] if s.split() else '')
    m = re.match(r'^(\d+\.)', s)
    return m.group(1) if m else ''


def _scraper_sanitize_filename(name):
    r"""Sanitize for file paths (matches script: replace / and \ with -)."""
    if not name:
        return ''
    return name.replace('/', '-').replace('\\', '-').strip()


@csrf_exempt
def scraper_fetch_and_save(request):
    """POST { group, base_url, params, headers, filename } - Fetch from API (with pagination, retry, dedupe) and save. Returns all questions in data."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        group = (body.get('group') or '').strip()
        website = (body.get('website') or body.get('site') or 'daricomma').strip()
        base_url = (body.get('base_url') or '').strip()
        params = body.get('params') or {}
        headers = body.get('headers') or {}
        filename = (body.get('filename') or 'data').strip()
        level1_name = (body.get('level1_name') or body.get('level1') or '').strip()
        level2_label = (body.get('level2_label') or body.get('level2') or '').strip()
        chapter_no = (body.get('chapter_no') or '').strip()
        session_id = (body.get('session_id') or '').strip()
        if not group or not base_url:
            return JsonResponse({'error': 'group and base_url are required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    if session_id and session_id in _scraper_aborted_sessions:
        return JsonResponse({'ok': False, 'error': 'Scraper stopped', 'data': []})
    dir_path = _scraper_site_dir(website, group)
    filename = _scraper_sanitize_filename(filename) or 'data'
    if not filename.endswith('.json'):
        filename = filename + '.json'
    file_path = os.path.join(dir_path, filename)
    if isinstance(params, dict):
        req_params = dict(params)
    else:
        req_params = {}
        for item in params if isinstance(params, list) else []:
            k = item.get('key') or item.get('name')
            v = item.get('value')
            if k:
                req_params[k] = v
    chapter_id = req_params.get('chapter_id', '')
    per_page = int(req_params.get('questionPerPage') or 200)
    if per_page <= 0:
        per_page = 200
    # Script: use Bearer token from the same browser session (localStorage) so API returns questions
    request_headers = dict(headers) if isinstance(headers, dict) else {}
    request_headers.setdefault('Accept', 'application/json')
    request_headers.setdefault('User-Agent', 'Mozilla/5.0')
    if session_id:
        driver = _scraper_sessions.get(session_id)
        if driver:
            try:
                local_storage = driver.execute_script('return window.localStorage;')
                if isinstance(local_storage, dict):
                    access_token = None
                    for k, v in local_storage.items():
                        if 'token' in (k or '').lower():
                            access_token = v
                            break
                    if access_token:
                        request_headers['Authorization'] = 'Bearer %s' % (access_token,)
                        _scraper_progress('Using token from session browser for API request.')
            except Exception:
                pass
    try:
        _scraper_progress('Fetching API: {}?chapter_id={} (paginated)'.format(
            base_url.rstrip('/'), chapter_id[:36] if chapter_id else ''))
        max_fetch_retries = 5
        wait_between_retries_sec = 10
        all_questions = []
        for fetch_attempt in range(max_fetch_retries):
            if session_id and session_id in _scraper_aborted_sessions:
                _scraper_progress('Scraper stopped by user.')
                return JsonResponse({'ok': False, 'error': 'Scraper stopped', 'data': []})
            all_questions = []
            seen_ids = set()
            total_questions = None
            page = 1
            page_retry = 0
            max_page_retry = 3
            while True:
                if session_id and session_id in _scraper_aborted_sessions:
                    _scraper_progress('Scraper stopped by user.')
                    return JsonResponse({'ok': False, 'error': 'Scraper stopped', 'data': []})
                # Like script: only page and questionPerPage (no chapter_id/subject in request)
                page_params = {'page': page, 'questionPerPage': per_page}
                resp = _scraper_safe_request(base_url, headers=request_headers, params=page_params, timeout=45)
                if not resp:
                    if page > 1 and page_retry < max_page_retry:
                        page_retry += 1
                        _scraper_progress('Page {} request failed, retrying ({}/{})...'.format(page, page_retry, max_page_retry))
                        time.sleep(2)
                        continue
                    break
                page_retry = 0
                try:
                    raw = resp.json()
                except Exception:
                    break
                data = raw.get('data', raw) if isinstance(raw, dict) else raw
                if total_questions is None and isinstance(raw, dict):
                    total_questions = raw.get('total') or raw.get('totalQuestions')
                if total_questions is None and isinstance(data, dict):
                    total_questions = data.get('total') or data.get('totalQuestions')
                # Parse questions like script: data.questions or raw.questions or data if list; fallbacks for other APIs
                questions = None
                if isinstance(data, dict):
                    questions = data.get('questions') or data.get('results') or data.get('items')
                if not questions and isinstance(raw, dict):
                    questions = raw.get('questions') or raw.get('results') or raw.get('items')
                if questions is None and isinstance(data, list):
                    questions = data
                if not questions:
                    questions = []
                if not questions:
                    break
                for q in questions:
                    if not isinstance(q, dict):
                        continue
                    qid = q.get('id')
                    if qid and qid not in seen_ids:
                        all_questions.append(q)
                        seen_ids.add(qid)
                _scraper_progress('Fetched page {}, {} questions (total so far {})'.format(page, len(questions), len(all_questions)))
                # Script: only break when no questions; otherwise continue to next page
                if len(questions) < per_page:
                    break
                page += 1
            if all_questions:
                break
            if session_id and session_id in _scraper_aborted_sessions:
                _scraper_progress('Scraper stopped by user.')
                return JsonResponse({'ok': False, 'error': 'Scraper stopped', 'data': []})
            if fetch_attempt < max_fetch_retries - 1:
                _scraper_progress('No questions returned, retrying ({}/{}) in {}s...'.format(
                    fetch_attempt + 2, max_fetch_retries, wait_between_retries_sec))
                for _ in range(wait_between_retries_sec):
                    if session_id and session_id in _scraper_aborted_sessions:
                        return JsonResponse({'ok': False, 'error': 'Scraper stopped', 'data': []})
                    time.sleep(1)
        if all_questions:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(all_questions, f, ensure_ascii=False, indent=2)
            _scraper_progress('Saved {} questions to {}'.format(len(all_questions), file_path))
            if level1_name or level2_label:
                csv_path = file_path.replace('.json', '.csv') if file_path.endswith('.json') else (file_path + '.csv')
                csv_header = ['ID', 'Subject', 'ChapterNo', 'Chapter', 'Topic', 'Question', 'Option 1', 'Option 2', 'Option 3', 'Option 4',
                              'Answer', 'Explanation', 'Question Type', 'Level', 'Subsources']
                with open(csv_path, 'w', newline='', encoding='utf-8-sig') as cf:
                    writer = csv.writer(cf)
                    writer.writerow(csv_header)
                    for idx, q in enumerate(all_questions, start=1):
                        if not isinstance(q, dict):
                            continue
                        options = q.get('option') or []
                        correct_index = q.get('mcq_solution_index')
                        correct_answer = options[correct_index] if isinstance(correct_index, int) and 0 <= correct_index < len(options) else ''
                        expl = _scraper_extract_text(q.get('answer_text')) or _scraper_extract_text(q.get('explanation_text'))
                        writer.writerow([
                            idx, level1_name, chapter_no, level2_label, _scraper_format_topics(q),
                            _scraper_extract_text(q.get('question_text')),
                            options[0] if len(options) > 0 else '', options[1] if len(options) > 1 else '',
                            options[2] if len(options) > 2 else '', options[3] if len(options) > 3 else '',
                            correct_answer, expl,
                            (q.get('question_type') or {}).get('name', ''),
                            (q.get('question_level') or {}).get('name', ''),
                            _scraper_format_subsources(q)
                        ])
                _scraper_progress('Saved chapter CSV: {}'.format(csv_path))
            _scraper_add_api_to_file(base_url.split('?')[0].strip())
            return JsonResponse({'ok': True, 'path': file_path, 'data': all_questions})
        empty_data = {'error': 'No data (after {} attempts)'.format(max_fetch_retries), 'data': [], 'chapter_id': chapter_id}
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(empty_data, f, ensure_ascii=False, indent=2)
        _scraper_progress('Saved (no data after {} retries) to {}'.format(max_fetch_retries, file_path))
        return JsonResponse({'ok': True, 'path': file_path, 'data': [], 'error': 'No questions returned (tried {} times)'.format(max_fetch_retries)})
    except Exception as e:
        logger.exception('scraper_fetch_and_save failed')
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump({'error': str(e), 'data': []}, f, ensure_ascii=False, indent=2)
            _scraper_progress('Saved (error) to {}'.format(file_path))
            return JsonResponse({'ok': True, 'path': file_path, 'data': [], 'error': str(e)})
        except Exception:
            return JsonResponse({'error': str(e), 'path': None, 'data': None}, status=500)


def _scraper_extract_text(editor_obj):
    """Extract plain text from editor JSON (blocks[].text)."""
    if not isinstance(editor_obj, dict):
        return ''
    blocks = editor_obj.get('blocks') or []
    return '\n'.join(block.get('text', '') for block in blocks if isinstance(block, dict))


def _scraper_format_topics(q):
    """Format topic names for CSV."""
    topic_names = []
    topics = q.get('topic')
    if isinstance(topics, list):
        for t in topics:
            if isinstance(t, dict) and t.get('name'):
                topic_names.append('"{}"'.format(t['name']))
    elif isinstance(topics, dict) and topics.get('name'):
        topic_names.append('"{}"'.format(topics['name']))
    return ', '.join(topic_names)


def _scraper_format_subsources(q):
    """Format question_subsources for CSV."""
    out = []
    for sub in (q.get('question_subsources') or []):
        if not isinstance(sub, dict):
            continue
        sub_source = sub.get('sub_source') or {}
        year_obj = sub.get('year') or {}
        short_name = sub_source.get('name', '')
        year_name = year_obj.get('name', '')
        if short_name and year_name:
            formatted = short_name + "'" + (year_name[-2:] if len(year_name) >= 2 else year_name)
            out.append('"{}"'.format(formatted))
    return ', '.join(out)


def _scraper_subject_file_has_questions(json_path):
    """Return True if json_path exists and contains at least one question (so we don't skip based on empty failed-run files)."""
    if not os.path.isfile(json_path):
        return False
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            content = json.load(f)
        if isinstance(content, list):
            return len(content) > 0
        if isinstance(content, dict):
            data = content.get('data', content)
            if isinstance(data, list):
                return len(data) > 0
            qs = content.get('questions')
            if not qs and isinstance(data, dict):
                qs = data.get('questions')
            return isinstance(qs, list) and len(qs) > 0
    except Exception:
        pass
    return False


@csrf_exempt
def scraper_file_exists(request):
    """GET ?group=...&filename=...&website=... (base name). Returns { exists: true } only if both .json and .csv exist AND json has at least one question (avoids skip from empty failed runs). Also returns path for UI."""
    if request.method != 'GET':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    group = (request.GET.get('group') or '').strip()
    filename = (request.GET.get('filename') or '').strip()
    website = (request.GET.get('website') or request.GET.get('site') or 'daricomma').strip()
    if not group or not filename:
        return JsonResponse({'error': 'group and filename are required'}, status=400)
    safe_name = re.sub(r'[^\w\-]', '_', _scraper_sanitize_filename(filename) or filename).strip('_') or 'data'
    dir_path = _scraper_site_dir(website, group)
    json_path = os.path.join(dir_path, safe_name + '.json')
    csv_path = os.path.join(dir_path, safe_name + '.csv')
    files_exist = os.path.isfile(json_path) and os.path.isfile(csv_path)
    has_questions = _scraper_subject_file_has_questions(json_path)
    exists = files_exist and has_questions
    return JsonResponse({'exists': exists, 'path': dir_path})


@csrf_exempt
def scraper_save_subject(request):
    """POST { group, level1_name, questions: [...] }. Writes level1_name.json and level1_name.csv (same format as Python script)."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        group = (body.get('group') or '').strip()
        website = (body.get('website') or body.get('site') or 'daricomma').strip()
        level1_name = (body.get('level1_name') or '').strip()
        questions = body.get('questions')
        if not isinstance(questions, list):
            questions = []
        if not group or not level1_name:
            return JsonResponse({'error': 'group and level1_name are required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    safe_name = re.sub(r'[^\w\-]', '_', _scraper_sanitize_filename(level1_name) or level1_name).strip('_') or 'subject'
    dir_path = _scraper_site_dir(website, group)
    json_path = os.path.join(dir_path, safe_name + '.json')
    csv_path = os.path.join(dir_path, safe_name + '.csv')
    try:
        # JSON: copy without internal fields (same as script)
        json_questions = []
        for q in questions:
            if isinstance(q, dict):
                qcopy = {k: v for k, v in q.items() if k not in ('_chapter', '_chapter_no', '_level1')}
                json_questions.append(qcopy)
            else:
                json_questions.append(q)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_questions, f, ensure_ascii=False, indent=2)
        csv_header = ['ID', 'Subject', 'ChapterNo', 'Chapter', 'Topic', 'Question', 'Option 1', 'Option 2', 'Option 3', 'Option 4',
                      'Answer', 'Explanation', 'Question Type', 'Level', 'Subsources']
        with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            writer.writerow(csv_header)
            for idx, q in enumerate(questions, start=1):
                if not isinstance(q, dict):
                    continue
                level1 = q.get('_level1', level1_name)
                chapter_no = q.get('_chapter_no', '')
                chapter = q.get('_chapter', '')
                options = q.get('option') or []
                correct_index = q.get('mcq_solution_index')
                correct_answer = options[correct_index] if isinstance(correct_index, int) and 0 <= correct_index < len(options) else ''
                expl = _scraper_extract_text(q.get('answer_text')) or _scraper_extract_text(q.get('explanation_text'))
                qtype = (q.get('question_type') or {}).get('name', '')
                qlevel = (q.get('question_level') or {}).get('name', '')
                writer.writerow([
                    idx, level1, chapter_no, chapter, _scraper_format_topics(q),
                    _scraper_extract_text(q.get('question_text')),
                    options[0] if len(options) > 0 else '', options[1] if len(options) > 1 else '',
                    options[2] if len(options) > 2 else '', options[3] if len(options) > 3 else '',
                    correct_answer, expl, qtype, qlevel, _scraper_format_subsources(q)
                ])
        _scraper_progress('Saved subject files: {}'.format(safe_name))
        return JsonResponse({'ok': True, 'path_json': json_path, 'path_csv': csv_path})
    except Exception as e:
        logger.exception('scraper_save_subject failed')
        return JsonResponse({'error': str(e)}, status=500)


# Persistent Selenium sessions for scraper (session_id -> driver). Headed browser for "Load website".
_scraper_sessions = {}
_scraper_aborted_sessions = set()  # session_ids for which user clicked Stop; in-flight fetch_and_save will exit
_SCRAPER_SESSION_TIMEOUT = 3600  # 1 hour


def _scraper_progress(msg):
    """Log and print scraper progress so it appears in the terminal (e.g. runserver)."""
    logger.info('[Scraper] %s', msg)
    try:
        print('[Scraper]', msg, flush=True)
    except Exception:
        pass


def _get_selenium_driver(headed=False, enable_performance_log=False):
    """Return Chrome WebDriver. headed=True shows the browser window. enable_performance_log=True for capturing API URLs from network."""
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        opts = Options()
        if not headed:
            opts.add_argument('--headless')
        opts.add_argument('--no-sandbox')
        opts.add_argument('--disable-dev-shm-usage')
        opts.add_argument('--disable-gpu')
        opts.add_argument('--disable-software-rasterizer')
        if enable_performance_log:
            opts.set_capability('goog:loggingPrefs', {'performance': 'ALL', 'browser': 'ALL'})
        try:
            from selenium.webdriver.chrome.service import Service
            from webdriver_manager.chrome import ChromeDriverManager
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=opts)
        except Exception:
            driver = webdriver.Chrome(options=opts)
        if enable_performance_log:
            try:
                driver.execute_cdp_cmd('Network.enable', {})
            except Exception:
                pass
        return driver
    except Exception as e:
        logger.warning('Selenium/Chrome not available: %s', e)
        return None


@csrf_exempt
def scraper_load_website(request):
    """POST { url, headless?: bool } - Open Chrome (headed or headless) and load URL. Returns session_id."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        url = (body.get('url') or '').strip()
        headless = bool(body.get('headless'))
        if not url:
            return JsonResponse({'error': 'url is required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    _scraper_progress('Opening browser (headless={})...'.format(headless))
    driver = _get_selenium_driver(headed=not headless, enable_performance_log=True)
    if not driver:
        return JsonResponse({'error': 'Selenium/Chrome not available. Install Chrome and: pip install selenium webdriver-manager'})
    try:
        import time
        import uuid
        _scraper_progress('Loading URL: {}'.format(url[:80] + ('...' if len(url) > 80 else '')))
        driver.get(url)
        wait_sec = 4 if headless else 2
        _scraper_progress('Waiting {}s for page load...'.format(wait_sec))
        time.sleep(wait_sec)
        session_id = str(uuid.uuid4())
        _scraper_sessions[session_id] = driver
        _scraper_progress('Session created: {}'.format(session_id[:8]))
        return JsonResponse({'session_id': session_id, 'url': url})
    except Exception as e:
        _scraper_progress('Load failed: {}'.format(e))
        try:
            driver.quit()
        except Exception:
            pass
        return JsonResponse({'error': str(e)})


@csrf_exempt
def scraper_navigate(request):
    """POST { session_id, url } - Navigate the session browser to url."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
        url = (body.get('url') or '').strip()
        if not session_id or not url:
            return JsonResponse({'error': 'session_id and url are required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    driver = _scraper_sessions.get(session_id)
    if not driver:
        return JsonResponse({'error': 'Session not found.'}, status=400)
    try:
        import time
        _scraper_progress('Navigating to: {}'.format(url[:70] + ('...' if len(url) > 70 else '')))
        driver.get(url)
        _scraper_progress('Waiting 4s for page load...')
        time.sleep(4)
        # If the site (e.g. Daricomma) redirects to a placeholder route, force the target URL again
        current = (driver.current_url or '')
        if '[subjectName]' in current or '[chapterId]' in current:
            _scraper_progress('Redirect detected; re-navigating to target URL...')
            driver.get(url)
            time.sleep(2)
        _scraper_progress('Navigate done.')
        return JsonResponse({'ok': True, 'url': url})
    except Exception as e:
        _scraper_progress('Navigate failed: {}'.format(e))
        return JsonResponse({'error': str(e)})


@csrf_exempt
def scraper_daricomma_login(request):
    """POST { session_id, login_url, username, password } - Navigate to login, fill form, submit."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
        login_url = (body.get('login_url') or '').strip()
        username = (body.get('username') or '').strip()
        password = (body.get('password') or '').strip()
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    if not session_id or not login_url:
        return JsonResponse({'error': 'session_id and login_url are required'}, status=400)
    driver = _scraper_sessions.get(session_id)
    if not driver:
        return JsonResponse({'error': 'Session not found.'}, status=400)
    try:
        import time
        from selenium.webdriver.common.by import By
        _scraper_progress('Loading login page...')
        driver.get(login_url)
        _scraper_progress('Waiting 4s for login form...')
        time.sleep(4)
        username_el = None
        pwd_el = None
        _scraper_progress('Looking for username/email field...')
        for sel in ['input[name="email"]', 'input[name="username"]', 'input[type="email"]', 'input#email', 'input#username']:
            try:
                username_el = driver.find_element(By.CSS_SELECTOR, sel)
                break
            except Exception:
                continue
        if not username_el:
            _scraper_progress('Trying Mantine form (first form text + password)...')
            # Daricomma sign-in uses Mantine: first field is mobile (type=text), second is password
            # Scope to first form (login tab) so we don't hit register form fields
            try:
                forms = driver.find_elements(By.CSS_SELECTOR, 'form')
                if forms:
                    form = forms[0]
                    username_el = form.find_element(By.CSS_SELECTOR, 'input[type="text"]')
                    pwd_el = form.find_element(By.CSS_SELECTOR, 'input[type="password"]')
            except Exception:
                pass
        if not username_el:
            return JsonResponse({'error': 'Could not find email/username field.'})
        if not pwd_el:
            try:
                pwd_el = driver.find_element(By.CSS_SELECTOR, 'input[name="password"], input[type="password"], input#password')
            except Exception:
                if username_el:
                    try:
                        form = username_el.find_element(By.XPATH, './ancestor::form[1]')
                        pwd_el = form.find_element(By.CSS_SELECTOR, 'input[type="password"]')
                    except Exception:
                        pass
        if not pwd_el:
            return JsonResponse({'error': 'Could not find password field.'})
        _scraper_progress('Filling username and password...')
        username_el.clear()
        username_el.send_keys(username)
        pwd_el.clear()
        pwd_el.send_keys(password)
        _scraper_progress('Submitting login form...')
        try:
            form = pwd_el.find_element(By.XPATH, './ancestor::form')
            form.submit()
        except Exception:
            for el in driver.find_elements(By.CSS_SELECTOR, 'button[type="submit"], input[type="submit"]'):
                try:
                    el.click()
                    break
                except Exception:
                    continue
        _scraper_progress('Waiting 3s after submit...')
        time.sleep(3)
        _scraper_progress('Login done.')
        return JsonResponse({'ok': True})
    except Exception as e:
        logger.exception('scraper_daricomma_login failed')
        return JsonResponse({'error': str(e)})


def _mantine_options(driver, combobox_index, previous_selections):
    """Get options from Mantine Select by index (0=first, 1=second). previous_selections for level 2."""
    import time
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    roots = driver.find_elements(By.CSS_SELECTOR, '.mantine-Select-root')
    if combobox_index >= len(roots):
        return []
    for i, val in enumerate(previous_selections):
        roots = driver.find_elements(By.CSS_SELECTOR, '.mantine-Select-root')
        if i >= len(roots):
            break
        try:
            cb = roots[i].find_element(By.CSS_SELECTOR, '[role="combobox"]')
            cb.click()
            time.sleep(0.9)
            listboxes = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"]')
            options_to_use = []
            for lb in listboxes:
                try:
                    if lb.is_displayed():
                        options_to_use = lb.find_elements(By.CSS_SELECTOR, '[role="option"]')
                        break
                except Exception:
                    pass
            if not options_to_use:
                options_to_use = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"] [role="option"]')
            for opt in options_to_use:
                ov = (opt.get_attribute('data-value') or opt.get_attribute('value') or '').strip()
                ot = (opt.text or '').strip()
                if ov == str(val).strip() or ot == str(val).strip():
                    opt.click()
                    break
            time.sleep(0.8)
        except Exception:
            try:
                driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.ESCAPE)
            except Exception:
                pass
        time.sleep(0.6)
    # After selecting Level1 (for Level2 dropdown), wait longer for API to populate options (script: wait_after_select=4)
    wait_after_selections = 4.0 if (combobox_index == 1 and previous_selections) else 1.5
    time.sleep(wait_after_selections)
    roots = driver.find_elements(By.CSS_SELECTOR, '.mantine-Select-root')
    if combobox_index >= len(roots):
        return []
    try:
        root = roots[combobox_index]
        combobox_el = root.find_element(By.CSS_SELECTOR, '[role="combobox"]')
        combobox_el.click()
        time.sleep(1.2)
        listbox_els = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"]')
        active_listbox = None
        for lb in listbox_els:
            try:
                if lb.is_displayed():
                    active_listbox = lb
                    break
            except Exception:
                pass
        if active_listbox is None:
            active_listbox = driver
        option_selector = '[role="listbox"] [role="option"]' if active_listbox == driver else '[role="option"]'
        try:
            option_els = active_listbox.find_elements(By.CSS_SELECTOR, option_selector)
        except Exception:
            option_els = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"] [role="option"]')
        opts = []
        # Script get_all_options: open dropdown, collect option text/value, close – no selecting each option
        for opt in (option_els or []):
            try:
                driver.execute_script('arguments[0].scrollIntoView({block: "nearest"});', opt)
                time.sleep(0.03)
            except Exception:
                pass
            text = (opt.text or '').strip()
            v = opt.get_attribute('data-value') or opt.get_attribute('value') or text
            opts.append({'value': v or text, 'text': text or v})
        driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.ESCAPE)
        return opts
    except Exception:
        try:
            driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.ESCAPE)
        except Exception:
            pass
        return []


@csrf_exempt
def scraper_capture_mantine(request):
    """POST { session_id, combobox_index (0|1), previous_selections: [] } - Get Mantine Select options."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
        combobox_index = int(body.get('combobox_index') or 0)
        previous_selections = body.get('previous_selections') or []
        if not session_id:
            return JsonResponse({'error': 'session_id is required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    driver = _scraper_sessions.get(session_id)
    if not driver:
        return JsonResponse({'error': 'Session not found.', 'options': []})
    try:
        _scraper_progress('Capturing Mantine options (combobox_index={})...'.format(combobox_index))
        opts = _mantine_options(driver, combobox_index, previous_selections)
        _scraper_progress('Got {} options.'.format(len(opts)))
        return JsonResponse({'options': opts})
    except Exception as e:
        _scraper_progress('Capture failed: {}'.format(e))
        logger.exception('scraper_capture_mantine failed')
        return JsonResponse({'error': str(e), 'options': []})


def _scraper_wait_for_chapter_url(driver, timeout=12, poll_interval=0.5):
    """After Level2 selection, wait until URL contains /academic/ and /chapter/ with non-empty chapter part (script: timeout=12)."""
    start = time.time()
    while (time.time() - start) < timeout:
        url = (driver.current_url or '').strip()
        if '/academic/' in url and '/chapter/' in url:
            parts = url.split('/chapter/', 1)
            if len(parts) == 2:
                rest = (parts[1].split('&')[0].split('?')[0] or '').strip()
                if rest:
                    return True
        time.sleep(poll_interval)
    return False


def _scraper_capture_question_api_url(driver, prefix, wait_after=1.5):
    """Capture the first request URL starting with prefix from performance log (script: wait_after=1.5)."""
    time.sleep(wait_after)
    try:
        logs = driver.get_log('performance')
    except Exception:
        return None
    for entry in logs:
        try:
            msg = json.loads(entry.get('message', '{}'))
            event = msg.get('message') or msg
            method = event.get('method')
            if method not in ('Network.requestWillBeSent', 'Network.responseReceived'):
                continue
            params = event.get('params') or {}
            req_url = ''
            if method == 'Network.requestWillBeSent':
                req_url = (params.get('request') or {}).get('url') or ''
            else:
                req_url = (params.get('response') or {}).get('url') or ''
            if req_url.startswith(prefix):
                return req_url.split('?')[0].strip()
        except Exception:
            continue
    return None


@csrf_exempt
def scraper_capture_question_url(request):
    """POST { session_id, level1_value, level2_value, level1_label?, level2_label?, api_base_prefix? } - Select Level1 & Level2 by visible text (script: uses option text). Wait for chapter URL, capture question API URL. Returns { url }."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
        level1_value = (body.get('level1_value') or body.get('level1') or '').strip()
        level2_value = (body.get('level2_value') or body.get('level2') or '').strip()
        level1_label = (body.get('level1_label') or '').strip()
        level2_label = (body.get('level2_label') or '').strip()
        api_base_prefix = (body.get('api_base_prefix') or 'https://api.daricomma.com/v2/question/').strip()
        if not session_id:
            return JsonResponse({'error': 'session_id is required'}, status=400)
        # Script selects by visible text (option text); prefer label when provided
        level1_selection = level1_label or level1_value
        level2_selection = level2_label or level2_value
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    driver = _scraper_sessions.get(session_id)
    if not driver:
        return JsonResponse({'error': 'Session not found.', 'url': ''})
    try:
        from selenium.webdriver.common.by import By
        from selenium.webdriver.common.keys import Keys
        _scraper_progress('Capturing question API URL for chapter...')
        # Apply Level1 then Level2 selection by visible text (script: select_mantine_option uses option_text)
        for i, val in enumerate([level1_selection, level2_selection]):
            if not val:
                continue
            roots = driver.find_elements(By.CSS_SELECTOR, '.mantine-Select-root')
            if i >= len(roots):
                break
            try:
                cb = roots[i].find_element(By.CSS_SELECTOR, '[role="combobox"]')
                cb.click()
                time.sleep(0.9)
                listboxes = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"]')
                options_to_use = []
                for lb in listboxes:
                    try:
                        if lb.is_displayed():
                            options_to_use = lb.find_elements(By.CSS_SELECTOR, '[role="option"]')
                            break
                    except Exception:
                        pass
                if not options_to_use:
                    options_to_use = driver.find_elements(By.CSS_SELECTOR, '[role="listbox"] [role="option"]')
                for opt in options_to_use:
                    ov = (opt.get_attribute('data-value') or opt.get_attribute('value') or '').strip()
                    ot = (opt.text or '').strip()
                    if ov == str(val).strip() or ot == str(val).strip():
                        opt.click()
                        break
                time.sleep(0.8)
            except Exception:
                try:
                    driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.ESCAPE)
                except Exception:
                    pass
            time.sleep(0.6)
        time.sleep(1.0)
        _scraper_wait_for_chapter_url(driver, timeout=12)
        captured = _scraper_capture_question_api_url(driver, prefix=api_base_prefix, wait_after=1.5)
        if captured:
            _scraper_progress('Captured API URL: {}...'.format(captured[:60]))
            return JsonResponse({'url': captured})
        return JsonResponse({'url': '', 'error': 'No question API URL found in network log'})
    except Exception as e:
        _scraper_progress('Capture question URL failed: {}'.format(e))
        logger.exception('scraper_capture_question_url failed')
        return JsonResponse({'error': str(e), 'url': ''})


@csrf_exempt
def scraper_capture_dropdown(request):
    """POST { session_id, level, previous_selections: [] } - Get options for level (1-based). Applies previous selections first."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
        level = int(body.get('level') or 1)
        previous_selections = body.get('previous_selections') or []
        if not session_id:
            return JsonResponse({'error': 'session_id is required'}, status=400)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    driver = _scraper_sessions.get(session_id)
    if not driver:
        return JsonResponse({'error': 'Session not found. Load website again.', 'options': []})
    try:
        import time
        from selenium.webdriver.support.ui import Select
        from selenium.webdriver.common.by import By
        selects = driver.find_elements(By.CSS_SELECTOR, 'select')
        if level < 1 or level > len(selects):
            return JsonResponse({'options': [], 'message': 'No select at this level.'})
        # Apply previous selections so level 2+ options are correct
        for i, val in enumerate(previous_selections):
            if i >= len(selects):
                break
            try:
                Select(selects[i]).select_by_value(val)
                time.sleep(0.5)
            except Exception:
                pass
        time.sleep(0.5)
        selects = driver.find_elements(By.CSS_SELECTOR, 'select')
        if level > len(selects):
            return JsonResponse({'options': []})
        sel = selects[level - 1]
        opts = []
        for opt in sel.find_elements(By.TAG_NAME, 'option'):
            v = opt.get_attribute('value')
            if v is None:
                v = ''
            text = (opt.text or '').strip() or v
            opts.append({'value': v, 'text': text})
        return JsonResponse({'options': opts})
    except Exception as e:
        logger.exception('scraper_capture_dropdown failed')
        return JsonResponse({'error': str(e), 'options': []})


@csrf_exempt
def scraper_close_session(request):
    """POST { session_id } - Close browser and remove session."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        session_id = (body.get('session_id') or '').strip()
    except Exception:
        session_id = ''
    if session_id:
        _scraper_aborted_sessions.add(session_id)
        driver = _scraper_sessions.pop(session_id, None)
        if driver:
            try:
                driver.quit()
            except Exception:
                pass
    return JsonResponse({'ok': True})


@csrf_exempt
def scraper_discover_dropdowns(request):
    """GET ?url=... - Open page with Selenium, find all <select>, return groups (one per dropdown)."""
    if request.method != 'GET':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    url = (request.GET.get('url') or '').strip()
    if not url:
        return JsonResponse({'error': 'url is required'}, status=400)
    driver = _get_selenium_driver()
    if not driver:
        return JsonResponse({
            'error': 'Selenium/Chrome not available. Install: 1) Chrome browser, 2) pip install selenium webdriver-manager. Or use "Capture by your activity" above instead.',
            'groups': []
        })
    try:
        import time
        driver.get(url)
        time.sleep(2)
        try:
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            from selenium.webdriver.common.by import By
            WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
            time.sleep(2)
        except Exception:
            pass
        selects = driver.find_elements('css selector', 'select')
        groups = []
        for i, sel in enumerate(selects):
            name = sel.get_attribute('name') or sel.get_attribute('id') or ('Dropdown_%d' % (i + 1))
            opts = []
            for opt in sel.find_elements('tag name', 'option'):
                v = opt.get_attribute('value')
                if v is None:
                    v = ''
                text = (opt.text or '').strip() or v
                opts.append({'value': v, 'text': text})
            groups.append({'name': name, 'options': opts})
        if not groups:
            return JsonResponse({
                'groups': [],
                'message': 'No native <select> dropdowns found on this page. Many sites use custom dropdowns (divs). Use "Capture by your activity" to add options manually.'
            })
        return JsonResponse({'groups': groups})
    except Exception as e:
        logger.exception('scraper_discover_dropdowns failed')
        return JsonResponse({'error': str(e), 'groups': []})
    finally:
        try:
            driver.quit()
        except Exception:
            pass


@csrf_exempt
def scraper_dynamic_dropdown(request):
    """POST { url, level_index, selections: [] } - Open page, set prior selects, return options for level level_index."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        body = json.loads(request.body.decode('utf-8'))
        url = (body.get('url') or '').strip()
        level_index = int(body.get('level_index') or 0)
        selections = body.get('selections') or []
        if not url:
            return JsonResponse({'error': 'url is required'}, status=400)
    except (json.JSONDecodeError, TypeError) as e:
        return JsonResponse({'error': str(e)}, status=400)
    driver = _get_selenium_driver()
    if not driver:
        return JsonResponse({'error': 'Selenium/Chrome not available.', 'options': []})
    try:
        driver.get(url)
        import time
        time.sleep(2)
        selects = driver.find_elements('css selector', 'select')
        if level_index <= 0 or level_index > len(selects):
            return JsonResponse({'options': []})
        try:
            from selenium.webdriver.support.ui import Select
            for i, val in enumerate(selections):
                if i >= len(selects):
                    break
                Select(selects[i]).select_by_value(val)
                time.sleep(0.6)
        except Exception:
            pass
        time.sleep(0.8)
        selects = driver.find_elements('css selector', 'select')
        sel = selects[level_index - 1] if level_index <= len(selects) else None
        if not sel:
            return JsonResponse({'options': []})
        opts = []
        for opt in sel.find_elements('tag name', 'option'):
            v = opt.get_attribute('value')
            if v is None:
                v = ''
            text = (opt.text or '').strip() or v
            opts.append({'value': v, 'text': text})
        return JsonResponse({'options': opts})
    except Exception as e:
        logger.exception('scraper_dynamic_dropdown failed')
        return JsonResponse({'error': str(e), 'options': []})
    finally:
        try:
            driver.quit()
        except Exception:
            pass


class PasswordUpdateView(APIView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        newpassword = request.data.get('newpassword')
        
        try:
            user = Customer.objects.get(username=username)
            
            # Verify old password
            if user.password.startswith('pbkdf2_') or user.password.startswith('argon2'):
                password_valid = user.check_password(password)
            else:
                # Legacy: compare plain text
                password_valid = (user.password == password)
            
            if password_valid:
                # Set new hashed password
                user.set_password(newpassword)
                user.save()
                token = self.generate_unique_key()
                CustomerToken.objects.filter(customer=user).delete()
                CustomerToken.objects.create(key=token, customer=user)
                return Response({'authToken': token, 'message': 'Password updated successfully'}, status=status.HTTP_200_OK)
            else:
                return Response({'error': 'Invalid current password'}, status=status.HTTP_401_UNAUTHORIZED)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        

    def generate_unique_key(self):
        length = 40
        characters = string.ascii_letters + string.digits
        key = ''.join(random.choice(characters) for _ in range(length))
        return key 



class MobileUpdateView(APIView):
    def post(self, request, *args, **kwargs):
        username = request.data.get('username')
        newusername = request.data.get('newusername')
        password = request.data.get('password')
        
        try:
            user = Customer.objects.get(username=username)
            
            # Verify password
            if user.password.startswith('pbkdf2_') or user.password.startswith('argon2'):
                password_valid = user.check_password(password)
            else:
                # Legacy: compare plain text
                password_valid = (user.password == password)
            
            if password_valid:
                # Check if new username already exists
                if Customer.objects.filter(username=newusername).exclude(pk=user.pk).exists():
                    return Response({'error': 'Mobile number already exists'}, status=status.HTTP_400_BAD_REQUEST)
                
                user.username = newusername
                user.save()
                token = self.generate_unique_key()
                CustomerToken.objects.filter(customer=user).delete()
                CustomerToken.objects.create(key=token, customer=user)
                return Response({'authToken': token, 'message': 'Mobile number updated successfully'}, status=status.HTTP_200_OK)
            else:
                return Response({'error': 'Invalid password'}, status=status.HTTP_401_UNAUTHORIZED)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        

    def generate_unique_key(self):
        length = 40
        characters = string.ascii_letters + string.digits
        key = ''.join(random.choice(characters) for _ in range(length))
        return key 


# ==============================================================================
# VERIFICATION VIEWS (Email + WhatsApp)
# ==============================================================================

class SendVerificationCodeView(APIView):
    """Send verification code via Telegram/Email/WhatsApp"""
    permission_classes = [PublicAccess]

    def post(self, request):
        username = request.data.get('username')
        if not username:
            return Response({'error': 'Phone number is required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        from .verification_service import send_verification_code
        result = send_verification_code(customer, purpose='verification')
        if result['success']:
            response_data = {
                'message': result.get('message', 'Verification code sent'),
                'method': result.get('method', 'unknown'),
                'expires_in': 600,
            }
            if result.get('requires_linking'):
                response_data['requires_linking'] = True
                response_data['bot_username'] = result.get('bot_username')
            return Response(response_data, status=status.HTTP_200_OK)
        return Response({
            'error': result['message'],
            'method': result.get('method'),
            'requires_linking': result.get('requires_linking', False),
            'bot_username': result.get('bot_username'),
        }, status=status.HTTP_400_BAD_REQUEST)


class VerifyCodeView(APIView):
    """Verify the verification code"""
    permission_classes = [PublicAccess]

    def post(self, request):
        username = request.data.get('username')
        code = request.data.get('code')
        if not username or not code:
            return Response({'error': 'Phone number and code are required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        from .verification_service import verify_code
        result = verify_code(customer, code)
        if result['success']:
            return Response({'message': result['message'], 'verified': True}, status=status.HTTP_200_OK)
        return Response({'error': result['message'], 'verified': False}, status=status.HTTP_400_BAD_REQUEST)


class SendPasswordResetCodeView(APIView):
    """Send password reset code via Email (primary) or WhatsApp (fallback)"""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def post(self, request):
        username = request.data.get('username')
        email = request.data.get('email')
        if not username:
            return Response({'success': False, 'error': 'Phone number is required', 'message': 'Phone number is required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'success': False, 'error': 'User not found', 'message': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        if email:
            from .verification_service import send_verification_to_email
            result = send_verification_to_email(customer, email, purpose='password_reset')
        else:
            from .verification_service import send_verification_code
            result = send_verification_code(customer, purpose='password_reset')
        if result['success']:
            return Response({
                'success': True,
                'message': result.get('message', 'Password reset code sent'),
                'method': result.get('method'),
                'expires_in': 600,
            }, status=status.HTTP_200_OK)
        return Response({
            'success': False,
            'message': result.get('message', 'Failed to send code'),
            'error': result.get('message'),
            'method': result.get('method'),
            'needs_email': result.get('needs_email', False),
            'needs_activation': result.get('needs_activation', False),
        }, status=status.HTTP_400_BAD_REQUEST)


class ResetPasswordWithCodeView(APIView):
    """Reset password using verification code"""
    permission_classes = [PublicAccess]

    def post(self, request):
        username = request.data.get('username')
        code = request.data.get('code')
        new_password = request.data.get('new_password')
        save_email = request.data.get('save_email')
        if not username or not code or not new_password:
            return Response({'error': 'Phone number, code, and new password are required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        from .verification_service import verify_code
        result = verify_code(customer, code)
        if not result['success']:
            return Response({'error': result['message']}, status=status.HTTP_400_BAD_REQUEST)
        customer.set_password(new_password)
        if save_email:
            customer.email = save_email
        customer.save()
        return Response({'message': 'Password reset successfully'}, status=status.HTTP_200_OK)


class UpdateEmailView(APIView):
    """Allow user to add/update their email"""
    permission_classes = [PublicAccess]

    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        email = request.data.get('email')
        if not username or not password or not email:
            return Response({'error': 'Phone number, password, and email are required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        if not customer.check_password(password):
            return Response({'error': 'Invalid password'}, status=status.HTTP_401_UNAUTHORIZED)
        if Customer.objects.filter(email=email).exclude(pk=customer.pk).exists():
            return Response({'error': 'This email is already registered'}, status=status.HTTP_400_BAD_REQUEST)
        customer.email = email
        customer.save(update_fields=['email'])
        return Response({'message': 'Email updated successfully'}, status=status.HTTP_200_OK)


class UpdateWhatsAppApiKeyView(APIView):
    """Allow user to save their CallMeBot API key for free WhatsApp notifications"""
    permission_classes = [PublicAccess]

    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        whatsapp_apikey = request.data.get('whatsapp_apikey')
        if not username or not password or not whatsapp_apikey:
            return Response({'error': 'Phone number, password, and WhatsApp API key are required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            customer = Customer.objects.get(username=username)
        except Customer.DoesNotExist:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        if not customer.check_password(password):
            return Response({'error': 'Invalid password'}, status=status.HTTP_401_UNAUTHORIZED)
        if hasattr(customer, 'whatsapp_apikey'):
            customer.whatsapp_apikey = whatsapp_apikey
            customer.save(update_fields=['whatsapp_apikey'])
        return Response({'message': 'WhatsApp API key saved successfully'}, status=status.HTTP_200_OK)


class GenerateDefaultPasswordView(APIView):
    """Generate default password preview (for frontend display)"""
    permission_classes = [PublicAccess]

    def post(self, request):
        full_name = request.data.get('fullName', '')
        year_of_birth = request.data.get('year_of_birth')
        if not full_name or not year_of_birth:
            return Response({'error': 'Full name and year of birth are required'}, status=status.HTTP_400_BAD_REQUEST)
        name_part = full_name[:3].strip()
        if len(name_part) > 0:
            name_part = name_part[0].upper() + name_part[1:].lower()
        default_password = f"{name_part}@{year_of_birth}"
        return Response({'default_password': default_password}, status=status.HTTP_200_OK)


# ----- Removed: view_order, view_ordered, view_canceled, update_shipped_status, move_*_orders, get_shipped_status -----
# ----- Removed: GroupViewSet, SubjectViewSet, ChapterViewSet, TopicViewSet, InstituteViewSet, YearViewSet, McqIctViewSet,
#      SubjectQuestionTablesView, SubjectQuestionDataView, GetGroupsByClassView, GetDepartmentsView, GetClassInfoView -----


def _university_departments_json_path():
    """Path to departments.json (university departments list for Teacher signup)."""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'departments.json')


def _append_department_to_json(dept_name, faculty=None):
    """Append a department to departments.json (e.g. from signup 'Others')."""
    if not dept_name or len(dept_name) > 200:
        return
    path = _university_departments_json_path()
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {'departments': []}
    departments = data.get('departments', [])
    if any((d.get('dept_name') or '').strip().lower() == dept_name.lower() for d in departments):
        return
    departments.append({
        'dept_code': 'CUSTOM_' + str(len(departments) + 1),
        'dept_name': dept_name,
        'faculty': faculty
    })
    data['departments'] = departments
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except OSError as e:
        logger.warning('Could not write departments.json: %s', e)


class UniversityDepartmentsView(APIView):
    """
    GET: Return university departments from departments.json (all aspects of knowledge, worldwide).
    Used for Teacher signup when Level = University. Not from database.
    POST: Append a new department (when user selects "Others" and enters name); adds to departments.json.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        path = _university_departments_json_path()
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            departments = data.get('departments', [])
            # Sort by dept_name ascending (Others is added by frontend at the end)
            departments = sorted(departments, key=lambda d: (d.get('dept_name') or '').strip().lower())
        except (FileNotFoundError, json.JSONDecodeError) as e:
            logger.warning('University departments JSON not found or invalid: %s', e)
            departments = []
        return Response({'departments': departments, 'count': len(departments)}, status=status.HTTP_200_OK)

    def post(self, request):
        """Append a department to departments.json (e.g. when user chose 'Others' and entered a name)."""
        dept_name = (request.data.get('dept_name') or request.data.get('department_name') or '').strip()
        if not dept_name or len(dept_name) > 200:
            return Response(
                {'error': 'dept_name is required and must be 1–200 characters'},
                status=status.HTTP_400_BAD_REQUEST
            )
        path = _university_departments_json_path()
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            data = {'departments': []}
        departments = data.get('departments', [])
        faculty = (request.data.get('faculty') or '').strip() or None
        dept_code = (request.data.get('dept_code') or '').strip()
        if not dept_code:
            dept_code = 'CUSTOM_' + str(len(departments) + 1)
        if any(d.get('dept_name', '').strip().lower() == dept_name.lower() for d in departments):
            return Response({'departments': data.get('departments', []), 'count': len(departments)}, status=status.HTTP_200_OK)
        departments.append({
            'dept_code': dept_code[:20],
            'dept_name': dept_name,
            'faculty': faculty
        })
        data['departments'] = departments
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except OSError as e:
            logger.warning('Could not write departments.json: %s', e)
            return Response({'error': 'Could not save department'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'departments': departments, 'count': len(departments)}, status=status.HTTP_201_CREATED)


# ----- Signup: levels, subjects, groups by country/level from cheradip_hsc.cheradip_subject -----

def _level_code_to_sql(level):
    """Map frontend level code (PSC, JSC, SSC, HSC) or level_tr to SQL WHERE fragment and params for cheradip_subject."""
    level = (level or '').strip()
    if not level:
        return None, []
    level_upper = level.upper()
    # Map standard codes to class_level conditions (country_id filtered separately)
    if level_upper == 'PSC':
        return " ( class_level IN (%s,%s,%s,%s,%s) OR ( level_tr LIKE %s AND level_tr NOT LIKE %s ) ) ", ['1', '2', '3', '4', '5', '%Primary%', '%Pre-primary%']
    if level_upper == 'JSC':
        return " ( class_level IN (%s,%s,%s) OR level_tr LIKE %s ) ", ['6', '7', '8', '%Junior%']
    if level_upper == 'SSC':
        return " ( class_level = %s OR level_tr LIKE %s OR level_tr LIKE %s ) ", ['9-10', '%SSC%', '%Secondary%']
    if level_upper == 'HSC':
        return " ( class_level = %s OR level_tr LIKE %s OR level_tr LIKE %s ) ", ['11-12', '%HSC%', '%Higher%']
    # Otherwise treat as level_tr (e.g. Pre-primary)
    return " level_tr = %s ", [level]


class LevelsByCountryView(APIView):
    """GET ?country_code=BD – distinct levels from cheradip_subject for signup Level dropdown. Returns level, level_tr, label."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        if not country_code:
            return Response({'levels': [], 'country_code': country_code}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'levels': [], 'country_code': country_code, 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        levels = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT level_tr, MIN(CASE "
                    "  WHEN class_level IS NOT NULL AND TRIM(COALESCE(class_level, '')) != '' "
                    "  THEN CAST(SUBSTRING_INDEX(CONCAT(TRIM(class_level), '-'), '-', 1) AS UNSIGNED) "
                    "  ELSE NULL END) AS sort_key FROM cheradip_subject "
                    "WHERE (country_id = %s OR country_id IS NULL) AND level_tr IS NOT NULL AND TRIM(COALESCE(level_tr, '')) != '' "
                    "GROUP BY level_tr ORDER BY sort_key, level_tr",
                    [country_code]
                )
                for row in cur.fetchall():
                    level_tr = (row[0] or '').strip()
                    if level_tr:
                        # Map level_tr to code for dropdown value when possible
                        code = level_tr
                        if level_tr.lower().startswith('pre-primary') or level_tr == 'Pre-primary':
                            code = 'Pre-primary'
                        elif 'Primary' in level_tr and 'Pre-primary' not in level_tr:
                            code = 'PSC'
                        elif 'Junior' in level_tr:
                            code = 'JSC'
                        elif 'Higher' in level_tr or 'HSC' in level_tr or 'Alim' in level_tr:
                            code = 'HSC'
                        elif 'Secondary' in level_tr or 'SSC' in level_tr or 'Dakhil' in level_tr:
                            code = 'SSC'
                        levels.append({'level': code, 'level_tr': level_tr, 'label': level_tr})
        except Exception as e:
            logger.exception('LevelsByCountryView: %s', e)
            return Response({'levels': [], 'country_code': country_code, 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'levels': levels, 'country_code': country_code}, status=status.HTTP_200_OK)


class SubjectsByCountryLevelView(APIView):
    """GET ?country_code=BD&level=PSC – subjects from cheradip_subject for Teacher signup."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        level = (request.query_params.get('level') or '').strip()
        if not country_code or not level:
            return Response({'subjects': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'subjects': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        level_sql, level_params = _level_code_to_sql(level)
        if level_sql is None:
            return Response({'subjects': []}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        subjects = []
        try:
            with conn.cursor() as cur:
                sql = (
                    "SELECT DISTINCT subject_tr, subject_code, subject_name FROM cheradip_subject "
                    "WHERE (country_id = %s OR country_id IS NULL) AND " + level_sql +
                    " AND subject_tr IS NOT NULL AND TRIM(COALESCE(subject_tr, '')) != '' ORDER BY subject_tr"
                )
                cur.execute(sql, [country_code] + level_params)
                for row in cur.fetchall():
                    subject_tr = (row[0] or '').strip()
                    if subject_tr:
                        subjects.append({
                            'subject_tr': subject_tr,
                            'subject_code': row[1] if len(row) > 1 else None,
                            'subject_name': (row[2] or '').strip() if len(row) > 2 else subject_tr,
                            'id': subject_tr,
                            'name': subject_tr,
                        })
        except Exception as e:
            logger.exception('SubjectsByCountryLevelView: %s', e)
            return Response({'subjects': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'subjects': subjects}, status=status.HTTP_200_OK)


class GroupsByCountryLevelView(APIView):
    """GET ?country_code=BD&level=PSC – distinct groups from cheradip_subject.groups for Student signup."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        level = (request.query_params.get('level') or '').strip()
        if not country_code or not level:
            return Response({'groups': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'groups': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        level_sql, level_params = _level_code_to_sql(level)
        if level_sql is None:
            return Response({'groups': []}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        seen = set()
        try:
            with conn.cursor() as cur:
                sql = (
                    "SELECT groups FROM cheradip_subject "
                    "WHERE (country_id = %s OR country_id IS NULL) AND " + level_sql +
                    " AND groups IS NOT NULL AND TRIM(COALESCE(groups, '')) != ''"
                )
                cur.execute(sql, [country_code] + level_params)
                for row in cur.fetchall():
                    for p in _parse_groups_column(row[0]):
                        if p:
                            seen.add(p)
        except Exception as e:
            logger.exception('GroupsByCountryLevelView: %s', e)
            return Response({'groups': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'groups': sorted(seen)}, status=status.HTTP_200_OK)


class GroupsByClassView(APIView):
    """GET ?class_code=9-10 – distinct groups from cheradip_subject for that class_level (Student signup Group dropdown)."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        class_code = (request.query_params.get('class_code') or '').strip()
        if not class_code:
            return Response({'groups': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'groups': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        seen = set()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT groups FROM cheradip_subject WHERE class_level = %s "
                    "AND groups IS NOT NULL AND TRIM(COALESCE(groups, '')) != ''",
                    [class_code]
                )
                for row in cur.fetchall():
                    for p in _parse_groups_column(row[0]):
                        if p:
                            seen.add(p)
        except Exception as e:
            logger.exception('GroupsByClassView: %s', e)
            return Response({'groups': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'groups': sorted(seen)}, status=status.HTTP_200_OK)


class SubjectsForDegreeView(APIView):
    """GET ?country_code=BD – subjects for University (Degree/Honours/Masters). Uses class_level 13-16 or honours DB if present."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        if not country_code:
            return Response({'subjects': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'subjects': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        subjects = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT DISTINCT subject_tr, subject_code, subject_name FROM cheradip_subject "
                    "WHERE (country_id = %s OR country_id IS NULL) AND subject_tr IS NOT NULL AND TRIM(COALESCE(subject_tr, '')) != '' "
                    "AND (class_level IN ('13','14','15','16','13-16') OR level_tr LIKE %s OR level_tr LIKE %s) ORDER BY subject_tr",
                    [country_code, '%Degree%', '%Honours%']
                )
                for row in cur.fetchall():
                    subject_tr = (row[0] or '').strip()
                    if subject_tr:
                        subjects.append({
                            'subject_tr': subject_tr,
                            'subject_code': row[1] if len(row) > 1 else None,
                            'subject_name': (row[2] or '').strip() if len(row) > 2 else subject_tr,
                            'id': row[1] if len(row) > 1 else subject_tr,
                            'name': subject_tr,
                        })
        except Exception as e:
            logger.exception('SubjectsForDegreeView: %s', e)
            return Response({'subjects': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'subjects': subjects}, status=status.HTTP_200_OK)


class ClassesByCountryView(APIView):
    """GET ?country_code=BD – distinct class_level from cheradip_subject for signup Class dropdown. Optional ?useHsc=1 same behaviour."""
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        country_code = (request.query_params.get('country_code') or '').strip().upper()
        if not country_code:
            return Response({'classes': [], 'country_code': country_code}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'classes': [], 'country_code': country_code, 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        classes = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT class_level FROM cheradip_subject "
                    "WHERE (country_id = %s OR country_id IS NULL) AND class_level IS NOT NULL AND TRIM(COALESCE(class_level, '')) != '' "
                    "GROUP BY class_level ORDER BY CAST(SUBSTRING_INDEX(CONCAT(TRIM(class_level), '-'), '-', 1) AS UNSIGNED), class_level",
                    [country_code]
                )
                for row in cur.fetchall():
                    cl = (row[0] or '').strip()
                    if cl:
                        classes.append({
                            'value': cl,
                            'label': cl,
                            'has_groups': cl in ('9-10', '11-12'),
                        })
        except Exception as e:
            logger.exception('ClassesByCountryView: %s', e)
            return Response({'classes': [], 'country_code': country_code, 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'classes': classes, 'country_code': country_code}, status=status.HTTP_200_OK)


# ----- Question section: levels, subjects, chapters from cheradip_hsc.cheradip_subject and subject question tables -----

class QuestionLevelsView(APIView):
    """
    GET: Distinct levels (level_tr) from cheradip_subject in cheradip_hsc database.
    Used for the question page first dropdown (like signUp level selection).
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        if 'hsc' not in connections:
            return Response({'levels': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        levels = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT level_tr, MIN(CASE "
                    "  WHEN class_level IS NOT NULL AND TRIM(COALESCE(class_level, '')) != '' "
                    "  THEN CAST(SUBSTRING_INDEX(CONCAT(TRIM(class_level), '-'), '-', 1) AS UNSIGNED) "
                    "  ELSE NULL END) AS sort_key FROM cheradip_subject "
                    "WHERE level_tr IS NOT NULL AND TRIM(COALESCE(level_tr, '')) != '' "
                    "GROUP BY level_tr"
                )
                rows = cur.fetchall()
                for row in rows:
                    level_tr = (row[0] or '').strip()
                    sort_key = row[1] if len(row) > 1 and row[1] is not None else 0
                    if level_tr:
                        levels.append({
                            'level': level_tr,
                            'level_tr': level_tr,
                            'label': level_tr,
                            'sort_order': sort_key,
                        })
                # Descending by class_level (highest first); then A-Z for same key
                levels.sort(key=lambda x: (-(x['sort_order'] or 0), x['level_tr']))
        except Exception as e:
            logger.exception('QuestionLevelsView: %s', e)
            return Response({'levels': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'levels': levels}, status=status.HTTP_200_OK)


class QuestionClassesView(APIView):
    """
    GET: Distinct class_level for a level from cheradip_subject in cheradip_hsc.
    Query param: level_tr. Ordered by class_level ascending (0, 1, 5, 8, 9-10, 11-12).
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        if not level_tr:
            return Response({'classes': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'classes': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        classes = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT class_level FROM cheradip_subject "
                    "WHERE level_tr = %s AND class_level IS NOT NULL AND TRIM(COALESCE(class_level, '')) != '' "
                    "GROUP BY class_level "
                    "ORDER BY MIN(CAST(SUBSTRING_INDEX(CONCAT(TRIM(class_level), '-'), '-', 1) AS UNSIGNED)), class_level",
                    [level_tr]
                )
                for row in cur.fetchall():
                    cl = (row[0] or '').strip()
                    if cl:
                        classes.append({'value': cl, 'label': cl})
        except Exception as e:
            logger.exception('QuestionClassesView: %s', e)
            return Response({'classes': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'classes': classes}, status=status.HTTP_200_OK)


def _parse_groups_column(raw):
    """Parse groups column value: JSON array e.g. [\"Science\", \"Humanities\", \"Business Studies\"] or fallback to comma/split."""
    if not raw or not (raw := str(raw).strip()):
        return []
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            return [str(x).strip() for x in parsed if str(x).strip()]
        if isinstance(parsed, str):
            return [p.strip() for p in parsed.replace('，', ',').split(',') if p.strip()]
        return []
    except (TypeError, ValueError, json.JSONDecodeError):
        return [p.strip() for p in raw.replace('，', ',').split(',') if p.strip()]


class QuestionGroupsView(APIView):
    """
    GET: Distinct groups from cheradip_subject.groups for level_tr and class_level.
    groups column is JSON array e.g. [\"Science\", \"Humanities\", \"Business Studies\"]; returns distinct values.
    Returns [] if no non-empty groups for that level/class.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        class_level = (request.query_params.get('class_level') or '').strip()
        if not level_tr:
            return Response({'groups': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'groups': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        seen = set()
        try:
            with conn.cursor() as cur:
                if class_level:
                    cur.execute(
                        "SELECT groups FROM cheradip_subject WHERE level_tr = %s AND class_level = %s AND groups IS NOT NULL AND TRIM(COALESCE(groups, '')) != ''",
                        [level_tr, class_level]
                    )
                else:
                    cur.execute(
                        "SELECT groups FROM cheradip_subject WHERE level_tr = %s AND groups IS NOT NULL AND TRIM(COALESCE(groups, '')) != ''",
                        [level_tr]
                    )
                for row in cur.fetchall():
                    for p in _parse_groups_column(row[0]):
                        if p:
                            seen.add(p)
        except Exception as e:
            logger.exception('QuestionGroupsView: %s', e)
            return Response({'groups': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'groups': sorted(seen)}, status=status.HTTP_200_OK)


class QuestionSubjectsView(APIView):
    """
    GET: Subjects from cheradip_subject in cheradip_hsc.
    Query params: level_tr (required), class_level (optional), group (optional).
    If group given, filter by FIND_IN_SET(group, groups) so only subjects for that group.
    Each item includes subject_tr, subject_name, subject_code (from DB), name (display: subject_name or subject_tr), sq.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        class_level = (request.query_params.get('class_level') or '').strip()
        group = (request.query_params.get('group') or '').strip()
        if not level_tr:
            return Response({'subjects': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'subjects': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        subjects = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_subject' AND column_name = 'sq'"
                )
                has_sq = cur.fetchone() is not None
                if has_sq:
                    sql = (
                        "SELECT level_tr, class_level, subject_tr, COALESCE(MAX(sq), 30) AS sq, "
                        "MAX(subject_name) AS subject_name, MAX(subject_code) AS subject_code FROM cheradip_subject "
                        "WHERE level_tr = %s AND subject_tr IS NOT NULL AND TRIM(COALESCE(subject_tr, '')) != '' "
                    )
                else:
                    sql = (
                        "SELECT level_tr, class_level, subject_tr, subject_name, subject_code FROM cheradip_subject "
                        "WHERE level_tr = %s AND subject_tr IS NOT NULL AND TRIM(COALESCE(subject_tr, '')) != '' "
                    )
                params = [level_tr]
                if class_level:
                    sql += " AND class_level = %s "
                    params.append(class_level)
                sql += " GROUP BY level_tr, class_level, subject_tr ORDER BY subject_tr " if has_sq else " ORDER BY subject_tr "
                cur.execute(sql, params)
                for row in cur.fetchall():
                    lt = (row[0] or '').strip()
                    cl = (row[1] or '').strip()
                    st = (row[2] or '').strip()
                    if has_sq:
                        sq_val = row[3] if len(row) > 3 and row[3] is not None else 30
                        subject_name = (row[4] or '').strip() if len(row) > 4 and row[4] is not None else ''
                        subject_code = (row[5] or '').strip() if len(row) > 5 and row[5] is not None else ''
                    else:
                        sq_val = 30
                        subject_name = (row[3] or '').strip() if len(row) > 3 and row[3] is not None else ''
                        subject_code = (row[4] or '').strip() if len(row) > 4 and row[4] is not None else ''
                    if not st:
                        continue
                    if group:
                        # Need to check groups column; re-query with groups or do in Python
                        pass
                    display_name = subject_name or st
                    subjects.append({
                        'level_tr': lt,
                        'class_level': cl,
                        'subject_tr': st,
                        'id': st,
                        'name': display_name,
                        'subject_name': subject_name,
                        'subject_code': subject_code,
                        'sq': int(sq_val) if sq_val is not None else 30,
                    })
                if group:
                    # Filter by group: groups column is JSON array; include subject only if selected group is in that array
                    cur.execute(
                        "SELECT class_level, subject_tr, groups FROM cheradip_subject "
                        "WHERE level_tr = %s AND subject_tr IS NOT NULL AND TRIM(COALESCE(subject_tr, '')) != '' " + (
                            " AND class_level = %s " if class_level else ""
                        ),
                        [level_tr] + ([class_level] if class_level else [])
                    )
                    group_lower = group.strip().lower()
                    allowed = set()
                    for row in cur.fetchall():
                        cl_r = (row[0] or '').strip()
                        st_r = (row[1] or '').strip()
                        gr = row[2]
                        if not st_r:
                            continue
                        parts = _parse_groups_column(gr)
                        if any(p.strip().lower() == group_lower for p in parts):
                            allowed.add((cl_r, st_r))
                    subjects = [s for s in subjects if (s['class_level'], s['subject_tr']) in allowed]
        except Exception as e:
            logger.exception('QuestionSubjectsView: %s', e)
            return Response({'subjects': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'subjects': subjects}, status=status.HTTP_200_OK)


class ExamSetListView(APIView):
    """
    GET: List exam sets from cheradip_exam_set (hsc) for regularexam page.
    Query params: level_tr, class_level, subject_tr (optional). Returns id, exam_type, set_key, name_label only (no qids).
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        """Returns all exam sets (level_tr, class_level, subject_tr included for client-side filtering)."""
        if 'hsc' not in connections:
            return Response({'exam_sets': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        exam_sets = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = 'cheradip_exam_set'"
                )
                if not cur.fetchone():
                    return Response({'exam_sets': []}, status=status.HTTP_200_OK)
                sql = (
                    "SELECT id, exam_type, set_key, name_label, level_tr, class_level, subject_tr FROM cheradip_exam_set WHERE db_alias = 'hsc' "
                    "ORDER BY exam_type, set_key "
                )
                cur.execute(sql)
                for row in cur.fetchall() or []:
                    exam_sets.append({
                        'id': row[0],
                        'exam_type': row[1] or '',
                        'set_key': row[2] or '',
                        'name_label': row[3] or '',
                        'level_tr': row[4] or '',
                        'class_level': row[5] or '',
                        'subject_tr': row[6] or '',
                    })
        except Exception as e:
            logger.exception('ExamSetListView: %s', e)
            return Response({'exam_sets': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'exam_sets': exam_sets}, status=status.HTTP_200_OK)


class ExamSetDetailView(APIView):
    """
    GET: Single exam set by id from cheradip_exam_set (hsc). Returns id, name_label, set_key, exam_type for session header.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request, pk):
        if 'hsc' not in connections:
            return Response({'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT id, exam_type, set_key, name_label, level_tr, class_level, subject_tr, qids_json FROM cheradip_exam_set WHERE id = %s AND db_alias = 'hsc'",
                    [pk]
                )
                row = cur.fetchone()
                if not row:
                    return Response({'error': 'Exam set not found'}, status=status.HTTP_404_NOT_FOUND)
                qids_json = row[7] if len(row) > 7 else None
                return Response({
                    'id': row[0],
                    'exam_type': row[1] or '',
                    'set_key': row[2] or '',
                    'name_label': row[3] or '',
                    'level_tr': row[4] or '',
                    'class_level': row[5] or '',
                    'subject_tr': row[6] or '',
                    'qids_json': qids_json or '',
                }, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('ExamSetDetailView: %s', e)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ExamSetQuestionsView(APIView):
    """
    GET: Questions for an exam set (by id). Reads qids_json from cheradip_exam_set, fetches from subject table.
    Returns up to 30 questions with qid, question, option_1..4, answer for client-side exam session.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request, pk):
        import json
        if 'hsc' not in connections:
            return Response({'questions': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT qids_json, level_tr, class_level, subject_tr FROM cheradip_exam_set WHERE id = %s AND db_alias = 'hsc'",
                    [pk]
                )
                row = cur.fetchone()
                if not row:
                    return Response({'questions': [], 'error': 'Exam set not found'}, status=status.HTTP_404_NOT_FOUND)
                qids_json, level_tr, class_level, subject_tr = row[0], row[1] or '', row[2] or '', row[3] or ''
                try:
                    qids = json.loads(qids_json) if qids_json else []
                except Exception:
                    qids = []
                if not qids:
                    return Response({'questions': []}, status=status.HTTP_200_OK)
                table_name = subject_question_table_name(level_tr, class_level, subject_tr)
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = %s",
                    [table_name]
                )
                if not cur.fetchone():
                    return Response({'questions': [], 'error': 'Subject table not found'}, status=status.HTTP_200_OK)
                limit = min(30, len(qids))
                qids_slice = qids[:limit]
                placeholders = ', '.join(['%s'] * len(qids_slice))
                tbl = table_name.replace('`', '``')
                cur.execute(
                    "SELECT qid, question, option_1, option_2, option_3, option_4, answer FROM `" + tbl + "` WHERE qid IN (" + placeholders + ")",
                    qids_slice
                )
                order_map = {qid: i for i, qid in enumerate(qids_slice)}
                rows = list(cur.fetchall() or [])
                rows.sort(key=lambda r: order_map.get(r[0], 999))
                questions = []
                for r in rows:
                    questions.append({
                        'qid': r[0],
                        'id': r[0],
                        'question': r[1] or '',
                        'option_1': r[2] or '',
                        'option_2': r[3] or '',
                        'option_3': r[4] or '',
                        'option_4': r[5] or '',
                        'answer': (r[6] or '').strip() if len(r) > 6 else '',
                    })
                return Response({'questions': questions}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('ExamSetQuestionsView: %s', e)
            return Response({'questions': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class QuestionChaptersView(APIView):
    """
    GET: Unique chapters from the subject question table in cheradip_hsc.
    Table name: cheradip_(level_tr)_(class_level)_(subject_tr) from subject_question_table_name().
    Query params: level_tr, class_level, subject_tr.
    Returns chapters with id (chapter_no or chapter) and name (chapter).
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        class_level = (request.query_params.get('class_level') or '').strip()
        subject_tr = (request.query_params.get('subject_tr') or '').strip()
        if not level_tr or not class_level or not subject_tr:
            return Response({'chapters': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'chapters': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        table_name = subject_question_table_name(level_tr, class_level, subject_tr)
        conn = connections['hsc']
        chapters = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = %s",
                    [table_name]
                )
                if not cur.fetchone():
                    return Response({'chapters': [], 'error': 'Table not found'}, status=status.HTTP_200_OK)
                cur.execute(
                    "SELECT DISTINCT chapter_no, chapter FROM `{}` "
                    "WHERE (chapter_no IS NOT NULL AND TRIM(COALESCE(chapter_no, '')) != '') "
                    "   OR (chapter IS NOT NULL AND TRIM(COALESCE(chapter, '')) != '') "
                    "ORDER BY CAST(COALESCE(NULLIF(TRIM(chapter_no), ''), '0') AS UNSIGNED), chapter_no, chapter".format(table_name)
                )
                seen = set()
                for row in cur.fetchall():
                    ch_no = (row[0] or '').strip()
                    ch = (row[1] or '').strip()
                    name = ch or ch_no or ''
                    if not name:
                        continue
                    key = (ch_no, ch)
                    if key in seen:
                        continue
                    seen.add(key)
                    chapters.append({'id': ch_no or name, 'name': name, 'chapter_no': ch_no or None})
        except Exception as e:
            logger.exception('QuestionChaptersView: %s', e)
            return Response({'chapters': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'chapters': chapters}, status=status.HTTP_200_OK)


class QuestionTopicsView(APIView):
    """
    GET: Unique topics from the subject question table in cheradip_hsc.
    Query params: level_tr, class_level, subject_tr; optional chapter (chapter_no or chapter name) to filter.
    Ordered by topic ascending (character order).
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        class_level = (request.query_params.get('class_level') or '').strip()
        subject_tr = (request.query_params.get('subject_tr') or '').strip()
        chapter = (request.query_params.get('chapter') or '').strip()
        if not level_tr or not class_level or not subject_tr:
            return Response({'topics': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'topics': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        table_name = subject_question_table_name(level_tr, class_level, subject_tr)
        conn = connections['hsc']
        topics = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = %s",
                    [table_name]
                )
                if not cur.fetchone():
                    return Response({'topics': []}, status=status.HTTP_200_OK)
                cur.execute(
                    "SELECT COLUMN_NAME FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = %s AND COLUMN_NAME = 'topic_no'",
                    [table_name]
                )
                has_topic_no = cur.fetchone() is not None
                cur.execute(
                    "SELECT COLUMN_NAME FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = %s AND COLUMN_NAME = 'chapter_no'",
                    [table_name]
                )
                has_chapter_no = cur.fetchone() is not None
                if has_topic_no and has_chapter_no and chapter:
                    cur.execute(
                        "SELECT DISTINCT chapter_no, topic_no, topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "AND (chapter_no = %s OR chapter = %s) "
                        "ORDER BY CAST(COALESCE(NULLIF(TRIM(chapter_no), ''), '0') AS UNSIGNED), CAST(COALESCE(NULLIF(TRIM(topic_no), ''), '0') AS UNSIGNED), topic".format(table_name),
                        [chapter, chapter]
                    )
                    for row in cur.fetchall():
                        ch_no, tno, t = (row[0] or '').strip(), (row[1] or '').strip(), (row[2] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t, 'topic_no': tno or None, 'chapter_no': ch_no or None})
                elif has_topic_no and has_chapter_no:
                    cur.execute(
                        "SELECT DISTINCT chapter_no, topic_no, topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "ORDER BY CAST(COALESCE(NULLIF(TRIM(chapter_no), ''), '0') AS UNSIGNED), CAST(COALESCE(NULLIF(TRIM(topic_no), ''), '0') AS UNSIGNED), topic".format(table_name)
                    )
                    for row in cur.fetchall():
                        ch_no, tno, t = (row[0] or '').strip(), (row[1] or '').strip(), (row[2] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t, 'topic_no': tno or None, 'chapter_no': ch_no or None})
                elif has_topic_no and chapter:
                    cur.execute(
                        "SELECT DISTINCT topic_no, topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "AND (chapter_no = %s OR chapter = %s) "
                        "ORDER BY CAST(COALESCE(NULLIF(TRIM(topic_no), ''), '0') AS UNSIGNED), topic_no, topic".format(table_name),
                        [chapter, chapter]
                    )
                    for row in cur.fetchall():
                        tno, t = (row[0] or '').strip(), (row[1] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t, 'topic_no': tno or None})
                elif has_topic_no:
                    cur.execute(
                        "SELECT DISTINCT topic_no, topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "ORDER BY CAST(COALESCE(NULLIF(TRIM(topic_no), ''), '0') AS UNSIGNED), topic_no, topic".format(table_name)
                    )
                    for row in cur.fetchall():
                        tno, t = (row[0] or '').strip(), (row[1] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t, 'topic_no': tno or None})
                elif chapter:
                    cur.execute(
                        "SELECT DISTINCT topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "AND (chapter_no = %s OR chapter = %s) ORDER BY topic".format(table_name),
                        [chapter, chapter]
                    )
                    for row in cur.fetchall():
                        t = (row[0] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t})
                else:
                    cur.execute(
                        "SELECT DISTINCT topic FROM `{}` "
                        "WHERE topic IS NOT NULL AND TRIM(COALESCE(topic, '')) != '' "
                        "ORDER BY topic".format(table_name)
                    )
                    for row in cur.fetchall():
                        t = (row[0] or '').strip()
                        if t:
                            topics.append({'id': t, 'name': t})
        except Exception as e:
            logger.exception('QuestionTopicsView: %s', e)
            return Response({'topics': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'topics': topics}, status=status.HTTP_200_OK)


class CheradipSourceListView(APIView):
    """
    GET: List institutes from cheradip_source (institute_code, institute_name, institute_type).
    Reads from HSC database. Used by More Filters to show institute_type dropdown and filter sources.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        if 'hsc' not in connections:
            return Response({'sources': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        conn = connections['hsc']
        db_name = conn.settings_dict.get('NAME', '')
        sources = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = %s AND table_name = %s",
                    [db_name, 'cheradip_source']
                )
                if not cur.fetchone():
                    return Response({'sources': []}, status=status.HTTP_200_OK)
                cur.execute(
                    "SELECT institute_code, institute_name, institute_type FROM cheradip_source ORDER BY institute_type, institute_code"
                )
                for row in cur.fetchall():
                    sources.append({
                        'institute_code': row[0].strip() if row[0] else '',
                        'institute_name': row[1].strip() if row[1] else '',
                        'institute_type': row[2].strip() if row[2] else '',
                    })
        except Exception as e:
            logger.exception('CheradipSourceListView: %s', e)
            return Response({'sources': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'sources': sources}, status=status.HTTP_200_OK)


class QuestionListView(APIView):
    """
    GET: List questions from the subject question table in cheradip_hsc, filtered by topic (and optional chapter).
    Query params: level_tr, class_level, subject_tr, topic; optional chapter.
    Returns questions with id, question, option_1..4, answer, chapter_no, chapter, topic, etc. for user to select.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def get(self, request):
        level_tr = (request.query_params.get('level_tr') or '').strip()
        class_level = (request.query_params.get('class_level') or '').strip()
        subject_tr = (request.query_params.get('subject_tr') or '').strip()
        topic = (request.query_params.get('topic') or '').strip()
        chapter = (request.query_params.get('chapter') or '').strip()
        if not level_tr or not class_level or not subject_tr or not topic:
            return Response({'questions': []}, status=status.HTTP_200_OK)
        if 'hsc' not in connections:
            return Response({'questions': [], 'error': 'HSC database not configured'}, status=status.HTTP_200_OK)
        table_name = subject_question_table_name(level_tr, class_level, subject_tr)
        conn = connections['hsc']
        questions = []
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = %s",
                    [table_name]
                )
                if not cur.fetchone():
                    return Response({'questions': []}, status=status.HTTP_200_OK)
                cur.execute(
                    "SELECT COLUMN_NAME FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = %s AND COLUMN_NAME IN ('qid', 'id', 'topic_no', 'subsource', 'explanation2', 'explanation3', 'level')",
                    [table_name]
                )
                col_set = {r[0] for r in cur.fetchall()}
                pk_col = 'qid' if 'qid' in col_set else 'id'
                mid = "chapter_no, chapter, topic_no, topic" if 'topic_no' in col_set else "chapter_no, chapter, topic"
                subsource_col = ", subsource" if 'subsource' in col_set else ""
                expl2_col = ", explanation2" if 'explanation2' in col_set else ""
                expl3_col = ", explanation3" if 'explanation3' in col_set else ""
                level_col = ", level" if 'level' in col_set else ""
                select_cols = f"{pk_col}, subject, {mid}, question, option_1, option_2, option_3, option_4, answer, explanation{expl2_col}{expl3_col}, type{level_col}{subsource_col}"
                if chapter:
                    cur.execute(
                        "SELECT {} FROM `{}` WHERE topic = %s AND (chapter_no = %s OR chapter = %s) ORDER BY {}".format(select_cols, table_name, pk_col),
                        [topic, chapter, chapter]
                    )
                else:
                    cur.execute(
                        "SELECT {} FROM `{}` WHERE topic = %s ORDER BY {}".format(select_cols, table_name, pk_col),
                        [topic]
                    )
                cols = [c[0] for c in cur.description]
                for row in cur.fetchall():
                    q = {}
                    for k, v in zip(cols, row):
                        q[k] = v.strip() if v is not None and isinstance(v, str) else v
                    questions.append(q)
        except Exception as e:
            logger.exception('QuestionListView: %s', e)
            return Response({'questions': [], 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return Response({'questions': questions}, status=status.HTTP_200_OK)


class PendingQuestionRequestView(APIView):
    """
    POST: Submit an edit request for an existing question. Body: qid, question, option_1..4, type,
    level_tr, class_level, subject_tr, chapter, topic, etc.
    Writes into cheradip_pending_question_request in the same DB as question_list (HSC).
    Requires Bearer token (CustomerToken) from login.
    """
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data or {}
        if not (data.get('question') or '').strip():
            return Response({'error': 'Missing or empty: question'}, status=status.HTTP_400_BAD_REQUEST)
        if 'hsc' not in connections:
            return Response({'error': 'HSC database not configured'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
        conn = connections['hsc']
        try:
            from django.utils import timezone
            now = timezone.now()
            raw_qid = data.get('qid')
            requested_qid = str(raw_qid).strip() if raw_qid is not None else ''
            requested_qid = requested_qid or None
            status_val = (data.get('status') or 'Update').strip()[:20] or 'Update'
            level_tr = (data.get('level_tr') or '')[:100]
            class_level = (data.get('class_level') or '')[:50]
            subject_tr = (data.get('subject_tr') or '').strip()[:255]
            if not subject_tr:
                subject_tr = ''
            table_val = (data.get('table') or '').strip()[:64] or None
            if not table_val and level_tr and class_level and subject_tr:
                from cheradip.subject_question_tables import subject_question_table_name
                table_val = subject_question_table_name(level_tr, class_level, subject_tr)
            chapter_no = (data.get('chapter_no') or '')[:50]
            chapter = (data.get('chapter') or '').strip()[:255]
            topic_no = (data.get('topic_no') or '')[:50]
            topic = (data.get('topic') or '').strip()[:255]
            question = (data.get('question') or '').strip()
            option_1 = (data.get('option_1') or '')[:500]
            option_2 = (data.get('option_2') or '')[:500]
            option_3 = (data.get('option_3') or '')[:500]
            option_4 = (data.get('option_4') or '')[:500]
            answer = (data.get('answer') or '')[:500]
            explanation = (data.get('explanation') or '')[:50000]
            explanation2 = (data.get('explanation2') or '')[:50000]
            explanation3 = (data.get('explanation3') or '')[:50000]
            type_val = (data.get('type') or '')[:100]
            level_val = (data.get('level') or '').strip()[:100] or None
            subsource_val = (data.get('subsource') or '').strip()[:255] or None
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_pending_question_request' AND column_name = 'requested_qid'"
                )
                has_requested_qid = cur.fetchone() is not None
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_pending_question_request' AND column_name = 'subsource'"
                )
                has_subsource = cur.fetchone() is not None
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_pending_question_request' AND column_name = 'level'"
                )
                has_level = cur.fetchone() is not None
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_pending_question_request' AND column_name = 'qid'"
                )
                has_qid_col = cur.fetchone() is not None
                cur.execute(
                    "SELECT 1 FROM information_schema.columns WHERE table_schema = DATABASE() AND table_name = 'cheradip_pending_question_request' AND column_name = 'table'"
                )
                has_table_col = cur.fetchone() is not None
                if has_requested_qid and requested_qid is not None:
                    if has_level and has_subsource and has_qid_col:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at, requested_qid, qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now, requested_qid, requested_qid]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at, requested_qid, qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now, requested_qid, requested_qid]
                            )
                    elif has_level and has_subsource:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at, requested_qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now, requested_qid]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at, requested_qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now, requested_qid]
                            )
                    elif has_qid_col:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at, requested_qid, qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now, requested_qid, requested_qid]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at, requested_qid, qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now, requested_qid, requested_qid]
                            )
                    else:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at, requested_qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now, requested_qid]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at, requested_qid)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now, requested_qid]
                            )
                else:
                    if has_level and has_subsource:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, level, subsource, status, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, level_val, subsource_val, status_val, now]
                            )
                    else:
                        if has_table_col:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (`table`, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [table_val, level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now]
                            )
                        else:
                            cur.execute(
                                """
                                INSERT INTO cheradip_pending_question_request
                                (level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type, status, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """,
                                [level_tr, class_level, subject_tr, chapter_no, chapter, topic_no, topic, question,
                                 option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3,
                                 type_val, status_val, now]
                            )
                pk = cur.lastrowid
            return Response({'id': pk, 'status': status_val, 'message': 'Edit request submitted.'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            logger.exception('PendingQuestionRequestView: %s', e)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PendingQuestionSubmitView(APIView):
    """
    POST: Submit a new question for approval. Body: level_tr, class_level, subject_tr, chapter_no, chapter,
    topic_no, topic, question, option_1..4, answer, explanation, type, etc.
    Creates a PendingQuestion with status=pending. When approved, it will be inserted into the HSC subject
    question table with qid = chapter_no_topic_no_0001, 0002, ...
    Requires Bearer token (CustomerToken) from login.
    """
    authentication_classes = [BearerTokenAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data or {}
        required = ['subject_tr', 'chapter', 'topic', 'question']
        for k in required:
            if not (data.get(k) or '').strip():
                return Response({'error': f'Missing or empty: {k}'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            obj = PendingQuestion.objects.create(
                level_tr=(data.get('level_tr') or '').strip(),
                class_level=(data.get('class_level') or '').strip(),
                subject_tr=(data.get('subject_tr') or '').strip(),
                chapter_no=(data.get('chapter_no') or '').strip(),
                chapter=(data.get('chapter') or '').strip(),
                topic_no=(data.get('topic_no') or '').strip(),
                topic=(data.get('topic') or '').strip(),
                question=(data.get('question') or '').strip(),
                option_1=(data.get('option_1') or '').strip()[:500],
                option_2=(data.get('option_2') or '').strip()[:500],
                option_3=(data.get('option_3') or '').strip()[:500],
                option_4=(data.get('option_4') or '').strip()[:500],
                answer=(data.get('answer') or '').strip()[:500],
                explanation=(data.get('explanation') or '')[:50000],
                explanation2=(data.get('explanation2') or '')[:50000],
                explanation3=(data.get('explanation3') or '')[:50000],
                type=(data.get('type') or '').strip()[:100],
                status=PendingQuestion.STATUS_PENDING,
            )
            return Response({'id': obj.id, 'status': obj.status, 'message': 'Question submitted for approval.'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            logger.exception('PendingQuestionSubmitView: %s', e)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PendingQuestionApproveView(APIView):
    """
    POST: Approve a pending question by id. Inserts it into the HSC subject question table with generated qid
    (chapter_no_topic_no_0001, ...) at the last position under that topic, then marks the PendingQuestion as approved.
    Body: id (required). Optional: approved_by for logging.
    """
    permission_classes = [PublicAccess]
    authentication_classes = []

    def post(self, request):
        pk = request.data.get('id') or request.query_params.get('id')
        if pk is None:
            return Response({'error': 'id is required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            pending = PendingQuestion.objects.get(pk=pk, status=PendingQuestion.STATUS_PENDING)
        except PendingQuestion.DoesNotExist:
            return Response({'error': 'Pending question not found or already processed'}, status=status.HTTP_404_NOT_FOUND)
        if 'hsc' not in connections:
            return Response({'error': 'HSC database not configured'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        table_name = subject_question_table_name(
            pending.level_tr or '',
            pending.class_level or '',
            pending.subject_tr or ''
        )
        qid = next_qid_for_chapter_topic(
            table_name,
            pending.chapter_no or '0',
            pending.topic_no or '0',
            using='hsc'
        )
        conn = connections['hsc']
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = %s",
                    [table_name]
                )
                if not cur.fetchone():
                    return Response({'error': f'Subject question table not found: {table_name}'}, status=status.HTTP_404_NOT_FOUND)
                cur.execute(
                    """INSERT INTO `{}` (qid, subject, chapter_no, chapter, topic_no, topic, question, option_1, option_2, option_3, option_4, answer, explanation, explanation2, explanation3, type, created_at, updated_at, updated_by)
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(6), NOW(6), %s)""".format(table_name),
                    [
                        qid, pending.subject_tr, pending.chapter_no or None, pending.chapter, pending.topic_no or None, pending.topic,
                        pending.question, pending.option_1 or None, pending.option_2 or None, pending.option_3 or None, pending.option_4 or None,
                        pending.answer or None, pending.explanation or None, pending.explanation2 or None, pending.explanation3 or None,
                        pending.type or None, request.data.get('approved_by') or ''
                    ]
                )
        except Exception as e:
            logger.exception('PendingQuestionApproveView insert: %s', e)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        from django.utils import timezone
        pending.status = PendingQuestion.STATUS_APPROVED
        pending.approved_at = timezone.now()
        pending.approved_qid = qid
        pending.save(update_fields=['status', 'approved_at', 'approved_qid'])
        return Response({'qid': qid, 'message': 'Question approved and added.'}, status=status.HTTP_200_OK)
